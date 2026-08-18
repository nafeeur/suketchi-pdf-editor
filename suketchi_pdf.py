"""
    Suketchi - A lightweight, free, and open-source PDF reader and editor.
    Copyright (C) 2026 Nafeeur Rahman
    
    This program is free software: you can redistribute it and/or modify
    it under the terms of the GNU General Public License as published by
    the Free Software Foundation, either version 3 of the License, or
    (at your option) any later version.

    This program is distributed in the hope that it will be useful,
    but WITHOUT ANY WARRANTY; without even the implied warranty of
    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
    GNU General Public License for more details.

    You should have received a copy of the GNU General Public License
    along with this program.  If not, see <https://www.gnu.org/licenses/>.
"""

from __future__ import annotations

import html
import json
import math
import os
import re
import sys
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

try:
    import pymupdf as fitz
except ImportError:
    import fitz
try:
    fitz.TOOLS.mupdf_display_errors(False)
except Exception:
    pass
try:
    fitz.TOOLS.mupdf_display_warnings(False)
except Exception:
    pass

import random

from PyQt6.QtCore import QBuffer, QIODevice, QObject, QPointF, QRectF, QSize, Qt, QThread, pyqtSignal, pyqtSlot
from PyQt6.QtGui import (
    QAction,
    QActionGroup,
    QColor,
    QFont,
    QFontDatabase,
    QFontMetrics,
    QGuiApplication,
    QImage,
    QIcon,
    QPainter,
    QBrush,
    QPainterPath,
    QPalette,
    QPen,
    QPixmap,
)
from PyQt6.QtWidgets import (
    QApplication,
    QProxyStyle,
    QStyle,
    QStyleOptionMenuItem,
    QStyleOptionToolButton,
    QStyleOptionViewItem,
    QCheckBox,
    QColorDialog,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QFrame,
    QGridLayout,
    QGroupBox,
    QHBoxLayout,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMenu,
    QMessageBox,
    QProgressBar,
    QPushButton,
    QScrollArea,
    QSizePolicy,
    QSpinBox,
    QSplitter,
    QStatusBar,
    QTabWidget,
    QTextEdit,
    QToolBar,
    QToolButton,
    QVBoxLayout,
    QWidget,
)


def rgb_from_qcolor(color: QColor) -> Tuple[float, float, float]:
    return color.redF(), color.greenF(), color.blueF()


def qcolor_from_pdf_int(value: int) -> QColor:
    """PyMuPDF spans return color as 0xRRGGBB integer."""
    r = (value >> 16) & 255
    g = (value >> 8) & 255
    b = value & 255
    return QColor(r, g, b)


def safe_filename(text: str, fallback: str = "export") -> str:
    cleaned = "".join(c if c.isalnum() or c in ("-", "_", ".") else "_" for c in text.strip())
    return cleaned or fallback


def clamp(value: float, low: float, high: float) -> float:
    return max(low, min(high, value))


def page_range_from_user(parent: QWidget, page_count: int) -> Optional[Tuple[int, int]]:
    text, ok = QInputDialog.getText(
        parent,
        "Page Range",
        f"Enter page range, for example 1-{page_count} or 3-7:",
    )
    if not ok or not text.strip():
        return None

    value = text.strip().replace(" ", "")
    try:
        if "-" in value:
            a, b = value.split("-", 1)
            start = int(a)
            end = int(b)
        else:
            start = end = int(value)
    except ValueError:
        QMessageBox.warning(parent, "Invalid Range", "Please enter a valid page range.")
        return None

    if start < 1 or end < start or end > page_count:
        QMessageBox.warning(parent, "Invalid Range", f"Range must be between 1 and {page_count}.")
        return None

    return start - 1, end - 1


INK = QColor(0, 0, 0)
PAPER = QColor(255, 255, 255)

APP_FONT_FAMILY = "Segoe UI"

GREY_MUTED = QColor(120, 120, 120)
GREY_FAINT = QColor(175, 175, 175)

# Hover fill used by the sketch style / tab bar. Kept as a module-level colour
# so dark mode can flip it alongside INK / PAPER.
HOVER = QColor(238, 238, 238)

# Theme palettes. Dark mode simply inverts ink and paper (and picks a hover /
# grey that reads well on a dark background) so the entire hand-drawn UI flips
# while keeping its sketch character.
_THEMES = {
    "light": {
        "ink": QColor(0, 0, 0),
        "paper": QColor(255, 255, 255),
        "hover": QColor(238, 238, 238),
        "grey_muted": QColor(120, 120, 120),
        "grey_faint": QColor(175, 175, 175),
    },
    "dark": {
        "ink": QColor(233, 233, 233),
        "paper": QColor(24, 24, 26),
        "hover": QColor(52, 52, 56),
        "grey_muted": QColor(150, 150, 150),
        "grey_faint": QColor(110, 110, 110),
    },
}

CURRENT_THEME = "light"


def apply_theme_colors(name: str):
    """Mutate the module-level colour objects in place so every painter that
    already holds a reference to INK / PAPER / etc. picks up the new theme."""
    global CURRENT_THEME
    theme = _THEMES.get(name, _THEMES["light"])
    CURRENT_THEME = name
    for target, key in (
        (INK, "ink"), (PAPER, "paper"), (HOVER, "hover"),
        (GREY_MUTED, "grey_muted"), (GREY_FAINT, "grey_faint"),
    ):
        c = theme[key]
        target.setRgb(c.red(), c.green(), c.blue())



def sketch_path(rect: QRectF, seed: int, jitter: float = 1.3, segments: int = 5) -> QPainterPath:
    """A wobbly rectangle that looks drawn by hand.

    Each edge becomes a short polyline whose points are nudged by a seeded
    random offset, then smoothed with quadratic curves so the line reads as a
    pen stroke rather than a jagged zigzag. The seed is derived from the widget
    and its geometry, so a control keeps the *same* wobble across repaints
    instead of shimmering every frame.
    """
    rnd = random.Random(seed)
    corners = [
        (rect.left(), rect.top()),
        (rect.right(), rect.top()),
        (rect.right(), rect.bottom()),
        (rect.left(), rect.bottom()),
    ]
    points = []
    for index in range(4):
        x0, y0 = corners[index]
        x1, y1 = corners[(index + 1) % 4]
        # Segment count scales with edge length. With a fixed count, long
        length = max(abs(x1 - x0), abs(y1 - y0))
        edge_segments = max(segments, int(length / 26))
        for step in range(edge_segments):
            t = step / edge_segments
            points.append((
                x0 + (x1 - x0) * t + rnd.uniform(-jitter, jitter),
                y0 + (y1 - y0) * t + rnd.uniform(-jitter, jitter),
            ))

    path = QPainterPath()
    path.moveTo(*points[0])
    for index in range(1, len(points)):
        px, py = points[index - 1]
        x, y = points[index]
        path.quadTo(px, py, (px + x) / 2, (py + y) / 2)
    path.closeSubpath()
    return path


def sketch_line(p1: QPointF, p2: QPointF, seed: int, jitter: float = 1.0) -> QPainterPath:
    """A single wobbly stroke between two points."""
    rnd = random.Random(seed)
    path = QPainterPath()
    path.moveTo(p1)
    steps = 6
    for step in range(1, steps + 1):
        t = step / steps
        x = p1.x() + (p2.x() - p1.x()) * t + rnd.uniform(-jitter, jitter)
        y = p1.y() + (p2.y() - p1.y()) * t + rnd.uniform(-jitter, jitter)
        path.lineTo(x, y)
    return path


def stable_seed(*parts) -> int:
    """Deterministic seed so a control's wobble never changes between frames."""
    return abs(hash(parts)) % (2 ** 31)


def draw_sketch_box(
    painter: QPainter,
    rect: QRectF,
    seed: int,
    fill: Optional[QColor] = None,
    width: float = 2.0,
    shadow: bool = False,
    shadow_offset: float = 4.0,
    jitter: float = 1.3,
):
    """Paint one hand-drawn box, optionally with the hard offset shadow."""
    painter.save()
    painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)

    if shadow:
        painter.setPen(Qt.PenStyle.NoPen)
        painter.setBrush(INK)
        painter.drawPath(sketch_path(rect.translated(shadow_offset, shadow_offset), seed, jitter))

    painter.setBrush(fill if fill is not None else Qt.BrushStyle.NoBrush)
    painter.setPen(QPen(INK, width, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
    painter.drawPath(sketch_path(rect, seed, jitter))
    painter.restore()


def render_signature_image(
    text: str,
    family: str,
    color: QColor,
    point_size: int = 96,
    scale: int = 3,
) -> QImage:
    """Render a typed signature to a transparent, tightly-cropped QImage.

    Retained as a utility (e.g. for previews / exports). Placed signatures are
    inserted as real embedded text rather than an image, so they stay
    searchable and movable.
    """
    text = (text or "").strip() or " "
    font = QFont(family)
    font.setPointSize(point_size * scale)
    font.setStyleStrategy(QFont.StyleStrategy.PreferAntialias)

    metrics = QFontMetrics(font)
    rect = metrics.boundingRect(text)
    pad = int(point_size * scale * 0.30)
    w = max(rect.width() + pad * 2, 1)
    h = max(rect.height() + pad * 2, 1)

    image = QImage(w, h, QImage.Format.Format_ARGB32)
    image.fill(Qt.GlobalColor.transparent)

    painter = QPainter(image)
    painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
    painter.setRenderHint(QPainter.RenderHint.TextAntialiasing, True)
    painter.setFont(font)
    painter.setPen(QPen(color))
    # Baseline positioning that accounts for the font's ascent/descent.
    painter.drawText(pad - rect.left(), pad - rect.top(), text)
    painter.end()
    return image


class SketchMenu(QMenu):
    """A QMenu that paints its own panel and hand-drawn selection.

    QStyleSheetStyle draws CE_MenuItem itself whenever a stylesheet is active
    anywhere in the widget's ancestry, and never delegates to the application
    style. So instead of styling the item, we paint the panel and the highlight
    box underneath, and let the base implementation draw only the text on top
    (the stylesheet keeps item backgrounds transparent).

    The popup is made a translucent, frameless, shadow-less window so the
    platform does not draw its own faint rectangular panel / drop shadow behind
    our hand-drawn box. Only the wobbly sketch box (filled with paper) shows.
    """

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setWindowFlags(
            self.windowFlags()
            | Qt.WindowType.FramelessWindowHint
            | Qt.WindowType.NoDropShadowWindowHint
        )
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        rect = QRectF(self.rect()).adjusted(1.5, 1.5, -1.5, -1.5)
        # Do NOT fill the whole rect: the window is translucent, so anything
        # outside the sketch box stays fully transparent. The box itself is
        # filled with paper, which supplies the opaque menu background.
        draw_sketch_box(painter, rect, stable_seed("menu", self.width(), self.height()),
                        fill=PAPER, width=2.0, jitter=1.0)

        action = self.activeAction()
        if action is not None and action.isEnabled() and not action.isSeparator():
            geo = QRectF(self.actionGeometry(action)).adjusted(3, 1, -3, -1)
            if geo.width() > 4 and geo.height() > 4:
                draw_sketch_box(painter, geo,
                                stable_seed("menusel", int(geo.width()), int(geo.y())),
                                fill=None, width=1.8, jitter=0.9)
        painter.end()
        super().paintEvent(event)


class SketchSwatchButton(QPushButton):
    """A colour-swatch button drawn as a hand-drawn sketch box.

    The Stroke / Fill / Text colour buttons used to be styled with a plain
    stylesheet border (a flat rectangle), which clashed with the hand-drawn
    look of the rest of the UI. This button instead paints a wobbly sketch box
    filled with the current swatch colour, matching every other control.
    """

    def __init__(self, text: str = "", parent=None):
        super().__init__(text, parent)
        self._swatch = QColor("#ffffff")
        # No stylesheet border: we paint the box ourselves so the proxy style
        # is not overridden by QSS box rendering.
        self.setStyleSheet("background: transparent; border: none;")
        self.setMinimumHeight(30)

    def set_swatch(self, color: QColor):
        self._swatch = QColor(color)
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)

        rect = QRectF(self.rect()).adjusted(2.5, 2.5, -3.5, -3.5)
        seed = stable_seed(id(self), self.width(), self.height())

        state = self.isDown() or self.isChecked()
        hovered = self.underMouse() and self.isEnabled()
        offset = 0.0 if state else 3.0
        target = rect.translated(1.5, 1.5) if state else rect

        draw_sketch_box(
            painter, target, seed,
            fill=self._swatch,
            width=2.0 if self.isEnabled() else 1.2,
            shadow=self.isEnabled() and not state,
            shadow_offset=offset,
            jitter=1.0,
        )

        # Contrast-aware label colour so text stays readable on any swatch.
        c = self._swatch
        luminance = 0.299 * c.red() + 0.587 * c.green() + 0.114 * c.blue()
        fg = QColor("#1f2733") if luminance > 150 else QColor("#ffffff")
        painter.setPen(QPen(fg))
        painter.drawText(target, Qt.AlignmentFlag.AlignCenter, self.text())
        painter.end()


class ThemeToggle(QWidget):
    """A little hand-drawn sun / moon button that flips light and dark mode.

    Sits in the bottom-right corner. It looks like someone doodled a sun (a
    circle with wobbly rays) in light mode, and a crescent moon with a couple of
    stars in dark mode — coherent with the sketch UI.
    """

    clicked = pyqtSignal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self._dark = False
        self._hover = False
        self.setFixedSize(34, 34)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setToolTip("Toggle dark mode")

    def set_dark(self, dark: bool):
        self._dark = bool(dark)
        self.setToolTip("Switch to light mode" if self._dark else "Switch to dark mode")
        self.update()

    def enterEvent(self, event):
        self._hover = True
        self.update()

    def leaveEvent(self, event):
        self._hover = False
        self.update()

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        r = QRectF(self.rect()).adjusted(3, 3, -3, -3)
        seed = stable_seed("themetoggle", self._dark)

        # The enclosing sketch box.
        draw_sketch_box(painter, r, seed,
                        fill=HOVER if self._hover else PAPER,
                        width=1.8, jitter=0.8)

        cx, cy = r.center().x(), r.center().y()
        painter.setPen(QPen(INK, 1.8, Qt.PenStyle.SolidLine,
                            Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))

        if not self._dark:
            # --- Sun: a wobbly circle with radiating rays. ---
            rad = 5.0
            circle = QRectF(cx - rad, cy - rad, rad * 2, rad * 2)
            painter.setBrush(Qt.BrushStyle.NoBrush)
            painter.drawPath(sketch_path(circle, seed + 1, 0.6))
            ray_in, ray_out = rad + 1.8, rad + 4.6
            for k in range(8):
                ang = math.radians(k * 45 + 10)
                x1 = cx + math.cos(ang) * ray_in
                y1 = cy + math.sin(ang) * ray_in
                x2 = cx + math.cos(ang) * ray_out
                y2 = cy + math.sin(ang) * ray_out
                painter.drawPath(sketch_line(QPointF(x1, y1), QPointF(x2, y2),
                                             seed + 2 + k, 0.35))
        else:
            # --- Moon: a crescent made by subtracting an offset circle. ---
            rad = 6.6
            outer = QPainterPath()
            outer.addEllipse(QPointF(cx - 0.5, cy), rad, rad)
            inner = QPainterPath()
            inner.addEllipse(QPointF(cx + 3.0, cy - 2.2), rad, rad)
            crescent = outer.subtracted(inner)
            painter.setBrush(INK)
            painter.setPen(QPen(INK, 1.3, Qt.PenStyle.SolidLine,
                                Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
            painter.drawPath(crescent)
            # A couple of little sketch stars.
            painter.setBrush(Qt.BrushStyle.NoBrush)
            for sx, sy, s in ((cx + 6.5, cy - 6.5, 1.7), (cx + 8, cy + 3, 1.3)):
                painter.drawPath(sketch_line(QPointF(sx - s, sy), QPointF(sx + s, sy),
                                             seed + 20 + int(sx), 0.3))
                painter.drawPath(sketch_line(QPointF(sx, sy - s), QPointF(sx, sy + s),
                                             seed + 30 + int(sy), 0.3))
        painter.end()


class SketchStyle(QProxyStyle):
    """Paints Qt's standard controls as hand-drawn boxes.

    Qt stylesheets cannot draw a wobbly border, so borders and backgrounds are
    deliberately left out of the stylesheet and painted here instead. The
    stylesheet still handles colour, font and padding.
    """

    def _seed_for(self, widget, rect) -> int:
        return stable_seed(id(widget) if widget is not None else 0,
                           rect.width(), rect.height())

    def drawPrimitive(self, element, option, painter, widget=None):
        rect = QRectF(option.rect).adjusted(1.5, 1.5, -1.5, -1.5)
        seed = self._seed_for(widget, option.rect)

        if element in (
            QStyle.PrimitiveElement.PE_PanelButtonCommand,
            QStyle.PrimitiveElement.PE_PanelButtonTool,
        ):
            if rect.width() <= 2 or rect.height() <= 2:
                return
            state = option.state
            on = bool(state & QStyle.StateFlag.State_On)
            sunken = bool(state & QStyle.StateFlag.State_Sunken)
            hovered = bool(state & QStyle.StateFlag.State_MouseOver)
            enabled = bool(state & QStyle.StateFlag.State_Enabled)

            if on or sunken:
                fill = INK
            elif hovered and enabled:
                fill = HOVER
            else:
                fill = PAPER

            shadow = enabled and not (on or sunken)
            offset = 0.0 if (on or sunken) else 3.5
            target = rect.translated(1.5, 1.5) if (on or sunken) else rect
            draw_sketch_box(painter, target, seed, fill=fill,
                            width=2.0 if enabled else 1.2,
                            shadow=shadow, shadow_offset=offset)
            return

        if element in (
            QStyle.PrimitiveElement.PE_FrameLineEdit,
            QStyle.PrimitiveElement.PE_PanelLineEdit,
        ):
            if element == QStyle.PrimitiveElement.PE_PanelLineEdit:
                draw_sketch_box(painter, rect, seed, fill=PAPER, width=1.8)
            return

        if element in (
            QStyle.PrimitiveElement.PE_Frame,
            QStyle.PrimitiveElement.PE_FrameGroupBox,
            QStyle.PrimitiveElement.PE_FrameTabWidget,
        ):
            draw_sketch_box(painter, rect, seed, fill=None, width=1.6, jitter=1.0)
            return

        if element in (
            QStyle.PrimitiveElement.PE_PanelMenu,
            QStyle.PrimitiveElement.PE_FrameMenu,
        ):
            # SketchMenu is a translucent popup that paints its own wobbly,
            # paper-filled box in paintEvent. Painting an opaque rectangle here
            # would show up as a faint rectangle behind that box, so for our
            # menus we skip the panel entirely. Any other (non-sketch) menu
            # still gets a hand-drawn outline.
            if isinstance(widget, SketchMenu):
                return
            draw_sketch_box(painter, QRectF(option.rect).adjusted(1.5, 1.5, -1.5, -1.5),
                            seed, fill=PAPER, width=2.0, jitter=1.0)
            return

        if element == QStyle.PrimitiveElement.PE_IndicatorCheckBox:
            draw_sketch_box(painter, rect, seed, fill=PAPER, width=1.8, jitter=0.9)
            if option.state & QStyle.StateFlag.State_On:
                painter.save()
                painter.setPen(QPen(INK, 2.0, Qt.PenStyle.SolidLine,
                                    Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
                painter.drawPath(sketch_line(
                    QPointF(rect.left() + rect.width() * 0.22, rect.top() + rect.height() * 0.52),
                    QPointF(rect.left() + rect.width() * 0.45, rect.top() + rect.height() * 0.78),
                    seed + 1, 0.6))
                painter.drawPath(sketch_line(
                    QPointF(rect.left() + rect.width() * 0.45, rect.top() + rect.height() * 0.78),
                    QPointF(rect.left() + rect.width() * 0.82, rect.top() + rect.height() * 0.22),
                    seed + 2, 0.6))
                painter.restore()
            return

        super().drawPrimitive(element, option, painter, widget)

    def drawControl(self, element, option, painter, widget=None):
        # Qt reuses these control elements with different option structs (a
        if element == QStyle.ControlElement.CE_MenuItem and isinstance(option, QStyleOptionMenuItem):
            menu_option = QStyleOptionMenuItem(option)

            if menu_option.menuItemType == QStyleOptionMenuItem.MenuItemType.Separator:
                rect = QRectF(option.rect)
                painter.save()
                painter.setPen(QPen(INK, 1.4, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
                y = rect.center().y()
                painter.drawPath(sketch_line(
                    QPointF(rect.left() + 10, y), QPointF(rect.right() - 10, y),
                    stable_seed("sep", int(rect.width()), int(rect.y())), 0.7))
                painter.restore()
                return

            selected = bool(option.state & QStyle.StateFlag.State_Selected) and \
                bool(option.state & QStyle.StateFlag.State_Enabled)
            if selected:
                rect = QRectF(option.rect).adjusted(3, 1.5, -3, -1.5)
                draw_sketch_box(
                    painter, rect,
                    stable_seed("menuitem", int(rect.width()), int(rect.y())),
                    fill=PAPER, width=1.8, jitter=1.0,
                )
                menu_option.state &= ~QStyle.StateFlag.State_Selected
                menu_option.palette.setColor(QPalette.ColorRole.HighlightedText, INK)
            super().drawControl(element, menu_option, painter, widget)
            return

        if element == QStyle.ControlElement.CE_ItemViewItem and isinstance(option, QStyleOptionViewItem):
            view_option = QStyleOptionViewItem(option)
            selected = bool(option.state & QStyle.StateFlag.State_Selected)
            if selected:
                rect = QRectF(option.rect).adjusted(5, 2, -9, -2)
                if rect.width() > 4 and rect.height() > 4:
                    draw_sketch_box(
                        painter, rect,
                        stable_seed("item", int(rect.width()), int(rect.y())),
                        fill=PAPER, width=1.8, jitter=1.0,
                        shadow=True, shadow_offset=2.5,
                    )
                view_option.state &= ~QStyle.StateFlag.State_Selected
                view_option.palette.setColor(QPalette.ColorRole.HighlightedText, INK)
                view_option.palette.setColor(QPalette.ColorRole.Highlight, Qt.GlobalColor.transparent)
            super().drawControl(element, view_option, painter, widget)
            return

        super().drawControl(element, option, painter, widget)

    def drawComplexControl(self, control, option, painter, widget=None):
        if control == QStyle.ComplexControl.CC_ToolButton and isinstance(option, QStyleOptionToolButton):
            # Toolbar buttons are auto-raise: Qt only paints their panel on
            rect = QRectF(option.rect).adjusted(1.5, 1.5, -1.5, -1.5)
            seed = self._seed_for(widget, option.rect)
            state = option.state
            on = bool(state & QStyle.StateFlag.State_On)
            sunken = bool(state & QStyle.StateFlag.State_Sunken)
            hovered = bool(state & QStyle.StateFlag.State_MouseOver)
            enabled = bool(state & QStyle.StateFlag.State_Enabled)

            if rect.width() > 2 and rect.height() > 2:
                if on or sunken:
                    fill = INK
                elif hovered and enabled:
                    fill = HOVER
                else:
                    fill = PAPER
                target = rect.translated(1.5, 1.5) if (on or sunken) else rect
                draw_sketch_box(
                    painter, target, seed, fill=fill,
                    width=2.0 if enabled else 1.2,
                    shadow=enabled and not (on or sunken), shadow_offset=3.5,
                )

            label = QStyleOptionToolButton(option)
            self.drawControl(QStyle.ControlElement.CE_ToolButtonLabel, label, painter, widget)
            return

        if control in (QStyle.ComplexControl.CC_ComboBox, QStyle.ComplexControl.CC_SpinBox):
            rect = QRectF(option.rect).adjusted(1.5, 1.5, -1.5, -1.5)
            seed = self._seed_for(widget, option.rect)
            draw_sketch_box(painter, rect, seed, fill=PAPER, width=1.8)
            if control == QStyle.ComplexControl.CC_ComboBox:
                painter.save()
                painter.setPen(QPen(INK, 1.8, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
                cx = rect.right() - 12
                cy = rect.center().y() - 1
                painter.drawPath(sketch_line(QPointF(cx - 4, cy - 2), QPointF(cx, cy + 3), seed + 3, 0.5))
                painter.drawPath(sketch_line(QPointF(cx, cy + 3), QPointF(cx + 4, cy - 2), seed + 4, 0.5))
                painter.restore()
            return
        super().drawComplexControl(control, option, painter, widget)


class Tool:
    SELECT = "Select"
    TEXT_BOX = "Add Text Box"
    EDIT_TEXT = "Edit Text Span"
    EDIT_BLOCK = "Edit Text Block"
    MOVE_TEXT = "Reposition Text"
    NOTE = "Add Comment"
    HIGHLIGHT = "Highlight"
    UNDERLINE = "Underline"
    STRIKEOUT = "Strikethrough"
    RECTANGLE = "Draw Rectangle"
    LINE = "Draw Line"
    INK = "Freehand Draw"
    REDACT = "Redact (Blackout)"
    CROP = "Crop Page"
    LINK = "Insert Hyperlink"
    IMAGE = "Insert Image"
    SIGN = "Add Signature"


DRAG_RECT_TOOLS = {
    Tool.TEXT_BOX,
    Tool.HIGHLIGHT,
    Tool.UNDERLINE,
    Tool.STRIKEOUT,
    Tool.RECTANGLE,
    Tool.REDACT,
    Tool.CROP,
    Tool.LINK,
}

CLICK_TOOLS = {
    Tool.EDIT_TEXT,
    Tool.EDIT_BLOCK,
    Tool.NOTE,
    Tool.IMAGE,
    Tool.SIGN,
}

LINE_TOOLS = {Tool.LINE}


@dataclass
class TextSpanHit:
    text: str
    bbox: fitz.Rect
    size: float
    color: QColor
    font: str


@dataclass
class TextBlockHit:
    text: str
    bbox: fitz.Rect


class TextBoxDialog(QDialog):
    def __init__(self, parent: QWidget, title: str, initial_text: str = ""):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(520, 360)

        layout = QVBoxLayout(self)
        self.editor = QTextEdit()
        self.editor.setPlainText(initial_text)
        self.editor.setPlaceholderText("Type the text that should appear inside the selected PDF box...")
        layout.addWidget(self.editor)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def text(self) -> str:
        return self.editor.toPlainText()


class SearchReplaceDialog(QDialog):
    def __init__(self, parent: QWidget):
        super().__init__(parent)
        self.setWindowTitle("Search and Replace Text")
        self.resize(520, 220)

        layout = QVBoxLayout(self)
        form = QFormLayout()

        self.find_input = QLineEdit()
        self.replace_input = QLineEdit()
        self.scope_combo = QComboBox()
        self.scope_combo.addItems(["Current Page", "Whole Document"])
        self.case_sensitive = QCheckBox("Case sensitive")

        form.addRow("Find:", self.find_input)
        form.addRow("Replace with:", self.replace_input)
        form.addRow("Scope:", self.scope_combo)
        form.addRow("", self.case_sensitive)

        note = QLabel(
            "This replaces visible text by redacting the old match area and drawing the replacement text. "
            "It is best for simple, horizontal text."
        )
        note.setWordWrap(True)
        note.setObjectName("SmallNote")

        layout.addLayout(form)
        layout.addWidget(note)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def values(self) -> Tuple[str, str, str, bool]:
        return (
            self.find_input.text(),
            self.replace_input.text(),
            self.scope_combo.currentText(),
            self.case_sensitive.isChecked(),
        )


class DocumentInfoDialog(QDialog):
    def __init__(self, parent: QWidget, info: Dict[str, str], page_count: int):
        super().__init__(parent)
        self.setWindowTitle("Document Information")
        self.resize(620, 460)

        layout = QVBoxLayout(self)
        self.info_box = QTextEdit()
        self.info_box.setReadOnly(True)

        data = {"page_count": page_count, **(info or {})}
        self.info_box.setPlainText(json.dumps(data, indent=2, ensure_ascii=False))
        layout.addWidget(self.info_box)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(self.reject)
        buttons.accepted.connect(self.accept)
        layout.addWidget(buttons)


class MetadataEditDialog(QDialog):
    def __init__(self, parent: QWidget, info: Dict[str, str]):
        super().__init__(parent)
        self.setWindowTitle("Edit PDF Metadata")
        self.resize(520, 340)

        layout = QVBoxLayout(self)
        form = QFormLayout()

        self.fields: Dict[str, QLineEdit] = {}
        for key in ["title", "author", "subject", "keywords", "creator", "producer"]:
            line = QLineEdit()
            line.setText(str((info or {}).get(key, "") or ""))
            self.fields[key] = line
            form.addRow(key.title() + ":", line)

        layout.addLayout(form)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def metadata(self) -> Dict[str, str]:
        return {k: v.text() for k, v in self.fields.items()}


class SignaturePreview(QWidget):
    """Shows the typed signature inside a hand-drawn box, coherent with the UI."""

    def __init__(self, parent=None):
        super().__init__(parent)
        self._text = ""
        self._family = APP_FONT_FAMILY
        self._color = QColor("#111111")
        self._size = 48
        self.setMinimumHeight(120)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)

    def set_signature(self, text: str, family: str, color: QColor, size: int = 48):
        self._text = text
        self._family = family
        self._color = QColor(color)
        self._size = max(8, int(size))
        self.update()

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        r = QRectF(self.rect()).adjusted(3, 3, -3, -3)
        draw_sketch_box(painter, r, stable_seed("sigpreview", int(r.width())),
                        fill=PAPER, width=1.8, jitter=0.8)

        text = (self._text or "").strip()
        if not text:
            painter.setPen(QPen(GREY_FAINT))
            painter.setFont(QFont(APP_FONT_FAMILY, 12))
            painter.drawText(r, Qt.AlignmentFlag.AlignCenter,
                             "Type your name to preview your signature")
            painter.end()
            return

        # Draw at the chosen point size, but scale down if it would overflow
        # the preview box so the whole signature always stays visible.
        inner = r.adjusted(16, 12, -16, -12)
        size = self._size
        font = QFont(self._family)
        font.setPointSize(size)
        metrics = QFontMetrics(font)
        tw = metrics.horizontalAdvance(text)
        th = metrics.height()
        if tw > 0 and th > 0:
            factor = min(inner.width() / tw, inner.height() / th, 1.0)
            size = max(8, int(size * factor))
            font.setPointSize(size)
        painter.setFont(font)
        painter.setPen(QPen(self._color))
        painter.drawText(inner, Qt.AlignmentFlag.AlignCenter, text)
        painter.end()


class SignatureDialog(QDialog):
    """Create a typed signature from a prebuilt script font, with a live,
    sketch-styled preview and a colour choice."""

    def __init__(self, parent: QWidget, default_name: str = ""):
        super().__init__(parent)
        self.setWindowTitle("Add Signature")
        self.resize(560, 460)

        self._color = QColor("#111111")
        self._families = load_signature_fonts()

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(10)

        form = QFormLayout()
        self.name_input = QLineEdit()
        self.name_input.setText(default_name)
        self.name_input.setPlaceholderText("Type your full name…")
        self.name_input.textChanged.connect(self._refresh)
        form.addRow("Name:", self.name_input)
        layout.addLayout(form)

        layout.addWidget(QLabel("Signature style:"))
        self.style_list = QListWidget()
        self.style_list.setIconSize(QSize(0, 0))
        self.style_list.setMaximumHeight(150)
        for family in self._families:
            item = QListWidgetItem(family)
            item.setData(Qt.ItemDataRole.UserRole, family)
            f = QFont(family)
            f.setPointSize(22)
            item.setFont(f)
            self.style_list.addItem(item)
        self.style_list.setCurrentRow(0)
        self.style_list.currentRowChanged.connect(lambda _i: self._refresh())
        layout.addWidget(self.style_list)

        # Size + colour row.
        options_row = QHBoxLayout()
        options_row.addWidget(QLabel("Font size:"))
        self.size_spin = QSpinBox()
        self.size_spin.setRange(12, 120)
        self.size_spin.setValue(48)
        self.size_spin.setSuffix(" pt")
        self.size_spin.valueChanged.connect(lambda _v: self._refresh())
        options_row.addWidget(self.size_spin)
        options_row.addSpacing(16)
        options_row.addWidget(QLabel("Ink colour:"))
        self.color_button = SketchSwatchButton("Colour")
        self.color_button.set_swatch(self._color)
        self.color_button.clicked.connect(self._choose_color)
        options_row.addWidget(self.color_button)
        options_row.addStretch()
        layout.addLayout(options_row)

        self.preview = SignaturePreview()
        layout.addWidget(self.preview)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        self._ok_button = buttons.button(QDialogButtonBox.StandardButton.Ok)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

        self._refresh()

    def _choose_color(self):
        color = QColorDialog.getColor(self._color, self, "Signature Colour")
        if color.isValid():
            self._color = color
            self.color_button.set_swatch(color)
            self._refresh()

    def _refresh(self):
        family = self.current_family()
        text = self.name_input.text()
        self.preview.set_signature(text, family, self._color, self.size_spin.value())
        if self._ok_button is not None:
            self._ok_button.setEnabled(bool(text.strip()))

    def current_family(self) -> str:
        item = self.style_list.currentItem()
        if item is not None:
            return item.data(Qt.ItemDataRole.UserRole)
        return self._families[0] if self._families else APP_FONT_FAMILY

    def result_values(self) -> Tuple[str, str, QColor, int]:
        return (self.name_input.text().strip(), self.current_family(),
                QColor(self._color), int(self.size_spin.value()))


class HeaderFooterDialog(QDialog):
    """Configure header/footer text in six slots with page-number tokens."""

    def __init__(self, parent: QWidget):
        super().__init__(parent)
        self.setWindowTitle("Header / Footer & Page Numbers")
        self.resize(520, 360)
        self._color = QColor("#333333")

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(8)

        hint = QLabel("Use tokens:  {page}  {pages}  {date}")
        hint.setObjectName("SmallNote")
        layout.addWidget(hint)

        form = QFormLayout()
        self.fields: Dict[str, QLineEdit] = {}
        labels = [
            ("hl", "Header left:"), ("hc", "Header center:"), ("hr", "Header right:"),
            ("fl", "Footer left:"), ("fc", "Footer center:"), ("fr", "Footer right:"),
        ]
        for key, label in labels:
            line = QLineEdit()
            self.fields[key] = line
            form.addRow(label, line)
        # Sensible default: page number centered in the footer.
        self.fields["fc"].setText("Page {page} of {pages}")
        layout.addLayout(form)

        opts = QHBoxLayout()
        opts.addWidget(QLabel("Font size:"))
        self.size_spin = QSpinBox()
        self.size_spin.setRange(6, 48)
        self.size_spin.setValue(10)
        opts.addWidget(self.size_spin)
        opts.addWidget(QLabel("Margin:"))
        self.margin_spin = QSpinBox()
        self.margin_spin.setRange(6, 120)
        self.margin_spin.setValue(24)
        opts.addWidget(self.margin_spin)
        opts.addWidget(QLabel("Color:"))
        self.color_button = SketchSwatchButton("Color")
        self.color_button.set_swatch(self._color)
        self.color_button.clicked.connect(self._choose_color)
        opts.addWidget(self.color_button)
        opts.addStretch()
        layout.addLayout(opts)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _choose_color(self):
        color = QColorDialog.getColor(self._color, self, "Text Color")
        if color.isValid():
            self._color = color
            self.color_button.set_swatch(color)

    def values(self) -> Dict:
        return {
            "slots": {k: v.text().strip() for k, v in self.fields.items()},
            "size": self.size_spin.value(),
            "margin": self.margin_spin.value(),
            "color": QColor(self._color),
        }


class FormFillDialog(QDialog):
    """Edit the values of a PDF's interactive form fields (AcroForm widgets).

    `fields` is a list of dicts: {page, index, name, type, value, choices}.
    """

    def __init__(self, parent: QWidget, fields: List[Dict]):
        super().__init__(parent)
        self.setWindowTitle("Fill Form Fields")
        self.resize(560, 520)
        self._fields = fields
        self._editors: List = []

        layout = QVBoxLayout(self)
        layout.setContentsMargins(14, 14, 14, 14)
        layout.setSpacing(8)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        inner = QWidget()
        form = QFormLayout(inner)
        form.setContentsMargins(4, 4, 4, 4)

        for f in fields:
            ftype = f["type"]
            label = f["name"] or f"(field on p.{f['page']+1})"
            if ftype == "checkbox":
                editor = QCheckBox()
                editor.setChecked(str(f["value"]).lower() not in ("off", "", "false", "no", "0"))
            elif ftype in ("combobox", "listbox") and f.get("choices"):
                editor = QComboBox()
                editor.setEditable(ftype == "combobox")
                editor.addItems([str(c) for c in f["choices"]])
                cur = str(f["value"] or "")
                idx = editor.findText(cur)
                if idx >= 0:
                    editor.setCurrentIndex(idx)
                elif cur:
                    editor.setEditText(cur)
            else:  # text and everything else
                editor = QLineEdit()
                editor.setText(str(f["value"] or ""))
            self._editors.append(editor)
            form.addRow(label + ":", editor)

        scroll.setWidget(inner)
        layout.addWidget(scroll)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def result_values(self) -> List[Dict]:
        out = []
        for f, editor in zip(self._fields, self._editors):
            if isinstance(editor, QCheckBox):
                val = editor.isChecked()
            elif isinstance(editor, QComboBox):
                val = editor.currentText()
            else:
                val = editor.text()
            out.append({**f, "new_value": val})
        return out


class PasswordDialog(QDialog):
    def __init__(self, parent: QWidget, title: str = "Password Required"):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.resize(360, 130)

        layout = QVBoxLayout(self)
        self.password = QLineEdit()
        self.password.setEchoMode(QLineEdit.EchoMode.Password)
        self.password.setPlaceholderText("Enter PDF password")
        layout.addWidget(QLabel("This PDF is encrypted."))
        layout.addWidget(self.password)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def value(self) -> str:
        return self.password.text()


class PageCanvas(QWidget):
    """Interactive PDF page preview. Coordinates passed to parent are image pixels."""

    def __init__(self, callback):
        super().__init__()
        self.callback = callback
        self.pixmap: Optional[QPixmap] = None
        self.zoom = 1.0
        self.tool = Tool.SELECT
        self.search_rects: List[QRectF] = []
        self.edit_rects: List[QRectF] = []
        self.crop_preview_rect: Optional[QRectF] = None
        self._brand: Optional[QPixmap] = None
        self.xray_enabled = False
        self.xray_text: List[QRectF] = []
        self.xray_image: List[QRectF] = []
        self.xray_vector: List[QRectF] = []
        self.spell_rects: List[QRectF] = []
        self.stroke_color = QColor("#257A6C")
        self.fill_color = QColor("#e7f0ee")
        self.text_color = QColor("#1c1a17")
        self.line_width = 2

        self._margin = 14
        self._drag_start: Optional[QPointF] = None
        self._drag_current: Optional[QPointF] = None
        self._ink_points: List[QPointF] = []
        self._pan_last: Optional[QPointF] = None

        self.setMouseTracking(True)
        self.setMinimumSize(850, 1000)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        # Accept keyboard focus so Delete can remove a selected annotation.
        self.setFocusPolicy(Qt.FocusPolicy.StrongFocus)

    def set_page(
        self,
        pixmap: QPixmap,
        zoom: float,
        search_rects: Optional[List[QRectF]] = None,
        edit_rects: Optional[List[QRectF]] = None,
        crop_preview_rect: Optional[QRectF] = None,
        xray: Optional[Dict[str, List[QRectF]]] = None,
        spell_rects: Optional[List[QRectF]] = None,
    ):
        self.pixmap = pixmap
        self.zoom = zoom
        self.search_rects = search_rects or []
        self.edit_rects = edit_rects or []
        self.crop_preview_rect = crop_preview_rect
        self.spell_rects = spell_rects or []
        xray = xray or {}
        self.xray_text = xray.get("text", [])
        self.xray_image = xray.get("image", [])
        self.xray_vector = xray.get("vector", [])
        self._drag_start = None
        self._drag_current = None
        self._ink_points = []
        self.resize(self.sizeHint())
        self.updateGeometry()
        self.update()

    def set_tool(self, tool: str):
        self.tool = tool
        self._drag_start = None
        self._drag_current = None
        self._ink_points = []
        self.update()

    def set_xray_enabled(self, enabled: bool):
        self.xray_enabled = bool(enabled)
        self.update()

    def set_style(self, stroke_color: QColor, fill_color: QColor, text_color: QColor, line_width: int):
        self.stroke_color = QColor(stroke_color)
        self.fill_color = QColor(fill_color)
        self.text_color = QColor(text_color)
        self.line_width = max(1, int(line_width))
        self.update()

    @staticmethod
    def _squiggle_path(rect: QRectF) -> QPainterPath:
        """A doodly, hand-drawn squiggly underline spanning the word's width,
        sitting just below its baseline."""
        path = QPainterPath()
        y = rect.bottom() + 2.0
        x0 = rect.left()
        x1 = rect.right()
        amp = 2.4          # squiggle height
        step = 4.5         # horizontal wavelength
        rnd = random.Random(stable_seed("spell", int(rect.left()), int(rect.bottom())))
        path.moveTo(x0, y)
        x = x0
        up = True
        while x < x1:
            nx = min(x + step, x1)
            midx = (x + nx) / 2
            # Alternating control points create the up/down wobble, with a
            # little random jitter so it looks drawn by hand.
            cy = y - amp + rnd.uniform(-0.5, 0.5) if up else y + amp + rnd.uniform(-0.5, 0.5)
            path.quadTo(midx, cy, nx, y + rnd.uniform(-0.4, 0.4))
            up = not up
            x = nx
        return path

    def sizeHint(self) -> QSize:
        if self.pixmap is None:
            return QSize(920, 1120)
        return QSize(self.pixmap.width() + self._margin * 2, self.pixmap.height() + self._margin * 2)

    def _page_rect(self) -> QRectF:
        if self.pixmap is None:
            return QRectF()
        x = max((self.width() - self.pixmap.width()) / 2, self._margin)
        y = self._margin
        return QRectF(x, y, self.pixmap.width(), self.pixmap.height())

    def _event_to_image_point(self, event) -> Optional[QPointF]:
        rect = self._page_rect()
        point = QPointF(event.position().x() - rect.x(), event.position().y() - rect.y())
        if 0 <= point.x() <= rect.width() and 0 <= point.y() <= rect.height():
            return point
        return None

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        painter.fillRect(self.rect(), PAPER)

        if self.pixmap is None:
            center = QRectF(self.rect())
            if self._brand is None:
                self._brand = load_brand_pixmap(128) or QPixmap()
            if not self._brand.isNull():
                painter.drawPixmap(
                    int(center.center().x() - self._brand.width() / 2),
                    int(center.center().y() - self._brand.height() - 34),
                    self._brand,
                )
            painter.setPen(GREY_MUTED)
            painter.setFont(QFont(APP_FONT_FAMILY, 20))
            painter.drawText(center.adjusted(0, 44, 0, 0), Qt.AlignmentFlag.AlignCenter, "Open a PDF to get started")
            painter.setPen(GREY_FAINT)
            painter.setFont(QFont(APP_FONT_FAMILY, 13))
            painter.drawText(center.adjusted(0, 108, 0, 0), Qt.AlignmentFlag.AlignCenter,
                             "Click here, drop a PDF, or press Ctrl+O")
            return

        page_rect = self._page_rect()

        page_seed = stable_seed("page", int(page_rect.width()), int(page_rect.height()))
        painter.setPen(Qt.PenStyle.NoPen)
        painter.setBrush(INK)
        painter.drawPath(sketch_path(page_rect.translated(6, 6), page_seed, 1.1))
        painter.setBrush(PAPER)
        painter.drawPath(sketch_path(page_rect, page_seed, 1.1))
        painter.drawPixmap(int(page_rect.x()), int(page_rect.y()), self.pixmap)
        # Outline goes on top: the page bitmap is opaque and would otherwise
        painter.setBrush(Qt.BrushStyle.NoBrush)
        painter.setPen(QPen(INK, 2.0, Qt.PenStyle.SolidLine,
                            Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
        painter.drawPath(sketch_path(page_rect, page_seed, 1.1))

        painter.save()
        painter.translate(page_rect.x(), page_rect.y())

        if self.xray_enabled:
            painter.fillRect(QRectF(0, 0, page_rect.width(), page_rect.height()),
                             QColor(255, 255, 255, 205))

            painter.setPen(QPen(INK, 1))
            painter.setBrush(QColor(0, 0, 0, 42))
            for rect in self.xray_text:
                painter.drawRect(rect)

            painter.setPen(QPen(INK, 1, Qt.PenStyle.DashLine))
            painter.setBrush(QBrush(INK, Qt.BrushStyle.BDiagPattern))
            for rect in self.xray_image:
                painter.drawRect(rect)

            painter.setPen(QPen(INK, 1.2, Qt.PenStyle.DotLine))
            painter.setBrush(Qt.BrushStyle.NoBrush)
            for rect in self.xray_vector:
                painter.drawRect(rect)

            self._paint_xray_legend(painter)

        if self.search_rects:
            painter.setPen(QPen(INK, 2))
            painter.setBrush(QColor(0, 0, 0, 38))
            for rect in self.search_rects:
                painter.drawPath(sketch_path(rect, stable_seed("find", rect.x(), rect.y()), 0.8))

        if self.tool in {Tool.EDIT_TEXT, Tool.EDIT_BLOCK, Tool.MOVE_TEXT} and self.edit_rects:
            painter.setPen(QPen(INK, 1.2, Qt.PenStyle.DashLine))
            painter.setBrush(QColor(0, 0, 0, 16))
            for rect in self.edit_rects:
                painter.drawRect(rect)

        if self.crop_preview_rect:
            painter.setPen(QPen(INK, 2, Qt.PenStyle.DashLine))
            painter.setBrush(QColor(0, 0, 0, 18))
            painter.drawRect(self.crop_preview_rect)

        # Spell-check: a hand-drawn squiggly underline beneath misspelled words.
        if self.spell_rects:
            pen = QPen(QColor(210, 40, 40), 2.6, Qt.PenStyle.SolidLine,
                       Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin)
            painter.setPen(pen)
            painter.setBrush(Qt.BrushStyle.NoBrush)
            for rect in self.spell_rects:
                painter.drawPath(self._squiggle_path(rect))

        if self._drag_start and self._drag_current:
            temp = QRectF(self._drag_start, self._drag_current).normalized()
            if self.tool == Tool.REDACT:
                painter.setPen(QPen(INK, 2.4, Qt.PenStyle.SolidLine))
                painter.setBrush(QBrush(INK, Qt.BrushStyle.FDiagPattern))
            else:
                painter.setPen(QPen(INK, 2, Qt.PenStyle.DashLine))
                painter.setBrush(QColor(0, 0, 0, 26))
            painter.drawPath(sketch_path(temp, stable_seed("drag", int(temp.width()), int(temp.height())), 1.0))

        if self._drag_start and self._drag_current and self.tool == Tool.LINE:
            painter.setPen(QPen(self.stroke_color, max(1.0, float(self.line_width))))
            painter.drawLine(self._drag_start, self._drag_current)

        if self._drag_start and self._drag_current and self.tool == Tool.MOVE_TEXT:
            painter.setPen(QPen(INK, 2.2, Qt.PenStyle.DashLine))
            painter.drawLine(self._drag_start, self._drag_current)
            painter.drawEllipse(self._drag_current, 5, 5)

        if len(self._ink_points) > 1:
            painter.setPen(QPen(self.stroke_color, max(1.0, float(self.line_width)), Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap, Qt.PenJoinStyle.RoundJoin))
            for a, b in zip(self._ink_points, self._ink_points[1:]):
                painter.drawLine(a, b)

        painter.restore()

    def _paint_xray_legend(self, painter: QPainter):
        """Small legend so the tints are self-explanatory."""
        entries = [
            ("solid", f"Editable text · {len(self.xray_text)}"),
            ("hatch", f"Image · {len(self.xray_image)}"),
            ("dotted", f"Vector · {len(self.xray_vector)}"),
        ]
        pad, swatch, row_h = 11, 11, 20
        painter.setFont(QFont(APP_FONT_FAMILY, 10))
        metrics = painter.fontMetrics()
        width = max(metrics.horizontalAdvance(label) for _, label in entries) + swatch + pad * 3
        height = row_h * len(entries) + pad

        box = QRectF(14, 14, width, height)
        draw_sketch_box(painter, box, stable_seed("legend", int(width), int(height)),
                        fill=PAPER, width=1.8)

        y = box.y() + pad / 2 + 4
        for texture, label in entries:
            marker = QRectF(box.x() + pad, y + 3, swatch, swatch)
            painter.setPen(QPen(INK, 1))
            if texture == "solid":
                painter.setBrush(QColor(0, 0, 0, 60))
            elif texture == "hatch":
                painter.setBrush(QBrush(INK, Qt.BrushStyle.BDiagPattern))
            else:
                painter.setPen(QPen(INK, 1.2, Qt.PenStyle.DotLine))
                painter.setBrush(Qt.BrushStyle.NoBrush)
            painter.drawRect(marker)
            painter.setPen(INK)
            painter.drawText(QRectF(box.x() + pad * 2 + swatch, y, width, row_h),
                             Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter, label)
            y += row_h

    def wheelEvent(self, event):
        """Ctrl/Cmd + wheel zooms. Normal wheel still scrolls via QScrollArea."""
        modifiers = event.modifiers()
        if modifiers & (Qt.KeyboardModifier.ControlModifier | Qt.KeyboardModifier.MetaModifier):
            self.callback("__zoom_delta__", event.angleDelta().y())
            event.accept()
            return
        super().wheelEvent(event)

    def mouseDoubleClickEvent(self, event):
        if self.pixmap is not None and event.button() == Qt.MouseButton.LeftButton:
            # In the Select/Move tool, double-click recolours the annotation
            # under the cursor; otherwise it fits the page width.
            if self.tool == Tool.MOVE_TEXT:
                p = self._event_to_image_point(event)
                if p is not None and self.callback("__recolor_annot__", p):
                    event.accept()
                    return
            self.callback("__fit_width__", None)
            event.accept()
            return
        super().mouseDoubleClickEvent(event)

    def keyPressEvent(self, event):
        # Delete/Backspace removes the currently selected annotation.
        if event.key() in (Qt.Key.Key_Delete, Qt.Key.Key_Backspace):
            if self.callback("__delete_selected_annot__", None):
                event.accept()
                return
        super().keyPressEvent(event)

    def mousePressEvent(self, event):
        if self.pixmap is None:
            if event.button() == Qt.MouseButton.LeftButton:
                self.callback("__open_pdf__", None)
                event.accept()
            return

        if event.button() in (Qt.MouseButton.MiddleButton, Qt.MouseButton.RightButton):
            self._pan_last = event.position()
            self.setCursor(Qt.CursorShape.ClosedHandCursor)
            event.accept()
            return

        if event.button() != Qt.MouseButton.LeftButton:
            return

        p = self._event_to_image_point(event)
        if p is None:
            return

        if self.tool == Tool.MOVE_TEXT:
            self.setFocus(Qt.FocusReason.MouseFocusReason)
            ok = bool(self.callback("__move_text_start__", p))
            if ok:
                self._drag_start = p
                self._drag_current = p
                self.update()
            return

        if self.tool in CLICK_TOOLS:
            self.callback(self.tool, p)
            return

        if self.tool in DRAG_RECT_TOOLS or self.tool in LINE_TOOLS:
            self._drag_start = p
            self._drag_current = p
            self.update()
            return

        if self.tool == Tool.INK:
            self._ink_points = [p]
            self.update()
            return

    def mouseMoveEvent(self, event):
        if self.pixmap is None:
            return

        if self._pan_last is not None:
            current = event.position()
            delta = current - self._pan_last
            self._pan_last = current
            self.callback("__pan__", delta)
            event.accept()
            return

        p = self._event_to_image_point(event)
        if p is None:
            return

        if self._drag_start and (self.tool in DRAG_RECT_TOOLS or self.tool in LINE_TOOLS or self.tool == Tool.MOVE_TEXT):
            self._drag_current = p
            self.update()
            return

        if self.tool == Tool.INK and self._ink_points:
            self._ink_points.append(p)
            self.update()

    def mouseReleaseEvent(self, event):
        if self.pixmap is None:
            return

        if event.button() in (Qt.MouseButton.MiddleButton, Qt.MouseButton.RightButton):
            self._pan_last = None
            self.setCursor(Qt.CursorShape.ArrowCursor)
            event.accept()
            return

        if event.button() != Qt.MouseButton.LeftButton:
            return

        p = self._event_to_image_point(event)
        if p is None:
            return

        if self._drag_start and self.tool == Tool.MOVE_TEXT:
            start = self._drag_start
            self._drag_start = None
            self._drag_current = None
            self.update()
            if (start - p).manhattanLength() > 4:
                self.callback("__move_text_end__", (start, p))
            return

        if self._drag_start and self.tool in DRAG_RECT_TOOLS:
            rect = QRectF(self._drag_start, p).normalized()
            self._drag_start = None
            self._drag_current = None
            self.update()
            if rect.width() >= 8 and rect.height() >= 8:
                self.callback(self.tool, rect)
            return

        if self._drag_start and self.tool in LINE_TOOLS:
            start = self._drag_start
            self._drag_start = None
            self._drag_current = None
            self.update()
            if (start - p).manhattanLength() > 8:
                self.callback(self.tool, (start, p))
            return

        if self.tool == Tool.INK and len(self._ink_points) > 1:
            points = list(self._ink_points)
            self._ink_points = []
            self.update()
            self.callback(self.tool, points)


def _open_worker_doc(source: Dict[str, object]):
    """Open an independent PyMuPDF document inside a worker thread."""
    if source.get("type") == "path":
        doc = fitz.open(str(source["path"]))
    else:
        doc = fitz.open(stream=source["data"], filetype="pdf")

    password = str(source.get("password") or "")
    if doc.needs_pass and password:
        doc.authenticate(password)
    return doc


class ThumbnailWorker(QObject):
    thumbnail_ready = pyqtSignal(int, int, bytes)
    progress = pyqtSignal(int, int, int)
    error = pyqtSignal(int, str)
    finished = pyqtSignal(int)

    def __init__(self, job_id: int, source: Dict[str, object], thumb_zoom: float = 0.17):
        super().__init__()
        self.job_id = job_id
        self.source = source
        self.thumb_zoom = thumb_zoom
        self.cancelled = False

    @pyqtSlot()
    def run(self):
        doc = None
        try:
            doc = _open_worker_doc(self.source)
            total = doc.page_count
            for page_index in range(total):
                if self.cancelled:
                    break
                page = doc[page_index]
                pix = page.get_pixmap(matrix=fitz.Matrix(self.thumb_zoom, self.thumb_zoom), alpha=False)
                self.thumbnail_ready.emit(self.job_id, page_index, pix.tobytes("png"))
                self.progress.emit(self.job_id, page_index + 1, total)
        except Exception as exc:
            self.error.emit(self.job_id, str(exc))
        finally:
            if doc:
                doc.close()
            self.finished.emit(self.job_id)


class SearchWorker(QObject):
    finished_results = pyqtSignal(int, str, object)
    progress = pyqtSignal(int, int, int)
    error = pyqtSignal(int, str)
    finished = pyqtSignal(int)

    def __init__(self, job_id: int, source: Dict[str, object], query: str):
        super().__init__()
        self.job_id = job_id
        self.source = source
        self.query = query
        self.cancelled = False

    @pyqtSlot()
    def run(self):
        doc = None
        results = []
        try:
            doc = _open_worker_doc(self.source)
            total = doc.page_count
            for page_index in range(total):
                if self.cancelled:
                    break
                for rect in doc[page_index].search_for(self.query):
                    results.append((page_index, rect.x0, rect.y0, rect.x1, rect.y1))
                self.progress.emit(self.job_id, page_index + 1, total)
            self.finished_results.emit(self.job_id, self.query, results)
        except Exception as exc:
            self.error.emit(self.job_id, str(exc))
        finally:
            if doc:
                doc.close()
            self.finished.emit(self.job_id)


class ImageExportWorker(QObject):
    progress = pyqtSignal(int, int, int)
    error = pyqtSignal(int, str)
    done = pyqtSignal(int, str, int)
    finished = pyqtSignal(int)

    def __init__(self, job_id: int, source: Dict[str, object], folder: str, base_name: str, ext: str, scale: float = 2.0):
        super().__init__()
        self.job_id = job_id
        self.source = source
        self.folder = folder
        self.base_name = base_name
        self.ext = ext
        self.scale = scale
        self.cancelled = False

    @pyqtSlot()
    def run(self):
        doc = None
        count = 0
        try:
            doc = _open_worker_doc(self.source)
            total = doc.page_count
            output_dir = Path(self.folder)
            output_dir.mkdir(parents=True, exist_ok=True)
            for page_index in range(total):
                if self.cancelled:
                    break
                pix = doc[page_index].get_pixmap(matrix=fitz.Matrix(self.scale, self.scale), alpha=False)
                out = output_dir / f"{self.base_name}_page_{page_index + 1:03d}.{self.ext}"
                pix.save(str(out))
                count += 1
                self.progress.emit(self.job_id, page_index + 1, total)
            self.done.emit(self.job_id, self.folder, count)
        except Exception as exc:
            self.error.emit(self.job_id, str(exc))
        finally:
            if doc:
                doc.close()
            self.finished.emit(self.job_id)


class DocumentTab:
    """All per-document state for one open PDF, so multiple PDFs can live in
    tabs at once (Adobe-Acrobat style). Editor-wide preferences such as the
    active tool, colours and line width are intentionally NOT stored here — they
    stay shared across tabs, matching how Acrobat behaves."""

    def __init__(self, doc, file_path, password=""):
        self.doc = doc
        self.file_path = file_path
        self.password = password
        self.current_page_index = 0
        self.zoom = 1.25
        self.is_dirty = False
        self.search_results: List[Tuple[int, "fitz.Rect"]] = []
        self.search_index = -1
        self.undo_stack: List[bytes] = []
        self.redo_stack: List[bytes] = []
        self.doc_version = 0
        self.render_cache: Dict[Tuple[int, float, int], QPixmap] = {}
        self.render_cache_order: List[Tuple[int, float, int]] = []
        self.xray_cache: Dict[Tuple[int, int, int], Dict[str, List[QRectF]]] = {}
        self.page_structure_changed = False
        # A freshly opened document defaults to Fit Page the first time it is
        # shown; afterwards the user's chosen zoom is preserved across switches.
        self.needs_fit_width = True
        # Signature placements for THIS document, tracked with their clean
        # source data so they can be re-drawn (moved) without reading text back
        # from the PDF, whose subset script-font encoding can be unreadable.
        self.signatures: List[Dict] = []

    def title(self) -> str:
        return Path(self.file_path).name if self.file_path else "Untitled"


class SketchTabBar(QWidget):
    """A hand-drawn document tab strip, coherent with the sketch UI.

    Each tab is painted as a wobbly sketch box (filled when active) with the
    document name and a small close cross. Signals let the window switch or
    close documents.
    """

    tab_selected = pyqtSignal(int)
    tab_closed = pyqtSignal(int)
    new_tab_requested = pyqtSignal()

    _TAB_H = 30
    _MIN_TAB_W = 120
    _MAX_TAB_W = 220
    _GAP = 8
    _PAD = 6
    _PLUS_W = 34

    def __init__(self, parent=None):
        super().__init__(parent)
        self._tabs: List[str] = []
        self._active = -1
        self._hover = -1
        self._hover_close = -1
        self._hover_plus = False
        self.setMouseTracking(True)
        self.setFixedHeight(self._TAB_H + 12)
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)

    def set_tabs(self, titles: List[str], active: int):
        self._tabs = list(titles)
        self._active = active
        self.setVisible(len(self._tabs) > 0)
        self.update()

    # -- geometry helpers ---------------------------------------------------
    def _tab_width(self) -> int:
        if not self._tabs:
            return self._MIN_TAB_W
        avail = self.width() - self._PAD * 2 - self._PLUS_W - self._GAP
        per = (avail - self._GAP * (len(self._tabs) - 1)) / max(1, len(self._tabs))
        return int(max(self._MIN_TAB_W, min(self._MAX_TAB_W, per)))

    def _tab_rect(self, index: int) -> QRectF:
        w = self._tab_width()
        x = self._PAD + index * (w + self._GAP)
        return QRectF(x, 6, w, self._TAB_H)

    def _plus_rect(self) -> QRectF:
        w = self._tab_width()
        x = self._PAD + len(self._tabs) * (w + self._GAP)
        return QRectF(x, 6, self._PLUS_W, self._TAB_H)

    def _close_rect(self, tab_rect: QRectF) -> QRectF:
        s = 16
        return QRectF(tab_rect.right() - s - 6,
                      tab_rect.center().y() - s / 2, s, s)

    # -- events -------------------------------------------------------------
    def mouseMoveEvent(self, event):
        pos = event.position()
        self._hover = -1
        self._hover_close = -1
        self._hover_plus = self._plus_rect().contains(pos)
        for i in range(len(self._tabs)):
            r = self._tab_rect(i)
            if r.contains(pos):
                self._hover = i
                if self._close_rect(r).contains(pos):
                    self._hover_close = i
                break
        self.update()

    def leaveEvent(self, event):
        self._hover = self._hover_close = -1
        self._hover_plus = False
        self.update()

    def mousePressEvent(self, event):
        pos = event.position()
        if self._plus_rect().contains(pos):
            self.new_tab_requested.emit()
            return
        for i in range(len(self._tabs)):
            r = self._tab_rect(i)
            if r.contains(pos):
                if self._close_rect(r).contains(pos):
                    self.tab_closed.emit(i)
                elif i != self._active:
                    self.tab_selected.emit(i)
                return
        # Middle-click closes a tab, like most tabbed apps.
        if event.button() == Qt.MouseButton.MiddleButton:
            for i in range(len(self._tabs)):
                if self._tab_rect(i).contains(pos):
                    self.tab_closed.emit(i)
                    return

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
        painter.fillRect(self.rect(), PAPER)

        fm = painter.fontMetrics()
        for i, title in enumerate(self._tabs):
            r = self._tab_rect(i)
            active = (i == self._active)
            hovered = (i == self._hover)
            seed = stable_seed("doctab", i, int(r.width()))
            fill = HOVER if (hovered and not active) else PAPER
            draw_sketch_box(
                painter, r.adjusted(1.5, 1.5, -1.5, -1.5), seed,
                fill=fill, width=2.0 if active else 1.5,
                shadow=active, shadow_offset=2.5, jitter=0.9,
            )

            close_r = self._close_rect(r)
            text_area = r.adjusted(10, 0, -(close_r.width() + 12), 0)
            elided = fm.elidedText(title, Qt.TextElideMode.ElideRight,
                                   int(text_area.width()))
            painter.setPen(QPen(INK if active else GREY_MUTED))
            painter.drawText(text_area, Qt.AlignmentFlag.AlignVCenter |
                             Qt.AlignmentFlag.AlignLeft, elided)

            # Close cross (drawn as two little sketchy strokes).
            cc = close_r.center()
            if self._hover_close == i:
                draw_sketch_box(painter, close_r.adjusted(1, 1, -1, -1),
                                seed + 7, fill=HOVER,
                                width=1.2, jitter=0.7)
            painter.setPen(QPen(INK, 1.6, Qt.PenStyle.SolidLine,
                                Qt.PenCapStyle.RoundCap))
            d = 4.0
            painter.drawPath(sketch_line(QPointF(cc.x() - d, cc.y() - d),
                                         QPointF(cc.x() + d, cc.y() + d),
                                         seed + 3, 0.5))
            painter.drawPath(sketch_line(QPointF(cc.x() + d, cc.y() - d),
                                         QPointF(cc.x() - d, cc.y() + d),
                                         seed + 4, 0.5))

        # The "new tab" (+) button.
        pr = self._plus_rect()
        draw_sketch_box(painter, pr.adjusted(1.5, 1.5, -1.5, -1.5),
                        stable_seed("doctab_plus", int(pr.width())),
                        fill=HOVER if self._hover_plus else PAPER,
                        width=1.5, jitter=0.8)
        pc = pr.center()
        painter.setPen(QPen(INK, 1.8, Qt.PenStyle.SolidLine, Qt.PenCapStyle.RoundCap))
        painter.drawPath(sketch_line(QPointF(pc.x() - 6, pc.y()),
                                     QPointF(pc.x() + 6, pc.y()),
                                     stable_seed("plush"), 0.5))
        painter.drawPath(sketch_line(QPointF(pc.x(), pc.y() - 6),
                                     QPointF(pc.x(), pc.y() + 6),
                                     stable_seed("plusv"), 0.5))
        painter.end()


class PdfStudioOverhaulPro(QMainWindow):
    def __init__(self):
        super().__init__()

        self.doc: Optional[fitz.Document] = None
        self.file_path: Optional[str] = None
        self.current_page_index = 0
        self.zoom = 1.25
        self.current_tool = Tool.SELECT
        self.annotation_color = QColor("#000000")
        self.fill_color = QColor("#ffffff")
        self.text_color = QColor("#000000")
        self.line_width = 2
        self.opacity_percent = 35
        self.is_dirty = False

        self.search_results: List[Tuple[int, fitz.Rect]] = []
        self.search_index = -1
        self.pending_image_path: Optional[str] = None
        self.pending_signature: Optional[Dict] = None
        self._selected_annot: Optional[Tuple[int, int]] = None
        self.signatures: List[Dict] = []
        self.undo_stack: List[bytes] = []
        self.redo_stack: List[bytes] = []
        self.max_history = 10
        self._pdf_password = ""
        self._workers: List[Tuple[QThread, QObject]] = []
        self._thumbnail_job_id = 0
        self._search_job_id = 0
        self._export_job_id = 0
        self._page_structure_changed = False

        self._doc_version = 0
        self._render_cache: Dict[Tuple[int, float, int], QPixmap] = {}
        self._render_cache_order: List[Tuple[int, float, int]] = []
        self.xray_enabled = False
        self._xray_cache: Dict[Tuple[int, int, int], Dict[str, List[QRectF]]] = {}
        self.spellcheck_enabled = False
        self._spell_cache: Dict[Tuple[int, int, int], List[QRectF]] = {}

        # Open documents, one per tab (Acrobat-style). self.doc etc. always
        # mirror the active tab so all existing logic keeps working unchanged.
        self.tabs: List[DocumentTab] = []
        self.active_tab_index = -1
        self._switching_tab = False

        self.setWindowTitle("Suketchi PDF Reader")
        self._apply_adaptive_geometry()

        self.setAcceptDrops(True)

        self._build_ui()
        self._apply_theme()
        self._update_contextual_panel()
        self._set_document_controls(False)

    def _apply_adaptive_geometry(self):
        """Size and center the window relative to the available screen area so
        the app fits any display size / resolution."""
        screen = self.screen() or QGuiApplication.primaryScreen()
        if screen is None:
            self.resize(1540, 950)
            return

        available = screen.availableGeometry()
        # Prefer ~85% of the available area, but never exceed the screen and
        # never shrink below a usable minimum.
        width = max(960, min(int(available.width() * 0.85), available.width()))
        height = max(600, min(int(available.height() * 0.85), available.height()))

        # A sensible minimum so widgets never get clipped on tiny displays.
        self.setMinimumSize(
            min(960, available.width()),
            min(600, available.height()),
        )

        self.resize(width, height)

        # Center within the available geometry.
        frame = self.frameGeometry()
        frame.moveCenter(available.center())
        self.move(frame.topLeft())

    def _reposition_theme_toggle(self):
        # The theme toggle now lives in the ribbon toolbar, so no manual
        # positioning is needed.
        pass

    def _build_ui(self):
        self.status = QStatusBar()
        self.setStatusBar(self.status)
        self.progress_bar = QProgressBar()
        self.progress_bar.setFixedWidth(190)
        self.progress_bar.setVisible(False)
        self.status.addPermanentWidget(self.progress_bar)

        self._build_top_bar()

        self.page_list = QListWidget()
        self.page_list.setIconSize(QSize(96, 130))
        self.page_list.currentRowChanged.connect(self.on_page_selected)
        # Drag-to-reorder pages (Combine/Organize).
        self.page_list.setDragDropMode(QListWidget.DragDropMode.InternalMove)
        self.page_list.setDefaultDropAction(Qt.DropAction.MoveAction)
        self.page_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        self.page_list.model().rowsMoved.connect(self._on_pages_reordered)
        self._suppress_page_reorder = False

        self.outline_list = QListWidget()
        self.outline_list.itemClicked.connect(self.go_to_outline_item)
        outline_panel = self._build_outline_panel()

        search_panel = self._build_search_panel()

        self.comments_list = QListWidget()
        self.comments_list.itemClicked.connect(self.go_to_comment_item)
        notes_panel = self._build_notes_panel()

        self.sidebar = QTabWidget()
        self.sidebar.setObjectName("SidebarTabs")
        self.sidebar.setMinimumWidth(238)
        self.sidebar.setMaximumWidth(320)
        self.sidebar.setUsesScrollButtons(False)
        self.sidebar.tabBar().setExpanding(True)
        self.sidebar.addTab(self.page_list, "Pages")
        self.sidebar.addTab(outline_panel, "Outline")
        self.sidebar.addTab(search_panel, "Find")
        self.sidebar.addTab(notes_panel, "Notes")

        self.canvas = PageCanvas(self.handle_canvas_action)
        self.scroll = QScrollArea()
        self.scroll.setWidgetResizable(False)
        self.scroll.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.scroll.setWidget(self.canvas)

        self.properties = self._build_properties_panel()
        self.properties.setMinimumWidth(210)
        self.properties.setMaximumWidth(250)

        body = QSplitter(Qt.Orientation.Horizontal)
        body.addWidget(self.sidebar)
        body.addWidget(self.scroll)
        body.addWidget(self.properties)
        body.setStretchFactor(0, 1)
        body.setStretchFactor(1, 7)
        body.setStretchFactor(2, 1)

        # Distribute the splitter proportionally to the current window width so
        # the layout adapts to any display size. The side panels are clamped to
        # their own min/max widths; the center canvas takes the remaining space.
        total = max(self.width(), 960)
        side = max(238, min(int(total * 0.17), 320))
        props = max(210, min(int(total * 0.15), 250))
        center = max(total - side - props, 400)
        body.setSizes([side, center, props])

        # Wrap the splitter so the window has breathing room on the left/right
        # (and top/bottom), coherent with the padded look of the rest of the UI
        # instead of the content sitting flush against the window edges.
        self.tab_bar = SketchTabBar()
        self.tab_bar.tab_selected.connect(self.switch_to_tab)
        self.tab_bar.tab_closed.connect(self.close_tab)
        self.tab_bar.new_tab_requested.connect(self.open_pdf)
        self.tab_bar.setVisible(False)

        container = QWidget()
        container.setObjectName("BodyContainer")
        outer = QVBoxLayout(container)
        outer.setContentsMargins(14, 6, 14, 12)
        outer.setSpacing(4)
        outer.addWidget(self.tab_bar)
        outer.addWidget(body, 1)

        self.setCentralWidget(container)

    def _build_top_bar(self):
        top = QToolBar("Ribbon")
        top.setMovable(False)
        top.setObjectName("Ribbon")
        top.setIconSize(QSize(16, 16))
        top.setContentsMargins(6, 4, 6, 4)
        # Disable the default right-click context menu on the ribbon (it would
        # otherwise offer to hide the toolbar, making "the menu go away").
        top.setContextMenuPolicy(Qt.ContextMenuPolicy.PreventContextMenu)
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, top)

        brand_pixmap = load_seal_pixmap(24)
        if brand_pixmap is not None:
            brand_icon = QLabel()
            brand_icon.setPixmap(brand_pixmap)
            brand_icon.setContentsMargins(2, 0, 6, 0)
            top.addWidget(brand_icon)
        brand_text = QLabel("Suketchi")
        brand_text.setObjectName("BrandText")
        top.addWidget(brand_text)
        top.addSeparator()

        self.open_action = self._add_action(top, "Open", self.open_pdf, "Ctrl+O")
        self.save_action = self._add_action(top, "Save", self.save_as_pdf, "Ctrl+S")
        for action in (self.open_action, self.save_action):
            button = top.widgetForAction(action)
            if button is not None:
                button.setObjectName("Primary")
        top.addSeparator()

        self.undo_action = self._add_action(top, "Undo", self.undo, "Ctrl+Z")
        self.redo_action = self._add_action(top, "Redo", self.redo, "Ctrl+Y")
        top.addSeparator()

        self.file_menu_button = self._menu_button("File", [
            ("Save Encrypted Copy…", self.save_encrypted_copy),
            ("Reduce File Size…", self.compress_pdf),
            None,
            ("Digitally Sign (.p12 / .pfx)…", self.digitally_sign),
            ("Verify Digital Signatures…", self.verify_signatures),
            None,
            ("Document Properties…", self.show_document_info),
            ("Edit Metadata…", self.edit_metadata),
        ])
        top.addWidget(self.file_menu_button)

        self.edit_menu_button = self._menu_button("Edit", [
            ("Find and Replace…", self.search_and_replace),
            ("Fill Form Fields…", self.fill_form_fields),
            None,
            ("Reflow Current Page", self.reflow_current_page_to_new_page),
            ("Reflow Whole Document…", self.reflow_document_to_new_pdf),
        ])
        top.addWidget(self.edit_menu_button)

        self.insert_menu_button = self._menu_button("Insert", [
            ("Image…", self.prepare_insert_image),
            ("Signature…", self.prepare_signature),
            ("Watermark…", self.add_watermark),
            ("Stamp…", self.add_stamp),
            None,
            ("Header / Footer & Page Numbers…", self.add_header_footer),
            ("Bates Numbering…", self.add_bates_numbering),
            None,
            ("Blank Page", self.insert_blank_page),
            ("Merge PDF After This Page…", self.insert_pdf_after_current),
        ])
        top.addWidget(self.insert_menu_button)

        self.pages_menu_button = self._menu_button("Pages", [
            ("Rotate Clockwise", self.rotate_current_page_clockwise),
            ("Rotate Counter-clockwise", self.rotate_current_page_counterclockwise),
            None,
            ("Move Page Up", self.move_page_up),
            ("Move Page Down", self.move_page_down),
            None,
            ("Duplicate Page", self.duplicate_current_page),
            ("Delete Page", self.delete_current_page),
            None,
            ("Insert File Before This Page…", lambda: self.insert_pdf_at(before=True)),
            ("Insert File After This Page…", lambda: self.insert_pdf_at(before=False)),
            None,
            ("Extract Page Range…", self.extract_page_range),
            ("Split Into Single Pages…", self.split_every_page),
        ])
        top.addWidget(self.pages_menu_button)

        self.export_menu_button = self._menu_button("Export", [
            ("Current Page as Image…", [
                ("PNG…", lambda: self.export_page_image("png")),
                ("JPEG…", lambda: self.export_page_image("jpg")),
            ]),
            ("All Pages as Images…", [
                ("PNG Folder…", lambda: self.export_all_pages_images("png")),
                ("JPEG Folder…", lambda: self.export_all_pages_images("jpg")),
            ]),
            None,
            ("Text", [
                ("Current Page…", self.export_current_page_text),
                ("Whole Document…", self.export_full_document_text),
            ]),
            ("Document As", [
                ("HTML…", self.export_full_document_html),
                ("Markdown…", self.export_full_document_markdown),
                ("Word DOCX…", self.export_full_document_docx),
            ]),
            None,
            ("Metadata as JSON…", self.export_metadata_json),
        ])
        top.addWidget(self.export_menu_button)

        spacer = QWidget()
        spacer.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        top.addWidget(spacer)

        self.xray_button = QToolButton()
        self.xray_button.setObjectName("ToolButton")
        self.xray_button.setText("X-ray")
        self.xray_button.setCheckable(True)
        self.xray_button.setToolTip(
            "Reveal what this page is really made of: editable text, images and vector art"
        )
        self.xray_button.toggled.connect(self.toggle_xray)
        top.addWidget(self.xray_button)

        self.spell_button = QToolButton()
        self.spell_button.setObjectName("ToolButton")
        self.spell_button.setText("Spell")
        self.spell_button.setCheckable(True)
        self.spell_button.setToolTip(
            "Spell check: underline misspelled words with a squiggly line"
        )
        self.spell_button.toggled.connect(self.toggle_spellcheck)
        top.addWidget(self.spell_button)
        top.addSeparator()

        self.zoom_out_action = self._add_action(top, "−", self.zoom_out, "Ctrl+-")
        self.zoom_combo = QComboBox()
        self.zoom_combo.setFixedWidth(104)
        self.zoom_combo.setToolTip("Zoom level")
        self.zoom_combo.addItems(["50%", "75%", "100%", "125%", "150%", "175%", "200%", "300%", "Fit Width", "Fit Page"])
        self.zoom_combo.setCurrentText("125%")
        self.zoom_combo.currentTextChanged.connect(self.handle_zoom_combo)
        top.addWidget(self.zoom_combo)
        self.zoom_in_action = self._add_action(top, "+", self.zoom_in, "Ctrl++")

        # Hand-drawn sun / moon dark-mode toggle, sitting right after the zoom
        # "+" button so it lives with the other trailing ribbon controls and
        # respects the toolbar's margins / spacing.
        self.theme_toggle = ThemeToggle()
        self.theme_toggle.clicked.connect(self.toggle_theme)
        self.theme_toggle.set_dark(CURRENT_THEME == "dark")
        top.addWidget(self.theme_toggle)

    def _menu_button(self, label: str, items) -> QToolButton:
        """Build a ribbon dropdown.

        Item forms, shared by every menu so they all behave identically:
            None                     -> separator
            (text, callable)         -> action
            (text, [sub-items...])   -> submenu
        """
        button = QToolButton()
        button.setText(f"{label} \u25be")
        button.setMinimumWidth(64)
        button.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        button.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        button.setMenu(self._build_menu(button, items))
        return button

    def _build_menu(self, parent: QWidget, items) -> QMenu:
        menu = SketchMenu(parent)
        for item in items:
            if item is None:
                menu.addSeparator()
                continue
            item_label, target = item
            if isinstance(target, (list, tuple)):
                submenu = self._build_menu(menu, target)
                submenu.setTitle(item_label)
                menu.addMenu(submenu)
            else:
                action = QAction(item_label, self)
                action.triggered.connect(target)
                menu.addAction(action)
        return menu

    def _add_action(self, toolbar: QToolBar, text: str, callback, shortcut: Optional[str] = None) -> QAction:
        action = QAction(text, self)
        action.triggered.connect(callback)
        if shortcut:
            action.setShortcut(shortcut)
        toolbar.addAction(action)
        return action

    def _build_search_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(8, 8, 8, 8)

        self.search_input = QLineEdit()
        self.search_input.setFixedHeight(31)
        self.search_input.setPlaceholderText("Find text...")
        self.search_input.returnPressed.connect(self.search_text)

        search_row = QHBoxLayout()
        search_row.setContentsMargins(0, 0, 0, 0)
        search_row.setSpacing(6)
        self.search_button = QPushButton("Find")
        self.search_button.setFixedHeight(31)
        self.search_button.clicked.connect(self.search_text)
        self.next_search_button = QPushButton("Next ▸")
        self.next_search_button.setFixedHeight(31)
        self.next_search_button.clicked.connect(self.go_to_next_search_result)
        search_row.addWidget(self.search_button)
        search_row.addWidget(self.next_search_button)

        self.search_results_list = QListWidget()
        self.search_results_list.itemClicked.connect(self.go_to_search_item)

        layout.addWidget(self.search_input)
        layout.addLayout(search_row)
        layout.addWidget(QLabel("Results"))
        layout.addWidget(self.search_results_list)
        return panel

    def _build_outline_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)
        layout.addWidget(self.outline_list)

        row1 = QHBoxLayout()
        row1.setContentsMargins(0, 0, 0, 0)
        row1.setSpacing(6)
        add_btn = QPushButton("Add")
        add_btn.setToolTip("Add a bookmark to the current page")
        add_btn.clicked.connect(self.add_bookmark)
        rename_btn = QPushButton("Rename")
        rename_btn.clicked.connect(self.rename_bookmark)
        del_btn = QPushButton("Delete")
        del_btn.clicked.connect(self.delete_bookmark)
        row1.addWidget(add_btn)
        row1.addWidget(rename_btn)
        row1.addWidget(del_btn)
        layout.addLayout(row1)

        row2 = QHBoxLayout()
        row2.setContentsMargins(0, 0, 0, 0)
        row2.setSpacing(6)
        up_btn = QPushButton("↑ Up")
        up_btn.clicked.connect(lambda: self.move_bookmark(-1))
        down_btn = QPushButton("↓ Down")
        down_btn.clicked.connect(lambda: self.move_bookmark(1))
        indent_btn = QPushButton("→ Indent")
        indent_btn.clicked.connect(lambda: self.indent_bookmark(1))
        outdent_btn = QPushButton("← Outdent")
        outdent_btn.clicked.connect(lambda: self.indent_bookmark(-1))
        row2.addWidget(up_btn)
        row2.addWidget(down_btn)
        layout.addLayout(row2)
        row3 = QHBoxLayout()
        row3.setContentsMargins(0, 0, 0, 0)
        row3.setSpacing(6)
        row3.addWidget(indent_btn)
        row3.addWidget(outdent_btn)
        layout.addLayout(row3)
        return panel

    def _build_notes_panel(self) -> QWidget:
        panel = QWidget()
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)
        layout.addWidget(self.comments_list)
        row = QHBoxLayout()
        row.setContentsMargins(0, 0, 0, 0)
        row.setSpacing(6)
        del_btn = QPushButton("Delete")
        del_btn.setToolTip("Delete the selected annotation")
        del_btn.clicked.connect(self.delete_selected_annotation)
        del_all_btn = QPushButton("Delete All")
        del_all_btn.clicked.connect(self.delete_all_annotations)
        row.addWidget(del_btn)
        row.addWidget(del_all_btn)
        layout.addLayout(row)
        return panel

    TOOL_SHORT_LABELS = {
        Tool.SELECT: "Select",
        Tool.TEXT_BOX: "Text Box",
        Tool.EDIT_TEXT: "Edit Span",
        Tool.EDIT_BLOCK: "Edit Block",
        Tool.MOVE_TEXT: "Move",
        Tool.NOTE: "Note",
        Tool.HIGHLIGHT: "Highlight",
        Tool.UNDERLINE: "Underline",
        Tool.STRIKEOUT: "Strike",
        Tool.RECTANGLE: "Rectangle",
        Tool.LINE: "Line",
        Tool.INK: "Freehand",
        Tool.REDACT: "Redact",
        Tool.CROP: "Crop",
        Tool.LINK: "Link",
        Tool.SIGN: "Signature",
    }

    TOOL_FIELDS = {
        Tool.TEXT_BOX: {"font", "text", "stroke"},
        Tool.EDIT_BLOCK: {"font", "text"},
        Tool.EDIT_TEXT: set(),
        Tool.MOVE_TEXT: set(),
        Tool.NOTE: set(),
        Tool.HIGHLIGHT: set(),
        Tool.UNDERLINE: {"stroke"},
        Tool.STRIKEOUT: {"stroke"},
        Tool.RECTANGLE: {"line", "opacity", "stroke", "fill"},
        Tool.LINE: {"line", "stroke"},
        Tool.INK: {"line", "stroke"},
        Tool.REDACT: set(),
        Tool.CROP: set(),
        Tool.LINK: set(),
        Tool.SELECT: set(),
        Tool.IMAGE: set(),
        Tool.SIGN: set(),
    }

    TOOL_HINTS = {
        Tool.SELECT: "Drag with the right or middle mouse button to pan. Ctrl+wheel zooms.",
        Tool.TEXT_BOX: "Drag a box, then type the text to draw inside it.",
        Tool.EDIT_TEXT: "Click a word to replace it. Its original size and colour are reused.",
        Tool.EDIT_BLOCK: "Click a paragraph to replace the whole block.",
        Tool.MOVE_TEXT: "Drag a signature, text box, or annotation. Double-click to recolour, Delete to remove.",
        Tool.NOTE: "Click to drop a sticky note comment.",
        Tool.HIGHLIGHT: "Drag across text to highlight it.",
        Tool.UNDERLINE: "Drag across text to underline it.",
        Tool.STRIKEOUT: "Drag across text to strike it through.",
        Tool.RECTANGLE: "Drag to draw a rectangle.",
        Tool.LINE: "Drag from start to end to draw a line.",
        Tool.INK: "Draw freehand with the left mouse button.",
        Tool.REDACT: "Drag an area to permanently remove its content.",
        Tool.CROP: "Drag the area to keep, then confirm the crop.",
        Tool.LINK: "Drag an area, then enter the URL it should open.",
        Tool.IMAGE: "Click where the chosen image should be placed.",
        Tool.SIGN: "Click where your signature should be placed.",
    }

    def _update_contextual_panel(self):
        """Show only the appearance rows the active tool uses."""
        fields = self.TOOL_FIELDS.get(self.current_tool, set())
        rows = {
            "font": self.font_size_spin,
            "line": self.line_width_spin,
            "opacity": self.opacity_spin,
            "text": self.text_color_button,
            "stroke": self.stroke_color_button,
            "fill": self.fill_color_button,
        }
        for name, widget in rows.items():
            index, _ = self.style_form.getWidgetPosition(widget)
            if index >= 0:
                self.style_form.setRowVisible(index, name in fields)
        self.style_group.setVisible(bool(fields))
        self.tool_hint.setText(self.TOOL_HINTS.get(self.current_tool, ""))

    def _make_tool_button(self, tool: str) -> QToolButton:
        button = QToolButton()
        button.setObjectName("ToolButton")
        button.setText(self.TOOL_SHORT_LABELS.get(tool, tool))
        button.setToolTip(tool)
        button.setCheckable(True)
        button.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        button.clicked.connect(lambda _checked=False, t=tool: self.set_tool(t))
        self.tool_buttons[tool] = button
        return button

    def _build_properties_panel(self) -> QWidget:
        panel = QWidget()
        panel.setObjectName("PropertiesPanel")
        layout = QVBoxLayout(panel)
        layout.setContentsMargins(8, 8, 8, 8)
        layout.setSpacing(6)

        title = QLabel("Tools")
        title.setObjectName("PanelTitle")
        layout.addWidget(title)

        self.tool_buttons: Dict[str, QToolButton] = {}
        self._tool_group = QActionGroup(self)
        self._tool_group.setExclusive(True)

        for group_name, tools in (
            ("Text", [Tool.EDIT_TEXT, Tool.EDIT_BLOCK, Tool.MOVE_TEXT, Tool.TEXT_BOX]),
            ("Markup", [Tool.HIGHLIGHT, Tool.UNDERLINE, Tool.STRIKEOUT, Tool.NOTE]),
            ("Draw", [Tool.RECTANGLE, Tool.LINE, Tool.INK]),
            ("Page", [Tool.REDACT, Tool.CROP, Tool.LINK]),
        ):
            group_box = QGroupBox(group_name)
            grid = QGridLayout(group_box)
            grid.setContentsMargins(8, 10, 8, 8)
            grid.setHorizontalSpacing(6)
            grid.setVerticalSpacing(6)
            for index, tool in enumerate(tools):
                grid.addWidget(self._make_tool_button(tool), index // 2, index % 2)
            layout.addWidget(group_box)

        select_button = self._make_tool_button(Tool.SELECT)
        select_button.setChecked(True)
        layout.insertWidget(1, select_button)

        self.tool_hint = QLabel()
        self.tool_hint.setObjectName("ToolHint")
        self.tool_hint.setWordWrap(True)
        layout.addWidget(self.tool_hint)

        self.style_group = QGroupBox("Appearance")
        style_group = self.style_group
        style_form = QFormLayout(style_group)
        self.style_form = style_form
        style_form.setContentsMargins(8, 10, 8, 8)
        style_form.setHorizontalSpacing(8)
        style_form.setVerticalSpacing(7)

        self.font_size_spin = QSpinBox()
        self.font_size_spin.setRange(6, 120)
        self.font_size_spin.setValue(14)

        self.line_width_spin = QSpinBox()
        self.line_width_spin.setRange(1, 20)
        self.line_width_spin.setValue(2)
        self.line_width_spin.valueChanged.connect(lambda value: setattr(self, "line_width", value))

        self.opacity_spin = QSpinBox()
        self.opacity_spin.setRange(5, 100)
        self.opacity_spin.setValue(35)
        self.opacity_spin.valueChanged.connect(lambda value: setattr(self, "opacity_percent", value))

        self.text_color_button = SketchSwatchButton("Text")
        self.text_color_button.clicked.connect(lambda: self.choose_color("text"))

        self.stroke_color_button = SketchSwatchButton("Stroke")
        self.stroke_color_button.clicked.connect(lambda: self.choose_color("stroke"))

        self.fill_color_button = SketchSwatchButton("Fill")
        self.fill_color_button.clicked.connect(lambda: self.choose_color("fill"))

        style_form.addRow("Font size:", self.font_size_spin)
        style_form.addRow("Line width:", self.line_width_spin)
        style_form.addRow("Opacity %:", self.opacity_spin)
        style_form.addRow("Text color:", self.text_color_button)
        style_form.addRow("Stroke color:", self.stroke_color_button)
        style_form.addRow("Fill color:", self.fill_color_button)
        layout.addWidget(style_group)


        layout.addStretch()
        self._refresh_color_buttons()
        return panel

    def _apply_theme(self):
        t = _THEMES.get(CURRENT_THEME, _THEMES["light"])
        paper = t["paper"].name()
        ink = t["ink"].name()
        muted = t["grey_muted"].name()
        faint = t["grey_faint"].name()
        # Text drawn on top of an INK-filled control (checked buttons, etc.)
        # must be the paper colour to stay readable in both themes.
        on = paper
        disabled = faint
        # A slightly lifted handle-hover derived from ink.
        hh = t["hover"].lighter(160).name() if CURRENT_THEME == "dark" else "#4a4a4a"
        sep = t["grey_faint"].name()
        self.setStyleSheet(f"""
            /* ============================================================
               Suketchi PDF Reader — hand-drawn. Ink / paper flip between the
               light and dark themes; the sketch character stays the same.
               ============================================================ */

            QMainWindow, QWidget {{
                background: {paper};
                color: {ink};
                font-size: 14px;
                font-weight: 400;
            }}

            QToolBar#Ribbon {{
                background: {paper};
                border: none;
                padding: 10px 14px;
                spacing: 6px;
            }}
            QToolBar#Ribbon::separator {{ background: transparent; width: 10px; }}

            /* Buttons: colour + spacing only. Outline comes from SketchStyle. */
            QToolButton, QPushButton {{
                color: {ink};
                padding: 7px 14px;
                min-height: 20px;
            }}
            QToolButton:disabled, QPushButton:disabled {{ color: {disabled}; }}
            QToolButton:checked, QPushButton:checked {{ color: {on}; }}
            QToolButton::menu-indicator {{ image: none; width: 0px; }}

            QToolButton#Primary, QPushButton#Primary {{ padding: 7px 18px; }}
            QToolButton#ToolButton {{ padding: 7px 8px; }}

            /* QComboBox is deliberately absent: any rule here would hand the
               widget to Qt's stylesheet engine and lose the hand-drawn box. */
            QSpinBox, QLineEdit {{
                color: {ink};
                padding: 6px 14px;
                min-height: 20px;
                selection-background-color: {ink};
                selection-color: {paper};
            }}
            QLineEdit::placeholder {{ color: {muted}; }}
            QSpinBox::up-button, QSpinBox::down-button {{ border: none; background: transparent; width: 0px; }}

            QComboBox QAbstractItemView {{
                background: {paper};
                color: {ink};
                border: 2px solid {ink};
                padding: 4px;
                outline: 0;
            }}

            /* SketchMenu paints the panel and the selection box itself, so
               item backgrounds stay transparent and text stays ink-coloured
               even when the row is active. */
            QMenu {{ background: transparent; border: none; padding: 7px; }}
            QMenu::item {{
                background: transparent;
                color: {ink};
                padding: 8px 30px 8px 16px;
                margin: 1px 2px;
            }}
            QMenu::item:selected {{ background: transparent; color: {ink}; }}
            QMenu::item:disabled {{ color: {disabled}; }}
            QMenu::separator {{ height: 1px; background: {sep}; margin: 6px 12px; }}

            QTabWidget#SidebarTabs {{ background: {paper}; }}
            QTabWidget::pane {{ background: {paper}; border: none; }}
            QTabBar::tab {{
                background: transparent;
                color: {muted};
                border: none;
                padding: 8px 7px;
                margin-right: 2px;
            }}
            QTabBar::tab:selected {{ color: {ink}; }}
            QTabBar::tab:hover:!selected {{ color: {ink}; }}

            QListWidget {{ background: {paper}; color: {ink}; border: none; outline: 0; }}
            /* Spacing only — the selection box itself is drawn by SketchStyle. */
            QListWidget::item {{ padding: 7px 8px; margin: 3px 6px; }}

            QTextEdit {{
                background: {paper};
                color: {ink};
                border: 2px solid {ink};
                padding: 7px;
                selection-background-color: {ink};
                selection-color: {paper};
            }}

            QGroupBox {{
                background: transparent;
                border: none;
                margin-top: 15px;
                padding: 10px 8px 8px 8px;
                color: {ink};
            }}
            QGroupBox::title {{
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 6px;
                color: {ink};
                font-size: 13px;
            }}

            QLabel#PanelTitle {{
                font-size: 20px;
                font-weight: 400;
                color: {ink};
                padding: 2px 2px 4px 2px;
            }}
            QLabel#ToolHint {{ color: {muted}; font-size: 13px; padding: 0 3px 6px 3px; }}
            QLabel#BrandText {{ color: {ink}; font-size: 19px; padding: 0 6px 0 0; }}
            QLabel#SmallNote {{ color: {muted}; font-size: 13px; }}

            QWidget#PropertiesPanel {{ background: {paper}; }}

            QSplitter::handle {{ background: {paper}; }}
            QSplitter::handle:horizontal {{ width: 2px; background: {ink}; }}

            QScrollArea {{ background: {paper}; border: none; }}
            QScrollBar:vertical, QScrollBar:horizontal {{
                background: transparent; border: none; margin: 2px;
                width: 12px; height: 12px;
            }}
            QScrollBar::handle:vertical, QScrollBar::handle:horizontal {{
                background: {ink}; border-radius: 6px; min-height: 34px; min-width: 34px;
            }}
            QScrollBar::handle:vertical:hover, QScrollBar::handle:horizontal:hover {{ background: {hh}; }}
            QScrollBar::add-line, QScrollBar::sub-line,
            QScrollBar::add-page, QScrollBar::sub-page {{
                width: 0px; height: 0px; background: none;
            }}

            QStatusBar {{ background: {paper}; color: {muted}; border: none; font-size: 13px; }}
            QStatusBar::item {{ border: none; }}

            QProgressBar {{
                background: {paper}; color: {ink};
                border: 2px solid {ink}; text-align: center;
                font-size: 12px; min-height: 16px;
            }}
            QProgressBar::chunk {{ background: {ink}; }}

            QDialog {{ background: {paper}; color: {ink}; }}
            QDialogButtonBox QPushButton {{ min-width: 88px; padding: 8px 18px; }}

            QCheckBox {{ color: {ink}; spacing: 8px; }}
            QCheckBox::indicator {{ width: 18px; height: 18px; background: transparent; border: none; }}

            QFormLayout QLabel {{ color: {ink}; font-size: 13px; }}

            QMessageBox {{ background: {paper}; color: {ink}; }}
            QMessageBox QPushButton {{ min-width: 80px; }}

            QToolTip {{
                background: {paper}; color: {ink};
                border: 2px solid {ink}; padding: 6px 9px;
            }}
        """)

    def toggle_theme(self):
        self.set_theme("light" if CURRENT_THEME == "dark" else "dark")

    def set_theme(self, name: str):
        """Flip the whole hand-drawn UI between light and dark by inverting the
        shared ink / paper colours, then re-applying the stylesheet, palette and
        repainting every widget."""
        apply_theme_colors(name)
        app = QApplication.instance()
        if app is not None:
            _apply_bw_palette(app)
        self._apply_theme()
        if getattr(self, "theme_toggle", None) is not None:
            self.theme_toggle.set_dark(name == "dark")
        # Repaint everything: custom-painted widgets hold INK / PAPER refs that
        # now point at the new colours.
        self.canvas.update()
        self.tab_bar.update()
        self.update()
        for w in self.findChildren(QWidget):
            w.update()

    def _set_document_controls(self, enabled: bool):
        widgets = [
            self.save_action,
            self.undo_action,
            self.redo_action,
            self.zoom_out_action,
            self.zoom_in_action,
            self.file_menu_button,
            self.edit_menu_button,
            self.insert_menu_button,
            self.pages_menu_button,
            self.export_menu_button,
            self.zoom_combo,
            self.search_input,
            self.search_button,
            self.next_search_button,
            self.font_size_spin,
            self.line_width_spin,
            self.opacity_spin,
            self.text_color_button,
            self.stroke_color_button,
            self.fill_color_button,
        ]
        widgets.append(self.xray_button)
        widgets.append(self.spell_button)
        widgets.extend(self.tool_buttons.values())
        for widget in widgets:
            widget.setEnabled(enabled)

        self.open_action.setEnabled(True)
        self._refresh_history_actions()

    def open_pdf(self):
        # Opening a document now adds a tab rather than replacing the current
        # one, so no discard confirmation is needed here.
        path, _ = QFileDialog.getOpenFileName(self, "Open PDF", "", "PDF Files (*.pdf)")
        if not path:
            return
        self.load_pdf(path)

    def load_pdf(self, path: str):
        """Open a PDF from a known path (file dialog, drag-and-drop, CLI) in a
        new tab. If the file is already open, just switch to its tab."""
        # Already open? Focus that tab instead of opening a duplicate.
        for i, tab in enumerate(self.tabs):
            if tab.file_path and os.path.abspath(tab.file_path) == os.path.abspath(path):
                self.switch_to_tab(i)
                return

        try:
            doc = fitz.open(path)
            opened_password = ""

            if doc.needs_pass:
                dialog = PasswordDialog(self)
                if dialog.exec() != QDialog.DialogCode.Accepted:
                    doc.close()
                    return
                opened_password = dialog.value()
                if not doc.authenticate(opened_password):
                    doc.close()
                    QMessageBox.warning(self, "Wrong Password", "The password was not accepted.")
                    return

            if doc.page_count == 0:
                raise ValueError("The selected PDF has no pages.")

            # Snapshot the currently active tab before switching context.
            self._snapshot_active_tab()

            tab = DocumentTab(doc, path, opened_password)
            self.tabs.append(tab)
            self.active_tab_index = len(self.tabs) - 1
            self._restore_active_tab()
            self._refresh_tab_bar()
            self.status.showMessage(f"Opened {path}")

        except Exception as exc:
            QMessageBox.critical(self, "Open Error", f"Could not open PDF:\n{exc}")

    # ---- Tab management ---------------------------------------------------
    def _snapshot_active_tab(self):
        """Copy the window's live per-document fields back into the active tab."""
        if not (0 <= self.active_tab_index < len(self.tabs)):
            return
        tab = self.tabs[self.active_tab_index]
        tab.doc = self.doc
        tab.file_path = self.file_path
        tab.password = self._pdf_password
        tab.current_page_index = self.current_page_index
        tab.zoom = self.zoom
        tab.is_dirty = self.is_dirty
        tab.search_results = list(self.search_results)
        tab.search_index = self.search_index
        tab.undo_stack = list(self.undo_stack)
        tab.redo_stack = list(self.redo_stack)
        tab.doc_version = self._doc_version
        tab.render_cache = self._render_cache
        tab.render_cache_order = self._render_cache_order
        tab.xray_cache = self._xray_cache
        tab.page_structure_changed = self._page_structure_changed
        tab.signatures = self.signatures

    def _restore_active_tab(self):
        """Load the active tab's state into the window's live fields and paint."""
        if not (0 <= self.active_tab_index < len(self.tabs)):
            self._clear_to_empty_state()
            return
        tab = self.tabs[self.active_tab_index]
        self.doc = tab.doc
        self.file_path = tab.file_path
        self._pdf_password = tab.password
        self.current_page_index = tab.current_page_index
        self.zoom = tab.zoom
        self.is_dirty = tab.is_dirty
        self.search_results = tab.search_results
        self.search_index = tab.search_index
        self.undo_stack = tab.undo_stack
        self.redo_stack = tab.redo_stack
        self._doc_version = tab.doc_version
        self._render_cache = tab.render_cache
        self._render_cache_order = tab.render_cache_order
        self._xray_cache = tab.xray_cache
        self._page_structure_changed = tab.page_structure_changed
        self.signatures = getattr(tab, "signatures", [])
        # Spell-check rects are per (doc-version, page, zoom); clear on switch so
        # a different document never reuses another's cached rects.
        self._spell_cache = {}

        self.search_results_list.clear()
        self._set_document_controls(True)
        self.refresh_sidebars()
        if tab.needs_fit_width:
            # First time this document is shown: default to Fit Page.
            tab.needs_fit_width = False
            self.zoom = tab.zoom
            self.fit_page()
        else:
            self.render_current_page()
        self._update_window_title()

    def _clear_to_empty_state(self):
        """No documents open: reset to the initial blank/no-doc state."""
        self.doc = None
        self.file_path = None
        self._pdf_password = ""
        self.current_page_index = 0
        self.zoom = 1.25
        self.is_dirty = False
        self.search_results = []
        self.search_index = -1
        self.undo_stack = []
        self.redo_stack = []
        self._doc_version = 0
        self._render_cache = {}
        self._render_cache_order = []
        self._xray_cache = {}
        self._page_structure_changed = False
        self.signatures = []
        self.search_results_list.clear()
        self.page_list.clear()
        self.outline_list.clear()
        self.comments_list.clear()
        self.canvas.set_page(QPixmap(), self.zoom)
        self._set_document_controls(False)
        self._update_window_title()

    def switch_to_tab(self, index: int):
        if index == self.active_tab_index or not (0 <= index < len(self.tabs)):
            return
        self._snapshot_active_tab()
        self.active_tab_index = index
        self._restore_active_tab()
        self._refresh_tab_bar()

    def close_tab(self, index: int):
        if not (0 <= index < len(self.tabs)):
            return
        # Make the target tab active so dirty checks apply to it.
        if index != self.active_tab_index:
            self.switch_to_tab(index)
        if self.is_dirty and not self._confirm_discard():
            return

        tab = self.tabs.pop(index)
        try:
            if tab.doc is not None:
                tab.doc.close()
        except Exception:
            pass

        if not self.tabs:
            self.active_tab_index = -1
            self._clear_to_empty_state()
            self._refresh_tab_bar()
            return

        # Pick a sensible neighbouring tab.
        self.active_tab_index = min(index, len(self.tabs) - 1)
        self._restore_active_tab()
        self._refresh_tab_bar()

    def _refresh_tab_bar(self):
        titles = [t.title() + (" •" if t.is_dirty else "") for t in self.tabs]
        # For the active tab, read the live window fields so a rename ("Save
        # As") or an unsaved-edit marker shows immediately, before any snapshot.
        if 0 <= self.active_tab_index < len(titles):
            base = Path(self.file_path).name if self.file_path else "Untitled"
            titles[self.active_tab_index] = base + (" •" if self.is_dirty else "")
        self.tab_bar.set_tabs(titles, self.active_tab_index)
        self.tab_bar.setVisible(len(self.tabs) > 0)



    def save_as_pdf(self):
        if self.doc is None:
            return

        default = "edited.pdf"
        if self.file_path:
            src = Path(self.file_path)
            default = str(src.with_name(f"{src.stem}_edited.pdf"))

        out, _ = QFileDialog.getSaveFileName(self, "Save Optimized PDF As", default, "PDF Files (*.pdf)")
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"

        try:
            if self.file_path and os.path.abspath(out) == os.path.abspath(self.file_path):
                QMessageBox.warning(self, "Use Different Name", "Choose a different output file name.")
                return
            self.doc.save(out, garbage=4, deflate=True, clean=True)
            self.file_path = out
            self.is_dirty = False
            self._update_window_title()
            self.status.showMessage(f"Saved {out}")
            QMessageBox.information(self, "Saved", f"PDF saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "Save Error", f"Could not save PDF:\n{exc}")

    def save_encrypted_copy(self):
        if self.doc is None:
            return

        user_pw, ok = QInputDialog.getText(self, "User Password", "Password required to open PDF:", QLineEdit.EchoMode.Password)
        if not ok:
            return
        owner_pw, ok = QInputDialog.getText(self, "Owner Password", "Owner password for permissions:", QLineEdit.EchoMode.Password)
        if not ok:
            return

        out, _ = QFileDialog.getSaveFileName(self, "Save Encrypted PDF", "encrypted.pdf", "PDF Files (*.pdf)")
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"

        try:
            self.doc.save(
                out,
                encryption=fitz.PDF_ENCRYPT_AES_256,
                user_pw=user_pw,
                owner_pw=owner_pw or user_pw,
                garbage=4,
                deflate=True,
            )
            QMessageBox.information(self, "Saved", f"Encrypted PDF saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "Encryption Error", f"Could not save encrypted copy:\n{exc}")

    def digitally_sign(self):
        """Apply a cryptographic (PKCS#7) digital signature using a PKCS#12
        (.p12/.pfx) certificate. The signed copy is written to a new file."""
        if self.doc is None:
            return
        if not pyhanko_available():
            QMessageBox.information(
                self, "Digital Signing Unavailable",
                "Cryptographic signing needs the pyHanko package.\n\n"
                "Install it with:\n    pip install pyHanko",
            )
            return

        cert_path, _ = QFileDialog.getOpenFileName(
            self, "Select Signing Certificate", "",
            "PKCS#12 Certificates (*.p12 *.pfx)")
        if not cert_path:
            return
        passphrase, ok = QInputDialog.getText(
            self, "Certificate Password", "Password for the certificate:",
            QLineEdit.EchoMode.Password)
        if not ok:
            return

        default = "signed.pdf"
        if self.file_path:
            src = Path(self.file_path)
            default = str(src.with_name(f"{src.stem}_signed.pdf"))
        out, _ = QFileDialog.getSaveFileName(self, "Save Signed PDF As", default, "PDF Files (*.pdf)")
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"

        try:
            from pyhanko.sign import signers
            from pyhanko.sign.signers import PdfSignatureMetadata
            from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter
            import io

            signer = signers.SimpleSigner.load_pkcs12(
                cert_path, passphrase=passphrase.encode("utf-8") if passphrase else None)
            if signer is None:
                raise ValueError("Could not load the certificate (wrong password?).")

            # Sign the current in-memory document (captures unsaved edits).
            pdf_bytes = self.doc.tobytes(garbage=4, deflate=True)
            reader = IncrementalPdfFileWriter(io.BytesIO(pdf_bytes))
            # Count existing signatures to pick a fresh, unique field name.
            try:
                from pyhanko.pdf_utils.reader import PdfFileReader
                existing = len(list(PdfFileReader(io.BytesIO(pdf_bytes)).embedded_signatures))
            except Exception:
                existing = 0
            field_name = "Signature%d" % (existing + 1)
            meta = PdfSignatureMetadata(field_name=field_name, reason="Approved", location="")
            with open(out, "wb") as outf:
                signers.sign_pdf(reader, meta, signer=signer, output=outf)

            QMessageBox.information(
                self, "Digitally Signed",
                f"Signed PDF saved:\n{out}\n\nField: {field_name}")
            self.status.showMessage(f"Digitally signed: {out}")
        except Exception as exc:
            QMessageBox.critical(self, "Signing Error", f"Could not sign the PDF:\n{exc}")

    def verify_signatures(self):
        """Report the cryptographic status of any digital signatures in the
        current PDF (intact / valid / trusted)."""
        if self.doc is None:
            return
        if not pyhanko_available():
            QMessageBox.information(
                self, "Verification Unavailable",
                "Signature verification needs the pyHanko package.\n\n"
                "Install it with:\n    pip install pyHanko",
            )
            return
        try:
            from pyhanko.pdf_utils.reader import PdfFileReader
            from pyhanko.sign.validation import validate_pdf_signature
            from pyhanko_certvalidator import ValidationContext
            import io

            # Verify the bytes exactly as stored on disk — re-serialising via
            # PyMuPDF would change the byte ranges and break the integrity check.
            if self.is_dirty or not self.file_path or not os.path.exists(self.file_path):
                QMessageBox.information(
                    self, "Verify Signatures",
                    "Please save (or open) the PDF file first, then verify.\n\n"
                    "Signatures must be checked against the exact bytes on disk.")
                return
            with open(self.file_path, "rb") as fh:
                reader = PdfFileReader(io.BytesIO(fh.read()))
            sigs = list(reader.embedded_signatures)
            if not sigs:
                QMessageBox.information(self, "Verify Signatures", "This PDF has no digital signatures.")
                return

            # No external trust roots configured → trust status will be False,
            # but integrity (intact) and validity are still meaningful.
            vc = ValidationContext(allow_fetching=False)
            lines = []
            for s in sigs:
                try:
                    status = validate_pdf_signature(s, vc)
                    signer_name = ""
                    try:
                        signer_name = status.signing_cert.subject.human_friendly
                    except Exception:
                        pass
                    lines.append(
                        f"• {s.field_name}: "
                        f"{'intact' if status.intact else 'MODIFIED'}, "
                        f"{'valid' if status.valid else 'invalid'}"
                        + (f"\n    Signer: {signer_name}" if signer_name else "")
                    )
                except Exception as exc:
                    lines.append(f"• {s.field_name}: could not validate ({exc})")

            QMessageBox.information(
                self, "Digital Signatures",
                "\n".join(lines) + "\n\n(Trust depends on your installed certificate authorities.)")
        except Exception as exc:
            QMessageBox.critical(self, "Verification Error", f"Could not verify signatures:\n{exc}")

    def _document_worker_source(self) -> Optional[Dict[str, object]]:
        """Create a thread-safe document source for background workers.

        A PyMuPDF Document must not be shared across threads. Workers either reopen the
        original file path or open a byte snapshot when the document has unsaved edits.
        """
        if self.doc is None:
            return None
        if self.file_path and not self.is_dirty:
            return {"type": "path", "path": self.file_path, "password": self._pdf_password}
        try:
            return {"type": "bytes", "data": self.doc.tobytes(garbage=1, deflate=True), "password": self._pdf_password}
        except Exception as exc:
            QMessageBox.critical(self, "Worker Error", f"Could not prepare PDF for background task:\n{exc}")
            return None

    def _start_worker(self, worker: QObject):
        thread = QThread(self)
        worker.moveToThread(thread)
        thread.started.connect(worker.run)
        thread.finished.connect(thread.deleteLater)
        worker.finished.connect(thread.quit)
        worker.finished.connect(worker.deleteLater)
        self._workers.append((thread, worker))

        def cleanup():
            self._workers[:] = [(t, w) for t, w in self._workers if t is not thread]
            if not self._workers:
                self.progress_bar.setVisible(False)

        thread.finished.connect(cleanup)
        thread.start()

    def _worker_progress(self, job_id: int, done: int, total: int, label: str):
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, max(total, 1))
        self.progress_bar.setValue(done)
        self.status.showMessage(f"{label}: {done}/{total}")

    def _worker_finished(self, job_id: int, message: str):
        self.status.showMessage(message)

    def _apply_thumbnail_from_worker(self, job_id: int, page_index: int, image_bytes: bytes):
        if job_id != self._thumbnail_job_id or self.doc is None:
            return
        if page_index < 0 or page_index >= self.page_list.count():
            return
        pixmap = QPixmap()
        if pixmap.loadFromData(image_bytes):
            self.page_list.item(page_index).setIcon(QIcon(pixmap))

    def _search_finished_from_worker(self, job_id: int, query: str, raw_results: object):
        if job_id != self._search_job_id or self.doc is None:
            return

        self.search_results.clear()
        self.search_results_list.clear()
        for page_index, x0, y0, x1, y1 in raw_results:
            self.search_results.append((int(page_index), fitz.Rect(x0, y0, x1, y1)))
            item = QListWidgetItem(f"Page {int(page_index) + 1}: {query}")
            item.setData(Qt.ItemDataRole.UserRole, len(self.search_results) - 1)
            self.search_results_list.addItem(item)

        if not self.search_results:
            self.render_current_page()
            QMessageBox.information(self, "Search", "No matches found.")
            return

        self.go_to_search_result(0)
        self.status.showMessage(f"Found {len(self.search_results)} result(s) for '{query}'")

    def _image_export_finished_from_worker(self, job_id: int, folder: str, count: int):
        if job_id != self._export_job_id:
            return
        QMessageBox.information(self, "Exported", f"Exported {count} page image(s) to:\n{folder}")

    def _compute_xray(self, page) -> Dict[str, List[QRectF]]:
        """Classify page content into text / image / vector regions.

        Cached per (document version, page, zoom) because get_drawings() is
        expensive on vector-heavy pages.
        """
        key = (self._doc_version, self.current_page_index, int(self.zoom * 1000))
        cached = self._xray_cache.get(key)
        if cached is not None:
            return cached

        text_rects: List[QRectF] = []
        image_rects: List[QRectF] = []
        vector_rects: List[QRectF] = []

        try:
            info = page.get_text("dict")
            for block in info.get("blocks", []):
                if block.get("type") == 0:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            if span.get("text", "").strip():
                                text_rects.append(self._pdf_rect_to_image_rect(fitz.Rect(span["bbox"])))
                elif block.get("bbox"):
                    image_rects.append(self._pdf_rect_to_image_rect(fitz.Rect(block["bbox"])))
        except Exception:
            pass

        try:
            for drawing in page.get_drawings():
                rect = drawing.get("rect")
                if rect is None or (rect.width < 2 and rect.height < 2):
                    continue
                if rect.width >= page.rect.width * 0.98 and rect.height >= page.rect.height * 0.98:
                    continue
                vector_rects.append(self._pdf_rect_to_image_rect(rect))
        except Exception:
            pass

        result = {"text": text_rects, "image": image_rects, "vector": vector_rects}
        self._xray_cache[key] = result
        if len(self._xray_cache) > 8:
            self._xray_cache.pop(next(iter(self._xray_cache)))
        return result

    def toggle_xray(self, enabled: bool):
        self.xray_enabled = bool(enabled)
        self.canvas.set_xray_enabled(self.xray_enabled)
        if self.doc is not None:
            self.render_current_page()
        if self.xray_enabled:
            self.status.showMessage("X-ray on — showing what this page actually contains")
        else:
            self.status.showMessage("X-ray off")

    def toggle_spellcheck(self, enabled: bool):
        enabled = bool(enabled)
        if enabled and not spell_checker_available():
            self.spell_button.blockSignals(True)
            self.spell_button.setChecked(False)
            self.spell_button.blockSignals(False)
            QMessageBox.information(
                self, "Spell Check Unavailable",
                "No dictionary is available for spell checking.\n\n"
                "Install one with:\n    pip install pyspellchecker",
            )
            return
        self.spellcheck_enabled = enabled
        if self.doc is not None:
            self.render_current_page()
        if enabled:
            self.status.showMessage("Spell check on — misspelled words are underlined")
        else:
            self.status.showMessage("Spell check off")

    def _compute_spellcheck(self, page) -> List[QRectF]:
        """Return image-space rects for every misspelled word on the page.

        Cached per (document version, page, zoom) since word extraction and
        dictionary lookups are relatively expensive.
        """
        key = (self._doc_version, self.current_page_index, int(self.zoom * 1000))
        cached = self._spell_cache.get(key)
        if cached is not None:
            return cached

        rects: List[QRectF] = []
        try:
            # words: (x0, y0, x1, y1, "word", block, line, word_no)
            for w in page.get_text("words"):
                x0, y0, x1, y1, token = w[0], w[1], w[2], w[3], w[4]
                # A PDF "word" may bundle punctuation; check each alpha run.
                if not any(is_word_misspelled(m.group(0)) for m in _WORD_RE.finditer(token)):
                    continue
                rects.append(self._pdf_rect_to_image_rect(fitz.Rect(x0, y0, x1, y1)))
        except Exception:
            pass

        self._spell_cache[key] = rects
        if len(self._spell_cache) > 12:
            self._spell_cache.pop(next(iter(self._spell_cache)))
        return rects

    def render_current_page(self):
        if self.doc is None:
            return

        try:
            page = self.doc[self.current_page_index]
            cache_key = (self._doc_version, self.current_page_index, int(self.zoom * 1000))
            qpix = self._render_cache.get(cache_key)
            if qpix is None:
                pix = page.get_pixmap(matrix=fitz.Matrix(self.zoom, self.zoom), alpha=False)
                image = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format.Format_RGB888).copy()
                qpix = QPixmap.fromImage(image)
                self._render_cache[cache_key] = qpix
                self._render_cache_order.append(cache_key)
                while len(self._render_cache_order) > 8:
                    old_key = self._render_cache_order.pop(0)
                    self._render_cache.pop(old_key, None)

            search_rects = [
                QRectF(r.x0 * self.zoom, r.y0 * self.zoom, r.width * self.zoom, r.height * self.zoom)
                for p, r in self.search_results
                if p == self.current_page_index
            ]

            edit_rects: List[QRectF] = []
            if self.current_tool == Tool.EDIT_TEXT:
                for span in self._iter_text_spans(page):
                    edit_rects.append(self._pdf_rect_to_image_rect(span.bbox))
            elif self.current_tool == Tool.EDIT_BLOCK:
                for block in self._iter_text_blocks(page):
                    edit_rects.append(self._pdf_rect_to_image_rect(block.bbox))
            if self.current_tool == Tool.MOVE_TEXT:
                # The Move tool only moves items you added; outline just those
                # (signatures and text boxes) on this page so they are easy to
                # grab, and nothing else.
                for sig in self.signatures:
                    if sig.get("page") == self.current_page_index:
                        edit_rects.append(self._pdf_rect_to_image_rect(fitz.Rect(sig["rect"])))

            crop_preview = None
            if page.cropbox != page.mediabox:
                crop_preview = self._pdf_rect_to_image_rect(page.cropbox)

            xray = self._compute_xray(page) if self.xray_enabled else None
            spell_rects = self._compute_spellcheck(page) if self.spellcheck_enabled else None

            self.canvas.set_style(self.annotation_color, self.fill_color, self.text_color, self.line_width)
            self.canvas.set_page(qpix, self.zoom, search_rects, edit_rects, crop_preview, xray, spell_rects)

            self.page_list.blockSignals(True)
            self.page_list.setCurrentRow(self.current_page_index)
            self.page_list.blockSignals(False)

            self._update_zoom_combo_text()
            self.status.showMessage(
                f"Page {self.current_page_index + 1}/{self.doc.page_count} | "
                f"Zoom {int(self.zoom * 100)}% | Tool: {self.current_tool}"
            )
        except Exception as exc:
            QMessageBox.critical(self, "Render Error", f"Could not render page:\n{exc}")

    def refresh_sidebars(self):
        self.refresh_page_thumbnails()
        self.refresh_outline()
        self.refresh_comments()

    def _refresh_light_sidebars(self):
        """Refresh cheap sidebars only. Do not restart the thumbnail worker."""
        self.refresh_outline()
        self.refresh_comments()

    def _update_current_thumbnail_inline(self):
        """Refresh only the visible current-page thumbnail after an annotation/edit."""
        if self.doc is None or self.current_page_index < 0 or self.current_page_index >= self.page_list.count():
            return
        try:
            page = self.doc[self.current_page_index]
            pix = page.get_pixmap(matrix=fitz.Matrix(0.17, 0.17), alpha=False)
            image = QImage(pix.samples, pix.width, pix.height, pix.stride, QImage.Format.Format_RGB888).copy()
            self.page_list.item(self.current_page_index).setIcon(QIcon(QPixmap.fromImage(image)))
        except Exception:
            pass

    def _mark_page_structure_changed(self):
        self._page_structure_changed = True

    def refresh_page_thumbnails(self):
        """Load thumbnails in a worker thread so large PDFs do not freeze the UI."""
        if self.doc is None:
            return

        self._thumbnail_job_id += 1
        job_id = self._thumbnail_job_id

        self.page_list.blockSignals(True)
        self.page_list.clear()
        for i in range(self.doc.page_count):
            item = QListWidgetItem(f"Page {i + 1}")
            item.setData(Qt.ItemDataRole.UserRole, i)
            self.page_list.addItem(item)
        self.page_list.setCurrentRow(self.current_page_index)
        self.page_list.blockSignals(False)

        source = self._document_worker_source()
        if not source:
            return

        worker = ThumbnailWorker(job_id, source, thumb_zoom=0.17)
        worker.thumbnail_ready.connect(self._apply_thumbnail_from_worker)
        worker.progress.connect(lambda jid, done, total: self._worker_progress(jid, done, total, "Loading thumbnails"))
        worker.error.connect(lambda jid, msg: self.status.showMessage(f"Thumbnail loading failed: {msg}"))
        worker.finished.connect(lambda jid: self._worker_finished(jid, "Thumbnails loaded"))
        self._start_worker(worker)

    def refresh_outline(self):
        self.outline_list.clear()
        if self.doc is None:
            return
        try:
            toc = self.doc.get_toc()
            if not toc:
                self.outline_list.addItem("No outline/bookmarks found.")
                return
            for idx, (level, title, page_num) in enumerate(toc):
                item = QListWidgetItem(("   " * max(level - 1, 0)) + "• " + title)
                # UserRole = target page (0-based); UserRole+1 = TOC row index.
                item.setData(Qt.ItemDataRole.UserRole, max(page_num - 1, 0))
                item.setData(Qt.ItemDataRole.UserRole + 1, idx)
                self.outline_list.addItem(item)
        except Exception:
            self.outline_list.addItem("Could not read outline.")

    def _current_toc(self) -> List[list]:
        try:
            return self.doc.get_toc() if self.doc is not None else []
        except Exception:
            return []

    def _selected_toc_index(self) -> int:
        item = self.outline_list.currentItem()
        if item is None:
            return -1
        idx = item.data(Qt.ItemDataRole.UserRole + 1)
        return idx if isinstance(idx, int) else -1

    def _apply_toc(self, toc: List[list], keep_index: int = -1):
        try:
            self.doc.set_toc(toc)
        except Exception as exc:
            QMessageBox.critical(self, "Bookmarks", f"Could not update bookmarks:\n{exc}")
            return
        self._mark_dirty("Bookmarks updated")
        self.refresh_outline()
        if 0 <= keep_index < self.outline_list.count():
            self.outline_list.setCurrentRow(keep_index)

    def add_bookmark(self):
        if self.doc is None:
            return
        title, ok = QInputDialog.getText(
            self, "Add Bookmark", "Bookmark title:",
            text=f"Page {self.current_page_index + 1}")
        if not ok or not title.strip():
            return
        toc = self._current_toc()
        # Insert after the selected entry (same level), else append at level 1.
        sel = self._selected_toc_index()
        entry = [1, title.strip(), self.current_page_index + 1]
        if 0 <= sel < len(toc):
            entry[0] = toc[sel][0]
            toc.insert(sel + 1, entry)
            keep = sel + 1
        else:
            toc.append(entry)
            keep = len(toc) - 1
        self._apply_toc(toc, keep)

    def rename_bookmark(self):
        if self.doc is None:
            return
        toc = self._current_toc()
        sel = self._selected_toc_index()
        if not (0 <= sel < len(toc)):
            QMessageBox.information(self, "Rename Bookmark", "Select a bookmark first.")
            return
        title, ok = QInputDialog.getText(self, "Rename Bookmark", "New title:", text=toc[sel][1])
        if not ok or not title.strip():
            return
        toc[sel][1] = title.strip()
        self._apply_toc(toc, sel)

    def delete_bookmark(self):
        if self.doc is None:
            return
        toc = self._current_toc()
        sel = self._selected_toc_index()
        if not (0 <= sel < len(toc)):
            QMessageBox.information(self, "Delete Bookmark", "Select a bookmark first.")
            return
        toc.pop(sel)
        toc = self._normalise_toc_levels(toc)
        self._apply_toc(toc, min(sel, len(toc) - 1))

    def move_bookmark(self, direction: int):
        if self.doc is None:
            return
        toc = self._current_toc()
        sel = self._selected_toc_index()
        j = sel + direction
        if not (0 <= sel < len(toc)) or not (0 <= j < len(toc)):
            return
        toc[sel], toc[j] = toc[j], toc[sel]
        # Repair the level hierarchy: the first entry must be level 1 and no
        # entry may be more than one level deeper than the one before it, or
        # set_toc rejects it.
        toc = self._normalise_toc_levels(toc)
        self._apply_toc(toc, j)

    @staticmethod
    def _normalise_toc_levels(toc: List[list]) -> List[list]:
        prev = 0
        for entry in toc:
            entry[0] = max(1, min(entry[0], prev + 1))
            prev = entry[0]
        return toc

    def indent_bookmark(self, delta: int):
        if self.doc is None:
            return
        toc = self._current_toc()
        sel = self._selected_toc_index()
        if not (0 <= sel < len(toc)):
            return
        new_level = toc[sel][0] + delta
        # Level 1 is the minimum; a child can be at most one deeper than the
        # entry above it (PDF TOC rule).
        max_level = (toc[sel - 1][0] + 1) if sel > 0 else 1
        new_level = max(1, min(new_level, max_level))
        toc[sel][0] = new_level
        self._apply_toc(toc, sel)

    def refresh_comments(self):
        self.comments_list.clear()
        if self.doc is None:
            return

        found = 0
        for page_index in range(self.doc.page_count):
            page = self.doc[page_index]
            for ordinal, annot in enumerate(page.annots()):
                found += 1
                try:
                    atype = annot.type[1]
                    info = annot.info or {}
                except Exception:
                    continue
                text = f"Page {page_index + 1}: {atype}"
                content = info.get("content") or info.get("title") or ""
                if content:
                    text += f" — {content[:60]}"
                item = QListWidgetItem(text)
                item.setData(Qt.ItemDataRole.UserRole, page_index)
                # Store the annotation's ORDINAL position on its page. Xrefs are
                # renumbered when the document is re-serialised (undo snapshot),
                # but the ordinal order is stable, so we delete by position.
                item.setData(Qt.ItemDataRole.UserRole + 1, ordinal)
                self.comments_list.addItem(item)
        if found == 0:
            self.comments_list.addItem("No comments or annotations found.")

    def on_page_selected(self, row: int):
        if self.doc is None or row < 0 or row >= self.doc.page_count:
            return
        self.current_page_index = row
        self.render_current_page()

    def go_to_outline_item(self, item: QListWidgetItem):
        if self.doc is None:
            return
        page_index = item.data(Qt.ItemDataRole.UserRole)
        if isinstance(page_index, int):
            self.current_page_index = clamp(page_index, 0, self.doc.page_count - 1)
            self.render_current_page()

    def go_to_comment_item(self, item: QListWidgetItem):
        if self.doc is None:
            return
        page_index = item.data(Qt.ItemDataRole.UserRole)
        if isinstance(page_index, int):
            self.current_page_index = page_index
            self.render_current_page()

    def delete_selected_annotation(self):
        if self.doc is None:
            return
        item = self.comments_list.currentItem()
        if item is None:
            QMessageBox.information(self, "Delete Annotation", "Select an annotation in the list first.")
            return
        page_index = item.data(Qt.ItemDataRole.UserRole)
        ordinal = item.data(Qt.ItemDataRole.UserRole + 1)
        if not isinstance(page_index, int) or not isinstance(ordinal, int) or ordinal < 0:
            return
        self._push_undo()
        page = self.doc[page_index]
        removed = False
        try:
            annots = list(page.annots())
            if 0 <= ordinal < len(annots):
                page.delete_annot(annots[ordinal])
                removed = True
        except Exception:
            removed = False
        try:
            self.doc.reload_page(page)
        except Exception:
            pass
        if removed:
            self._mark_dirty("Annotation deleted", refresh_sidebars=True)
        else:
            QMessageBox.information(self, "Delete Annotation", "That annotation could not be found (it may have already been removed).")

    def delete_all_annotations(self):
        if self.doc is None:
            return
        total = sum(1 for pi in range(self.doc.page_count) for _ in self.doc[pi].annots())
        if total == 0:
            QMessageBox.information(self, "Delete All Annotations", "This document has no annotations.")
            return
        if QMessageBox.question(
            self, "Delete All Annotations",
            f"Remove all {total} annotation(s) from the document?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        ) != QMessageBox.StandardButton.Yes:
            return
        self._push_undo()
        for pi in range(self.doc.page_count):
            page = self.doc[pi]
            # Repeatedly delete the first annotation until none remain; fetching
            # a fresh list each time avoids stale handles after mutation.
            guard = 0
            while True:
                annots = list(page.annots())
                if not annots or guard > 10000:
                    break
                try:
                    page.delete_annot(annots[0])
                except Exception:
                    break
                guard += 1
            try:
                self.doc.reload_page(page)
            except Exception:
                pass
        self._mark_dirty("All annotations deleted", refresh_sidebars=True)

    def set_tool(self, tool: str):
        self.current_tool = tool
        # Drop any pending placement payloads if the user leaves that tool, so a
        # stale signature/image is not dropped by a later click.
        if tool != Tool.SIGN:
            self.pending_signature = None
        if tool != Tool.IMAGE:
            self.pending_image_path = None
        self.canvas.set_tool(tool)
        for name, button in self.tool_buttons.items():
            button.blockSignals(True)
            button.setChecked(name == tool)
            button.blockSignals(False)
        self._update_contextual_panel()
        self.render_current_page()
        self.status.showMessage(f"{tool} selected")

    def choose_color(self, target: str):
        current = self.text_color if target == "text" else self.annotation_color if target == "stroke" else self.fill_color
        color = QColorDialog.getColor(current, self, f"Choose {target.title()} Color")
        if not color.isValid():
            return
        if target == "text":
            self.text_color = color
        elif target == "stroke":
            self.annotation_color = color
        else:
            self.fill_color = color
        self._refresh_color_buttons()

    def _refresh_color_buttons(self):
        self.text_color_button.set_swatch(self.text_color)
        self.stroke_color_button.set_swatch(self.annotation_color)
        self.fill_color_button.set_swatch(self.fill_color)

    def handle_canvas_action(self, tool: str, payload):
        if tool == "__open_pdf__":
            self.open_pdf()
            return True
        if self.doc is None:
            return False

        try:
            if tool == "__zoom_delta__":
                self.zoom_by_wheel(int(payload))
                return True
            if tool == "__pan__":
                self.pan_view(payload)
                return True
            if tool == "__fit_width__":
                self.fit_width()
                return True
            if tool == "__move_text_start__":
                return self.begin_move_text(payload)
            if tool == "__move_text_end__":
                self.end_move_text(payload)
                return True
            if tool == "__recolor_annot__":
                return self.recolor_annotation_at(payload)
            if tool == "__delete_selected_annot__":
                return self.delete_selected_annot_on_canvas()

            if tool == Tool.TEXT_BOX:
                self.add_text_box(payload)
            elif tool == Tool.EDIT_TEXT:
                self.edit_existing_text_span(payload)
            elif tool == Tool.EDIT_BLOCK:
                self.edit_existing_text_block(payload)
            elif tool == Tool.NOTE:
                self.add_note(payload)
            elif tool == Tool.HIGHLIGHT:
                self.add_markup_annotation(payload, "highlight")
            elif tool == Tool.UNDERLINE:
                self.add_markup_annotation(payload, "underline")
            elif tool == Tool.STRIKEOUT:
                self.add_markup_annotation(payload, "strikeout")
            elif tool == Tool.RECTANGLE:
                self.add_rectangle(payload)
            elif tool == Tool.LINE:
                self.add_line(payload)
            elif tool == Tool.INK:
                self.add_ink(payload)
            elif tool == Tool.REDACT:
                self.apply_redaction(payload)
            elif tool == Tool.CROP:
                self.crop_page(payload)
            elif tool == Tool.LINK:
                self.add_url_link(payload)
            elif tool == Tool.IMAGE:
                self.insert_image_at(payload)
            elif tool == Tool.SIGN:
                self.place_signature_at(payload)
            return True
        except Exception as exc:
            QMessageBox.critical(self, "Tool Error", f"{tool} failed:\n{exc}")
            return False

    def zoom_by_wheel(self, delta: int):
        if self.doc is None:
            return
        factor = 1.10 if delta > 0 else 1 / 1.10
        self.zoom = clamp(self.zoom * factor, 0.25, 5.0)
        self.render_current_page()

    def pan_view(self, delta: QPointF):
        self.scroll.horizontalScrollBar().setValue(self.scroll.horizontalScrollBar().value() - int(delta.x()))
        self.scroll.verticalScrollBar().setValue(self.scroll.verticalScrollBar().value() - int(delta.y()))

    def begin_move_text(self, image_point: QPointF) -> bool:
        pdf_point = fitz.Point(image_point.x() / self.zoom, image_point.y() / self.zoom)
        self._moving_image_hit = None
        self._moving_signature_index = None
        self._moving_annot_ordinal = None

        # 1. Tracked items (signatures / text boxes) — redrawn from clean data.
        sig_index = self._find_signature_at_point(self.current_page_index, pdf_point)
        if sig_index is not None:
            self._moving_text_hit = None
            self._moving_signature_index = sig_index
            kind = self.signatures[sig_index].get("kind", "signature")
            self.status.showMessage("Move: dragging %s" % ("text box" if kind == "textbox" else "signature"))
            return True

        # 2. Native PDF annotations (highlight, note, rectangle, ink, …). These
        # move cleanly via set_rect, so they can be repositioned in place.
        ordinal = self._find_annot_at_point(self.current_page_index, pdf_point)
        if ordinal is not None:
            self._moving_text_hit = None
            self._moving_annot_ordinal = ordinal
            self._selected_annot = (self.current_page_index, ordinal)
            self.status.showMessage("Move: dragging annotation (Del to delete, double-click to recolour)")
            return True

        self._moving_text_hit = None
        self._selected_annot = None
        self.status.showMessage("Move: click a signature, text box or annotation.")
        return False

    def _find_annot_at_point(self, page_index: int, point: fitz.Point) -> Optional[int]:
        """Return the ordinal index of the top-most annotation whose rectangle
        contains the point on the given page (or None)."""
        try:
            page = self.doc[page_index]
        except Exception:
            return None
        tol = fitz.Rect(point.x - 2, point.y - 2, point.x + 2, point.y + 2)
        best = None
        for ordinal, annot in enumerate(page.annots()):
            try:
                r = fitz.Rect(annot.rect)
            except Exception:
                continue
            if r.is_empty or r.is_infinite:
                continue
            if r.contains(point) or r.intersects(tol):
                best = ordinal  # later annots draw on top
        return best

    def _find_image_at_point(self, page, point: fitz.Point) -> Optional[Dict]:
        """Return info for the top-most image whose placement contains the
        click, so signatures and inserted images can be moved. Later-drawn
        images sit on top, so we scan in reverse draw order."""
        try:
            infos = page.get_image_info(xrefs=True)
        except Exception:
            infos = []
        tol = fitz.Rect(point.x - 2, point.y - 2, point.x + 2, point.y + 2)
        chosen = None
        chosen_bbox = None
        for info in infos:  # draw order; keep the last (top-most) match
            bbox = fitz.Rect(info["bbox"])
            if bbox.is_empty or bbox.is_infinite:
                continue
            if bbox.contains(point) or bbox.intersects(tol):
                chosen = info
                chosen_bbox = bbox
        if chosen is None:
            return None
        xref = int(chosen.get("xref", 0))
        if xref <= 0:
            return None
        # Build a Pixmap from the xref so any transparency (SMask) is preserved;
        # extract_image alone loses the alpha and would re-insert as an opaque
        # (black) rectangle.
        try:
            pix = fitz.Pixmap(self.doc, xref)
            if pix.width <= 0 or pix.height <= 0:
                return None
            png_bytes = pix.tobytes("png")
            width, height = pix.width, pix.height
        except Exception:
            try:
                extracted = self.doc.extract_image(xref)
                png_bytes = extracted.get("image")
                width = int(extracted.get("width", 0) or 0)
                height = int(extracted.get("height", 0) or 0)
            except Exception:
                return None
        if not png_bytes or width <= 0 or height <= 0:
            return None
        return {"bbox": chosen_bbox, "bytes": png_bytes, "w": width,
                "h": height, "xref": xref}

    def end_move_text(self, image_points: Tuple[QPointF, QPointF]):
        start, end = image_points
        if self.zoom <= 0:
            return
        dx = (end.x() - start.x()) / self.zoom
        dy = (end.y() - start.y()) / self.zoom

        # --- Moving a tracked item (signature or added text box) ---
        # These are always redrawn from clean source data, so their text, font,
        # size and colour are preserved no matter how many times they move.
        sig_index = getattr(self, "_moving_signature_index", None)
        if sig_index is not None:
            self._moving_signature_index = None
            if abs(dx) < 0.5 and abs(dy) < 0.5:
                return
            if not (0 <= sig_index < len(self.signatures)):
                return
            sig = self.signatures[sig_index]
            page = self.doc[self.current_page_index]
            old_rect = fitz.Rect(sig["rect"])
            self._push_undo()
            page = self.doc[self.current_page_index]

            if sig.get("kind") == "signature" and sig.get("png"):
                # Image signature: remove the old image, then place the SAME PNG
                # at the shifted rect. Re-using the stored bytes means it never
                # corrupts, and redaction keeps the file from accumulating
                # hidden copies. (We always re-insert, so removal is safe.)
                new_rect = old_rect + (dx, dy, dx, dy)
                try:
                    page.add_redact_annot(old_rect + (-2, -2, 2, 2), fill=(1, 1, 1))
                    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_REMOVE)
                    self.doc.reload_page(page)
                    page = self.doc[self.current_page_index]
                except Exception:
                    try:
                        page.draw_rect(old_rect + (-2, -2, 2, 2), color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                    except Exception:
                        pass
                page.insert_image(new_rect, stream=sig["png"],
                                  keep_proportion=False, overlay=True)
                sig = dict(sig)
                sig["rect"] = [new_rect.x0, new_rect.y0, new_rect.x1, new_rect.y1]
                self.signatures[sig_index] = sig
                self._mark_dirty("Signature moved", refresh_sidebars=True)
                return

            # Text box: erase old text and redraw fresh Helvetica text.
            size = float(sig.get("size", 48))
            erase = old_rect + (-2, -2, size * 0.3, 2)
            page.add_redact_annot(erase, fill=(1, 1, 1))
            page.apply_redactions()
            self.signatures.pop(sig_index)
            col = sig.get("color", [17, 17, 17])
            color = QColor(col[0], col[1], col[2])
            self._draw_signature(
                page, self.current_page_index,
                old_rect.x0 + dx + size * 0.1, old_rect.y0 + dy + size * 0.2,
                sig["name"], sig.get("fontfile"), size, color,
                kind="textbox", align=sig.get("align", 0),
            )
            self._mark_dirty("Text box moved", refresh_sidebars=True)
            return

        # --- Moving a native PDF annotation via set_rect ---
        ordinal = getattr(self, "_moving_annot_ordinal", None)
        if ordinal is not None:
            self._moving_annot_ordinal = None
            if abs(dx) < 0.5 and abs(dy) < 0.5:
                return
            page = self.doc[self.current_page_index]
            annots = list(page.annots())
            if not (0 <= ordinal < len(annots)):
                return
            try:
                old = fitz.Rect(annots[ordinal].rect)
            except Exception:
                return
            self._push_undo()
            # Re-fetch after the undo snapshot renumbers objects.
            page = self.doc[self.current_page_index]
            annots = list(page.annots())
            if not (0 <= ordinal < len(annots)):
                return
            try:
                annots[ordinal].set_rect(old + (dx, dy, dx, dy))
                annots[ordinal].update()
            except Exception as exc:
                QMessageBox.warning(self, "Move Annotation", f"Could not move this annotation:\n{exc}")
                return
            self._selected_annot = (self.current_page_index, ordinal)
            self._mark_dirty("Annotation moved", refresh_sidebars=True)
            return

        # Nothing tracked was under the cursor: do nothing (existing PDF body
        # text and images are intentionally not movable to avoid corruption).
        self._moving_text_hit = None
        self._moving_image_hit = None

    def recolor_annotation_at(self, image_point: QPointF) -> bool:
        """Recolour the annotation under the cursor (double-click, Move tool)."""
        if self.doc is None:
            return False
        pdf_point = fitz.Point(image_point.x() / self.zoom, image_point.y() / self.zoom)
        ordinal = self._find_annot_at_point(self.current_page_index, pdf_point)
        if ordinal is None:
            return False
        page = self.doc[self.current_page_index]
        annots = list(page.annots())
        if not (0 <= ordinal < len(annots)):
            return False
        # Read the current colour BEFORE any undo snapshot (which renumbers and
        # invalidates live annotation handles).
        try:
            colors = annots[ordinal].colors or {}
            stroke = colors.get("stroke")
            initial = QColor.fromRgbF(*stroke) if stroke else QColor("#000000")
        except Exception:
            initial = QColor("#000000")

        color = QColorDialog.getColor(initial, self, "Annotation Colour")
        if not color.isValid():
            return True  # handled (user cancelled) — don't fall back to zoom

        self._push_undo()
        rgb = (color.redF(), color.greenF(), color.blueF())
        # Re-fetch the annotation fresh after the undo snapshot.
        page = self.doc[self.current_page_index]
        annots = list(page.annots())
        if not (0 <= ordinal < len(annots)):
            return True
        try:
            annots[ordinal].set_colors(stroke=rgb)
            annots[ordinal].update()
        except Exception as exc:
            QMessageBox.warning(self, "Recolour Annotation", f"Could not recolour:\n{exc}")
            return True
        self._selected_annot = (self.current_page_index, ordinal)
        self._mark_dirty("Annotation recoloured", refresh_sidebars=True)
        return True

    def delete_selected_annot_on_canvas(self) -> bool:
        """Delete the annotation selected on the canvas (Delete key, Move tool)."""
        sel = getattr(self, "_selected_annot", None)
        if self.doc is None or not sel:
            return False
        page_index, ordinal = sel
        if page_index != self.current_page_index:
            return False
        page = self.doc[page_index]
        annots = list(page.annots())
        if not (0 <= ordinal < len(annots)):
            self._selected_annot = None
            return False
        self._push_undo()
        try:
            page.delete_annot(annots[ordinal])
        except Exception:
            return False
        try:
            self.doc.reload_page(page)
        except Exception:
            pass
        self._selected_annot = None
        self._mark_dirty("Annotation deleted", refresh_sidebars=True)
        return True

    def edit_existing_text_span(self, image_point: QPointF):
        page = self.doc[self.current_page_index]
        pdf_point = fitz.Point(image_point.x() / self.zoom, image_point.y() / self.zoom)

        hit = self._find_span_at_point(page, pdf_point)
        if not hit:
            QMessageBox.information(self, "No Text Found", "No editable text span found at this point.")
            return

        dialog = TextBoxDialog(self, "Replace Text Span", hit.text)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        replacement = dialog.text()
        if replacement == hit.text:
            return

        self._push_undo()
        font_size = max(6, hit.size)
        erase_rect = hit.bbox + (-1, -1, 1, 1)
        draw_rect = self._expanded_text_rect(page, erase_rect, replacement, font_size)
        self._replace_rect_with_text(
            page=page,
            rect=erase_rect,
            text=replacement,
            fontsize=font_size,
            color=hit.color if hit.color.isValid() else self.text_color,
            fill=(1, 1, 1),
            align=0,
            draw_rect=draw_rect,
        )
        self._mark_dirty("Text span replaced")

    def edit_existing_text_block(self, image_point: QPointF):
        page = self.doc[self.current_page_index]
        pdf_point = fitz.Point(image_point.x() / self.zoom, image_point.y() / self.zoom)

        hit = self._find_block_at_point(page, pdf_point)
        if not hit:
            QMessageBox.information(self, "No Text Block Found", "No editable text block found at this point.")
            return

        dialog = TextBoxDialog(self, "Replace Text Block", hit.text)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        replacement = dialog.text()
        if replacement == hit.text:
            return

        self._push_undo()
        font_size = self.font_size_spin.value()
        erase_rect = hit.bbox + (-2, -2, 2, 2)
        draw_rect = self._expanded_text_rect(page, erase_rect, replacement, font_size)
        self._replace_rect_with_text(
            page=page,
            rect=erase_rect,
            text=replacement,
            fontsize=font_size,
            color=self.text_color,
            fill=(1, 1, 1),
            align=0,
            draw_rect=draw_rect,
        )
        self._mark_dirty("Text block replaced")

    def search_and_replace(self):
        if self.doc is None:
            return

        dialog = SearchReplaceDialog(self)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        find_text, replace_text, scope, case_sensitive = dialog.values()
        if not find_text:
            return

        pages = [self.current_page_index] if scope == "Current Page" else list(range(self.doc.page_count))
        total = 0

        self._push_undo()
        for page_index in pages:
            page = self.doc[page_index]
            rects = page.search_for(find_text)
            if case_sensitive and rects:
                exact = []
                for rect in rects:
                    try:
                        found = page.get_textbox(rect)
                    except Exception:
                        found = ""
                    if find_text in found:
                        exact.append(rect)
                rects = exact
            if not rects:
                continue

            for rect in rects:
                page.add_redact_annot(rect + (-1, -1, 1, 1), fill=(1, 1, 1))
            page.apply_redactions()

            for rect in rects:
                size = clamp(rect.height * 0.72, 6, 24)
                draw_rect = self._expanded_text_rect(page, rect + (0, -1, max(80, rect.width * 2), 5), replace_text, size)
                self._safe_insert_textbox(page, draw_rect, replace_text, size, self.text_color, align=0)
                total += 1

        if total == 0:
            self.undo_stack.pop() if self.undo_stack else None
            QMessageBox.information(self, "Search and Replace", "No matches found.")
            self._refresh_history_actions()
            return

        self._mark_dirty(f"Replaced {total} text match(es)", refresh_sidebars=True)

    def reflow_current_page_to_new_page(self):
        if self.doc is None:
            return

        page = self.doc[self.current_page_index]
        text = page.get_text("text", sort=True).strip()
        if not text:
            QMessageBox.information(self, "Reflow", "No selectable text found on this page.")
            return

        self._push_undo()
        self._insert_reflowed_text_pages(
            text=text,
            start_at=self.current_page_index + 1,
            page_size=(page.rect.width, page.rect.height),
            title=f"Reflowed copy of page {self.current_page_index + 1}",
        )
        self.current_page_index += 1
        self._mark_dirty("Reflowed page inserted after current page", refresh_sidebars=True)

    def reflow_document_to_new_pdf(self):
        if self.doc is None:
            return

        out, _ = QFileDialog.getSaveFileName(self, "Save Reflowed Text PDF", "reflowed_document.pdf", "PDF Files (*.pdf)")
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"

        try:
            new_doc = fitz.open()
            for i, page in enumerate(self.doc):
                text = page.get_text("text", sort=True).strip()
                if not text:
                    text = f"[Page {i + 1}: no selectable text found]"
                self._insert_reflowed_text_pages_into_doc(
                    doc=new_doc,
                    text=text,
                    page_size=(page.rect.width, page.rect.height),
                    title=f"Page {i + 1}",
                )
            new_doc.save(out, garbage=4, deflate=True)
            new_doc.close()
            QMessageBox.information(self, "Reflow Complete", f"Reflowed PDF saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "Reflow Error", f"Could not create reflowed PDF:\n{exc}")

    def _replace_rect_with_text(
        self,
        page,
        rect: fitz.Rect,
        text: str,
        fontsize: float,
        color: QColor,
        fill: Tuple[float, float, float] = (1, 1, 1),
        align: int = 0,
        draw_rect: Optional[fitz.Rect] = None,
    ):
        """Erase old visible PDF text and redraw replacement text safely.

        The erase rectangle and draw rectangle are intentionally separate. Text
        spans in PDFs are often too small for replacement text; using the same
        rectangle can make PyMuPDF draw nothing.
        """
        erase_rect = fitz.Rect(rect)
        page.add_redact_annot(erase_rect, fill=fill)
        page.apply_redactions()

        if not text or not str(text).strip():
            return

        final_draw_rect = draw_rect or self._expanded_text_rect(page, erase_rect, text, fontsize)
        self._safe_insert_textbox(page, final_draw_rect, text, fontsize, color, align=align)

    def _expanded_text_rect(self, page, rect: fitz.Rect, text: str, fontsize: float) -> fitz.Rect:
        """Return a draw box large enough that replacement text does not vanish.

        PyMuPDF's insert_textbox returns a negative value when the text does not
        fit. This method estimates a safe width/height before drawing.
        """
        font_size = float(clamp(fontsize, 6, 72))
        source = str(text or "")
        lines = source.splitlines() or [source]
        non_empty = [line if line else " " for line in lines]

        try:
            max_line_width = max(fitz.get_text_length(line, fontname="helv", fontsize=font_size) for line in non_empty)
        except Exception:
            max_line_width = max(len(line) for line in non_empty) * font_size * 0.56

        needed_w = max(rect.width, max_line_width + font_size * 0.8 + 6)
        needed_h = max(rect.height, len(non_empty) * font_size * 1.45 + 6)

        page_rect = page.rect
        x0 = clamp(rect.x0, page_rect.x0, page_rect.x1 - 8)
        y0 = clamp(rect.y0 - font_size * 0.15, page_rect.y0, page_rect.y1 - 8)
        x1 = min(max(rect.x1, x0 + needed_w), page_rect.x1 - 2)
        y1 = min(max(rect.y1 + font_size * 0.45, y0 + needed_h), page_rect.y1 - 2)

        if x1 - x0 < min(needed_w, page_rect.width - 4):
            x0 = max(page_rect.x0 + 2, x1 - min(needed_w, page_rect.width - 4))

        return fitz.Rect(x0, y0, x1, y1)

    def _safe_insert_textbox(self, page, rect: fitz.Rect, text: str, fontsize: float, color: QColor, align: int = 0, fontfile: Optional[str] = None):
        """Insert text safely and never treat negative insert_textbox return as success.

        insert_textbox does not always raise an exception. When text does not
        fit, it returns a negative number and draws nothing. That was the cause
        of the disappearing text bug.

        If ``fontfile`` is given (e.g. a signature's original script font), it is
        embedded and tried first so the moved text keeps its typeface.
        """
        if not text or not str(text).strip():
            return 0

        font_size = float(clamp(fontsize, 6, 72))
        rgb = rgb_from_qcolor(color if color and color.isValid() else QColor(17, 24, 39))
        text_value = str(text)
        draw_rect = self._expanded_text_rect(page, fitz.Rect(rect), text_value, font_size)
        last_error = None

        # Build the list of fonts to attempt. A custom fontfile (if valid) goes
        # first, then the built-in Helvetica fallbacks. A fresh unique alias per
        # call forces PyMuPDF to (re-)embed the font, which is important after a
        # redaction may have dropped a previous embedding.
        embedded_alias = None
        if fontfile and Path(fontfile).exists():
            self._font_alias_counter = getattr(self, "_font_alias_counter", 0) + 1
            embedded_alias = "sig%d" % self._font_alias_counter
        font_attempts = ([(embedded_alias, fontfile)] if embedded_alias else []) + \
                        [("helv", None), ("Helvetica", None), (None, None)]

        for candidate_rect in (draw_rect, fitz.Rect(rect)):
            for fontname, ffile in font_attempts:
                try:
                    kwargs = dict(
                        fontsize=font_size,
                        color=rgb,
                        align=align,
                        overlay=True,
                    )
                    if fontname:
                        kwargs["fontname"] = fontname
                    if ffile:
                        kwargs["fontfile"] = ffile
                    rc = page.insert_textbox(candidate_rect, text_value, **kwargs)
                    if rc is None or rc >= 0:
                        return rc
                    last_error = RuntimeError(f"insert_textbox returned {rc}; text did not fit")
                except Exception as exc:
                    last_error = exc

        try:
            x = draw_rect.x0
            y = draw_rect.y0 + font_size
            line_gap = font_size * 1.25
            for raw_line in text_value.splitlines() or [text_value]:
                line = raw_line if raw_line else " "
                if y > page.rect.y1 - 2:
                    break
                inserted = False
                for fontname in ("helv", "Helvetica", None):
                    try:
                        kwargs = dict(fontsize=font_size, color=rgb, overlay=True)
                        if fontname:
                            kwargs["fontname"] = fontname
                        page.insert_text(fitz.Point(x, y), line, **kwargs)
                        inserted = True
                        break
                    except Exception as exc:
                        last_error = exc
                if not inserted:
                    raise last_error or RuntimeError("insert_text fallback failed")
                y += line_gap
            return 0
        except Exception as exc:
            raise RuntimeError(f"Could not draw text. First error: {last_error}; fallback error: {exc}")

    def _iter_text_spans(self, page) -> Iterable[TextSpanHit]:
        info = page.get_text("dict")
        for block in info.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text.strip():
                        continue
                    yield TextSpanHit(
                        text=text,
                        bbox=fitz.Rect(span["bbox"]),
                        size=float(span.get("size", 11)),
                        color=qcolor_from_pdf_int(int(span.get("color", 0))),
                        font=span.get("font", "helv"),
                    )

    def _iter_text_blocks(self, page) -> Iterable[TextBlockHit]:
        blocks = page.get_text("blocks", sort=True)
        for block in blocks:
            if len(block) < 5:
                continue
            x0, y0, x1, y1, text = block[:5]
            if str(text).strip():
                yield TextBlockHit(text=str(text).strip(), bbox=fitz.Rect(x0, y0, x1, y1))

    def _find_span_at_point(self, page, point: fitz.Point) -> Optional[TextSpanHit]:
        candidates = [span for span in self._iter_text_spans(page) if span.bbox.contains(point)]
        if not candidates:
            tolerance = fitz.Rect(point.x - 3, point.y - 3, point.x + 3, point.y + 3)
            candidates = [span for span in self._iter_text_spans(page) if span.bbox.intersects(tolerance)]
        if not candidates:
            return None
        return min(candidates, key=lambda s: s.bbox.get_area())

    def _find_block_at_point(self, page, point: fitz.Point) -> Optional[TextBlockHit]:
        candidates = [block for block in self._iter_text_blocks(page) if block.bbox.contains(point)]
        if not candidates:
            tolerance = fitz.Rect(point.x - 4, point.y - 4, point.x + 4, point.y + 4)
            candidates = [block for block in self._iter_text_blocks(page) if block.bbox.intersects(tolerance)]
        if not candidates:
            return None
        return min(candidates, key=lambda b: b.bbox.get_area())

    def _insert_reflowed_text_pages(
        self,
        text: str,
        start_at: int,
        page_size: Tuple[float, float],
        title: str,
    ):
        width, height = page_size
        chunks = self._paginate_text_for_reflow(text, width, height)
        for offset, chunk in enumerate(chunks):
            page = self.doc.new_page(pno=start_at + offset, width=width, height=height)
            self._draw_reflow_page(page, chunk, title if offset == 0 else f"{title} continued")

    def _insert_reflowed_text_pages_into_doc(
        self,
        doc,
        text: str,
        page_size: Tuple[float, float],
        title: str,
    ):
        width, height = page_size
        chunks = self._paginate_text_for_reflow(text, width, height)
        for offset, chunk in enumerate(chunks):
            page = doc.new_page(width=width, height=height)
            self._draw_reflow_page(page, chunk, title if offset == 0 else f"{title} continued")

    def _paginate_text_for_reflow(self, text: str, width: float, height: float) -> List[str]:
        margin = 54
        usable_width = max(width - margin * 2, 250)
        chars_per_line = max(50, int(usable_width / 5.4))
        lines_per_page = max(28, int((height - margin * 2 - 34) / 15))

        lines: List[str] = []
        for para in text.splitlines():
            para = para.strip()
            if not para:
                lines.append("")
                continue
            lines.extend(textwrap.wrap(para, width=chars_per_line) or [""])

        chunks = []
        for i in range(0, len(lines), lines_per_page):
            chunks.append("\n".join(lines[i:i + lines_per_page]))
        return chunks or [""]

    def _draw_reflow_page(self, page, text: str, title: str):
        rect = page.rect
        margin = 54
        page.insert_text(
            fitz.Point(margin, margin - 18),
            title,
            fontsize=13,
            fontname="helv",
            color=(0.10, 0.20, 0.45),
        )
        page.draw_line(
            fitz.Point(margin, margin - 8),
            fitz.Point(rect.width - margin, margin - 8),
            color=(0.70, 0.75, 0.82),
            width=0.8,
        )
        page.insert_textbox(
            fitz.Rect(margin, margin + 8, rect.width - margin, rect.height - margin),
            text,
            fontsize=11,
            fontname="helv",
            color=(0.05, 0.05, 0.05),
            align=0,
        )

    def add_text_box(self, image_rect: QRectF):
        dialog = TextBoxDialog(self, "Add Text Box")
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        text = dialog.text().strip()
        if not text:
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = self._image_rect_to_pdf_rect(image_rect)
        # Register the text box like a signature so it can be moved cleanly with
        # the Move tool (re-drawn from its original text/size/colour).
        self._draw_signature(
            page, self.current_page_index,
            rect.x0, rect.y0, text, None,
            float(self.font_size_spin.value()), self.text_color,
            kind="textbox", align=0,
        )
        self._mark_dirty("Text box added", refresh_sidebars=True)

    def add_note(self, image_point: QPointF):
        text, ok = QInputDialog.getMultiLineText(self, "Sticky Note", "Note:")
        if not ok or not text.strip():
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        x, y = self._image_point_to_pdf_xy(image_point)
        annot = page.add_text_annot(fitz.Point(x, y), text.strip())
        annot.set_info(title="PDF Studio Note", content=text.strip())
        annot.update()
        self._mark_dirty("Sticky note added", refresh_sidebars=True)

    def add_markup_annotation(self, image_rect: QRectF, kind: str):
        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = self._image_rect_to_pdf_rect(image_rect)

        if kind == "highlight":
            annot = page.add_highlight_annot(rect)
            annot.set_colors(stroke=rgb_from_qcolor(QColor("#b4b4b4")))
        elif kind == "underline":
            annot = page.add_underline_annot(rect)
            annot.set_colors(stroke=rgb_from_qcolor(self.annotation_color))
        else:
            annot = page.add_strikeout_annot(rect)
            annot.set_colors(stroke=rgb_from_qcolor(self.annotation_color))
        annot.update()
        self._mark_dirty(f"{kind.title()} added", refresh_sidebars=True)

    def add_rectangle(self, image_rect: QRectF):
        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = self._image_rect_to_pdf_rect(image_rect)
        annot = page.add_rect_annot(rect)
        annot.set_colors(stroke=rgb_from_qcolor(self.annotation_color), fill=rgb_from_qcolor(self.fill_color))
        annot.set_border(width=self.line_width)
        annot.set_opacity(self.opacity_percent / 100)
        annot.update()
        self._mark_dirty("Rectangle added", refresh_sidebars=True)

    def add_line(self, image_points: Tuple[QPointF, QPointF]):
        self._push_undo()
        page = self.doc[self.current_page_index]
        p1, p2 = image_points
        annot = page.add_line_annot(
            fitz.Point(p1.x() / self.zoom, p1.y() / self.zoom),
            fitz.Point(p2.x() / self.zoom, p2.y() / self.zoom),
        )
        annot.set_colors(stroke=rgb_from_qcolor(self.annotation_color))
        annot.set_border(width=self.line_width)
        annot.update()
        self._mark_dirty("Line added", refresh_sidebars=True)

    def add_ink(self, image_points: List[QPointF]):
        """Add freehand ink.

        PyMuPDF expects ink paths as a sequence of strokes, where each stroke is
        a sequence of plain float coordinate pairs: [[(x, y), (x, y), ...]].
        Passing fitz.Point objects can crash on newer PyMuPDF builds with:
        "arg must be seq of seq of float pairs".
        """
        if not image_points or len(image_points) < 2:
            return

        self._push_undo()
        page = self.doc[self.current_page_index]

        stroke = [
            (float(p.x() / self.zoom), float(p.y() / self.zoom))
            for p in image_points
        ]

        compact_stroke = []
        last = None
        for point in stroke:
            rounded = (round(point[0], 2), round(point[1], 2))
            if rounded != last:
                compact_stroke.append(point)
                last = rounded

        if len(compact_stroke) < 2:
            return

        try:
            annot = page.add_ink_annot([compact_stroke])
            annot.set_colors(stroke=rgb_from_qcolor(self.annotation_color))
            annot.set_border(width=self.line_width)
            annot.update()
        except Exception:
            color = rgb_from_qcolor(self.annotation_color)
            for start, end in zip(compact_stroke, compact_stroke[1:]):
                page.draw_line(
                    fitz.Point(start[0], start[1]),
                    fitz.Point(end[0], end[1]),
                    color=color,
                    width=self.line_width,
                    overlay=True,
                )

        self._mark_dirty("Ink added", refresh_sidebars=True)

    def apply_redaction(self, image_rect: QRectF):
        response = QMessageBox.warning(
            self,
            "Apply Redaction",
            "This permanently removes visible content inside the selected area in the edited copy. Continue?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if response != QMessageBox.StandardButton.Yes:
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = self._image_rect_to_pdf_rect(image_rect)
        page.add_redact_annot(rect, fill=(0, 0, 0))
        page.apply_redactions()
        self._mark_dirty("Redaction applied", refresh_sidebars=True)

    def crop_page(self, image_rect: QRectF):
        response = QMessageBox.question(
            self,
            "Crop Page",
            "Crop the current page to the selected area?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if response != QMessageBox.StandardButton.Yes:
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = self._image_rect_to_pdf_rect(image_rect)
        page.set_cropbox(rect)
        self._mark_dirty("Page cropped", refresh_sidebars=True)

    def add_url_link(self, image_rect: QRectF):
        url, ok = QInputDialog.getText(self, "Add URL Link", "URL:")
        if not ok or not url.strip():
            return
        url = url.strip()
        if not (url.startswith("http://") or url.startswith("https://") or url.startswith("mailto:")):
            url = "https://" + url

        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = self._image_rect_to_pdf_rect(image_rect)
        page.insert_link({"kind": fitz.LINK_URI, "from": rect, "uri": url})
        page.draw_rect(rect, color=(0.1, 0.25, 0.9), width=0.8, overlay=True)
        self._mark_dirty("URL link area added")

    def prepare_insert_image(self):
        if self.doc is None:
            return
        path, _ = QFileDialog.getOpenFileName(self, "Choose Image", "", "Images (*.png *.jpg *.jpeg *.bmp *.webp)")
        if not path:
            return
        self.pending_image_path = path
        self.set_tool(Tool.IMAGE)
        self.status.showMessage("Click where you want to place the image.")

    def insert_image_at(self, image_point: QPointF):
        if not self.pending_image_path:
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        x, y = self._image_point_to_pdf_xy(image_point)
        img = QImage(self.pending_image_path)
        if img.isNull():
            raise ValueError("Could not read selected image.")

        width = min(240, page.rect.width * 0.45)
        height = width * img.height() / max(img.width(), 1)
        rect = fitz.Rect(x, y, x + width, y + height)
        page.insert_image(rect, filename=self.pending_image_path, overlay=True)

        self.pending_image_path = None
        self.set_tool(Tool.SELECT)
        self._mark_dirty("Image inserted", refresh_sidebars=True)

    def prepare_signature(self):
        if self.doc is None:
            return
        default = ""
        try:
            info = self.doc.metadata or {}
            default = info.get("author", "") or ""
        except Exception:
            default = ""
        dialog = SignatureDialog(self, default_name=default)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        name, family, color, size = dialog.result_values()
        if not name:
            return
        # Signatures are inserted as real (searchable, movable) text drawn with
        # the chosen script font, so they behave like any other text: they can
        # be repositioned with the Move tool and never turn into opaque boxes.
        self.pending_signature = {
            "name": name,
            "family": family,
            "color": QColor(color),
            "size": int(size),
            "fontfile": signature_font_file(family),
        }
        self.set_tool(Tool.SIGN)
        self.status.showMessage("Click where you want to place your signature.")

    def place_signature_at(self, image_point: QPointF):
        if not self.pending_signature:
            return
        sig = self.pending_signature
        name = sig.get("name", "")
        if not name:
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        x, y = self._image_point_to_pdf_xy(image_point)
        self._draw_signature(
            page, self.current_page_index, x, y, name,
            sig.get("fontfile"), float(sig.get("size", 48)),
            sig.get("color", QColor("#111111")),
        )

        self.pending_signature = None
        self.set_tool(Tool.SELECT)
        self._mark_dirty("Signature added", refresh_sidebars=True)

    def _draw_signature(self, page, page_index, x, y, name, fontfile, size, color, kind="signature", align=0):
        """Draw an added item and record it in the per-document registry so it
        can be moved later from clean source data.

        Signatures are drawn as a transparent IMAGE (rendered once from the
        script font). Re-embedding a subset script font on every move corrupts
        its text encoding, so images are used instead — they move losslessly.
        Text boxes stay as real Helvetica text (which does not corrupt).
        """
        size = float(size)

        if kind == "signature":
            # Render the signature to a crisp transparent PNG and place it.
            family = None
            # Recover a family name for QFont from the font file, if given.
            image = None
            if fontfile:
                fam = self._family_for_fontfile(fontfile)
                image = render_signature_image(name, fam, color)
            if image is None or image.isNull():
                image = render_signature_image(name, APP_FONT_FAMILY, color)
            png = self._qimage_to_png(image)
            aspect = image.width() / max(image.height(), 1)
            height = max(12.0, size)
            width = height * aspect
            max_w = page.rect.width * 0.95
            if width > max_w:
                width = max_w
                height = width / max(aspect, 0.01)
            rect = fitz.Rect(x, y, x + width, y + height)
            page.insert_image(rect, stream=png, keep_proportion=False, overlay=True)
            self.signatures.append({
                "page": int(page_index),
                "rect": [rect.x0, rect.y0, rect.x1, rect.y1],
                "name": name,
                "fontfile": fontfile,
                "size": size,
                "color": [color.red(), color.green(), color.blue()],
                "kind": "signature",
                "align": align,
                "png": png,          # stored bytes → lossless re-placement
                "img_w": image.width(),
                "img_h": image.height(),
            })
            return

        # --- text box: real Helvetica text (safe to redraw) ---
        self._font_alias_counter = getattr(self, "_font_alias_counter", 0) + 1
        fontname = "sig%d" % self._font_alias_counter
        lines = name.splitlines() or [name]
        try:
            text_len = max(fitz.get_text_length(ln, fontname="helv", fontsize=size) for ln in lines)
        except Exception:
            text_len = max(len(ln) for ln in lines) * size * 0.5
        n_lines = len(lines)
        box_w = min(max(text_len + size, size * 2), page.rect.width * 0.95)
        box_h = size * (1.4 * n_lines + 0.4)
        rect = fitz.Rect(x, y, x + box_w, y + box_h)
        rc = page.insert_textbox(rect, name, fontsize=size, fontname="helv",
                                 color=rgb_from_qcolor(color), align=align, overlay=True)
        if rc < 0:
            rect = fitz.Rect(x, y, x + page.rect.width * 0.95, y + box_h * 1.5)
            page.insert_textbox(rect, name, fontsize=size, fontname="helv",
                                color=rgb_from_qcolor(color), align=align, overlay=True)
        pad_x = size * 0.4
        placed_rect = fitz.Rect(
            x - pad_x * 0.25, y - size * 0.2,
            x + min(text_len + pad_x, box_w) + pad_x,
            y + size * (1.4 * n_lines + 0.3),
        )
        self.signatures.append({
            "page": int(page_index),
            "rect": [placed_rect.x0, placed_rect.y0, placed_rect.x1, placed_rect.y1],
            "name": name,
            "fontfile": None,
            "size": size,
            "color": [color.red(), color.green(), color.blue()],
            "kind": "textbox",
            "align": align,
        })

    @staticmethod
    def _qimage_to_png(image) -> bytes:
        buf = QBuffer()
        buf.open(QIODevice.OpenModeFlag.WriteOnly)
        image.save(buf, "PNG")
        data = bytes(buf.data())
        buf.close()
        return data

    @staticmethod
    def _family_for_fontfile(fontfile: str) -> str:
        """Map a bundled signature font file back to its Qt family name."""
        for fam, path in _SIGNATURE_FONT_PATHS.items():
            if path == fontfile:
                return fam
        return APP_FONT_FAMILY

    def _find_signature_at_point(self, page_index: int, point: fitz.Point) -> Optional[int]:
        """Return the index in self.signatures of a placed signature under the
        point on the given page, or None."""
        tol = fitz.Rect(point.x - 3, point.y - 3, point.x + 3, point.y + 3)
        best = None
        for i, sig in enumerate(self.signatures):
            if sig.get("page") != page_index:
                continue
            r = fitz.Rect(sig["rect"])
            if r.contains(point) or r.intersects(tol):
                best = i  # later placements are on top
        return best

    def add_stamp(self):
        presets = ["DRAFT", "APPROVED", "CONFIDENTIAL", "REVIEWED", "PAID", "VOID", "FINAL"]
        text, ok = QInputDialog.getItem(self, "Add Stamp", "Stamp:", presets, 0, True)
        if not ok or not text.strip():
            return

        self._push_undo()
        page = self.doc[self.current_page_index]
        rect = page.rect
        stamp_rect = fitz.Rect(rect.width * 0.20, rect.height * 0.38, rect.width * 0.80, rect.height * 0.52)
        page.draw_rect(stamp_rect, color=rgb_from_qcolor(self.annotation_color), width=2.5, overlay=True)
        page.insert_textbox(
            stamp_rect,
            text.strip(),
            fontsize=34,
            fontname="helv",
            color=rgb_from_qcolor(self.annotation_color),
            align=1,
            overlay=True,
        )
        self._mark_dirty("Stamp added")

    def add_watermark(self):
        text, ok = QInputDialog.getText(self, "Watermark", "Watermark text:")
        if not ok or not text.strip():
            return

        # ValueError: bad rotate value. This version uses a large centered
        self._push_undo()
        for page in self.doc:
            rect = page.rect
            font_size = max(28, min(rect.width, rect.height) / 9)
            box = fitz.Rect(rect.width * 0.08, rect.height * 0.43, rect.width * 0.92, rect.height * 0.57)
            self._safe_insert_textbox(
                page,
                box,
                text.strip(),
                fontsize=font_size,
                color=QColor(155, 163, 175),
                align=1,
            )
        self._mark_dirty("Watermark added to all pages")

    def compress_pdf(self):
        """Save a size-reduced copy: garbage-collect, deflate streams, and
        optionally downsample large images."""
        if self.doc is None:
            return
        downsample = QMessageBox.question(
            self, "Reduce File Size",
            "Also downsample large images? (Smaller file, slightly lower image quality.)",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        ) == QMessageBox.StandardButton.Yes

        default = "compressed.pdf"
        if self.file_path:
            src = Path(self.file_path)
            default = str(src.with_name(f"{src.stem}_compressed.pdf"))
        out, _ = QFileDialog.getSaveFileName(self, "Save Reduced PDF As", default, "PDF Files (*.pdf)")
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"

        try:
            work = self.doc
            temp = None
            if downsample:
                # Work on a copy so the open document is untouched.
                temp = fitz.open("pdf", self.doc.tobytes())
                self._downsample_images(temp)
                work = temp

            work.save(out, garbage=4, deflate=True, deflate_images=True,
                      deflate_fonts=True, clean=True)
            if temp is not None:
                temp.close()

            before = os.path.getsize(self.file_path) if self.file_path and os.path.exists(self.file_path) else None
            after = os.path.getsize(out)
            msg = f"Saved reduced PDF:\n{out}\n\nNew size: {self._human_size(after)}"
            if before:
                pct = (1 - after / before) * 100 if before else 0
                msg += f"\nOriginal: {self._human_size(before)}  ({pct:.0f}% smaller)"
            self.status.showMessage(f"Reduced PDF saved: {self._human_size(after)}")
            QMessageBox.information(self, "Reduce File Size", msg)
        except Exception as exc:
            QMessageBox.critical(self, "Reduce File Size", f"Could not compress PDF:\n{exc}")

    @staticmethod
    def _human_size(n: int) -> str:
        size = float(n)
        for unit in ("B", "KB", "MB", "GB"):
            if size < 1024 or unit == "GB":
                return f"{size:.1f} {unit}" if unit != "B" else f"{int(size)} B"
            size /= 1024
        return f"{size:.1f} GB"

    def _downsample_images(self, doc, max_dim: int = 1600, quality: int = 70):
        """Re-encode oversized images as smaller JPEGs to cut file size."""
        seen = set()
        for page in doc:
            for info in page.get_images(full=True):
                xref = info[0]
                if xref in seen:
                    continue
                seen.add(xref)
                try:
                    pix = fitz.Pixmap(doc, xref)
                    if pix.width <= max_dim and pix.height <= max_dim:
                        continue
                    scale = max_dim / max(pix.width, pix.height)
                    new_w = max(1, int(pix.width * scale))
                    new_h = max(1, int(pix.height * scale))
                    if pix.n > 4:  # CMYK etc → convert to RGB
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    pix = fitz.Pixmap(pix, new_w, new_h)  # scale
                    img_bytes = pix.tobytes("jpeg", jpg_quality=quality)
                    doc.update_stream(xref, img_bytes)  # best-effort
                except Exception:
                    continue

    # ---- Header / footer & page numbering ---------------------------------
    def add_header_footer(self):
        if self.doc is None:
            return
        dialog = HeaderFooterDialog(self)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        cfg = dialog.values()
        if not any(cfg["slots"].values()):
            return

        from datetime import datetime
        date_str = datetime.now().strftime("%Y-%m-%d")
        total = self.doc.page_count
        margin = float(cfg["margin"])
        size = float(cfg["size"])
        color = QColor(cfg["color"])

        def expand(txt, page_no):
            return (txt.replace("{page}", str(page_no))
                       .replace("{pages}", str(total))
                       .replace("{date}", date_str))

        self._push_undo()
        for i in range(total):
            page = self.doc[i]
            r = page.rect
            y_top = margin
            y_bot = r.height - margin - size
            positions = {
                "hl": (r.x0 + margin, y_top, 0),
                "hc": (r.x0, y_top, 1),
                "hr": (r.x0, y_top, 2),
                "fl": (r.x0 + margin, y_bot, 0),
                "fc": (r.x0, y_bot, 1),
                "fr": (r.x0, y_bot, 2),
            }
            for key, text in cfg["slots"].items():
                if not text:
                    continue
                x, y, align = positions[key]
                content = expand(text, i + 1)
                box = fitz.Rect(margin, y, r.width - margin, y + size * 1.6)
                self._safe_insert_textbox(page, box, content, fontsize=size,
                                          color=color, align=align)
        self._mark_dirty("Header/footer added to all pages", refresh_sidebars=True)

    def add_bates_numbering(self):
        if self.doc is None:
            return
        prefix, ok = QInputDialog.getText(self, "Bates Numbering", "Prefix (optional):", text="")
        if not ok:
            return
        start, ok = QInputDialog.getInt(self, "Bates Numbering", "Start number:", 1, 0, 10_000_000)
        if not ok:
            return
        digits, ok = QInputDialog.getInt(self, "Bates Numbering", "Zero-padded digits:", 6, 1, 12)
        if not ok:
            return

        self._push_undo()
        size = 9.0
        for i in range(self.doc.page_count):
            page = self.doc[i]
            r = page.rect
            label = f"{prefix}{str(start + i).zfill(digits)}"
            box = fitz.Rect(r.width - 220, r.height - 26, r.width - 12, r.height - 8)
            self._safe_insert_textbox(page, box, label, fontsize=size,
                                      color=QColor(40, 40, 40), align=2)
        self._mark_dirty("Bates numbering applied", refresh_sidebars=True)

    # ---- Interactive form fields (AcroForm) -------------------------------
    def fill_form_fields(self):
        if self.doc is None:
            return
        if not getattr(self.doc, "is_form_pdf", False):
            QMessageBox.information(self, "Fill Form Fields",
                                    "This PDF has no interactive form fields.")
            return

        type_map = {
            fitz.PDF_WIDGET_TYPE_TEXT: "text",
            fitz.PDF_WIDGET_TYPE_CHECKBOX: "checkbox",
            fitz.PDF_WIDGET_TYPE_COMBOBOX: "combobox",
            fitz.PDF_WIDGET_TYPE_LISTBOX: "listbox",
            fitz.PDF_WIDGET_TYPE_RADIOBUTTON: "radio",
        }
        fields: List[Dict] = []
        for pi in range(self.doc.page_count):
            page = self.doc[pi]
            for wi, wdg in enumerate(page.widgets()):
                fields.append({
                    "page": pi,
                    "index": wi,
                    "name": wdg.field_name or "",
                    "type": type_map.get(wdg.field_type, "text"),
                    "value": wdg.field_value,
                    "choices": list(getattr(wdg, "choice_values", None) or []),
                })
        if not fields:
            QMessageBox.information(self, "Fill Form Fields",
                                    "This PDF has no fillable fields.")
            return

        dialog = FormFillDialog(self, fields)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        results = dialog.result_values()

        self._push_undo()
        changed = 0
        # Re-fetch widgets by page + index so we write to live objects.
        for pi in range(self.doc.page_count):
            page = self.doc[pi]
            widgets = list(page.widgets())
            for r in results:
                if r["page"] != pi or not (0 <= r["index"] < len(widgets)):
                    continue
                wdg = widgets[r["index"]]
                try:
                    if r["type"] == "checkbox":
                        wdg.field_value = bool(r["new_value"])
                    else:
                        wdg.field_value = str(r["new_value"])
                    wdg.update()
                    changed += 1
                except Exception:
                    continue
        self._mark_dirty(f"Form fields updated ({changed})", refresh_sidebars=True)

    def rotate_current_page_clockwise(self):
        self._rotate_current_page(90)

    def rotate_current_page_counterclockwise(self):
        self._rotate_current_page(-90)

    def _rotate_current_page(self, delta: int):
        if self.doc is None:
            return
        self._push_undo()
        page = self.doc[self.current_page_index]
        page.set_rotation((page.rotation + delta) % 360)
        self._mark_dirty("Page rotated")

    def delete_current_page(self):
        if self.doc is None or self.doc.page_count <= 1:
            return
        response = QMessageBox.question(
            self,
            "Delete Page",
            f"Delete page {self.current_page_index + 1}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if response != QMessageBox.StandardButton.Yes:
            return
        self._push_undo()
        self.doc.delete_page(self.current_page_index)
        self.current_page_index = min(self.current_page_index, self.doc.page_count - 1)
        self._mark_page_structure_changed()
        self._mark_dirty("Page deleted", refresh_sidebars=True)

    def duplicate_current_page(self):
        if self.doc is None:
            return

        try:
            self._push_undo()

            # PyMuPDF cannot insert pages from the same open document object
            temp_doc = fitz.open()
            temp_doc.insert_pdf(
                self.doc,
                from_page=self.current_page_index,
                to_page=self.current_page_index,
            )
            self.doc.insert_pdf(
                temp_doc,
                from_page=0,
                to_page=0,
                start_at=self.current_page_index + 1,
            )
            temp_doc.close()

            self.current_page_index += 1
            self._mark_page_structure_changed()
            self._mark_dirty("Page duplicated", refresh_sidebars=True)
        except Exception as exc:
            QMessageBox.critical(self, "Duplicate Page Error", f"Could not duplicate page:\n{exc}")

    def _on_pages_reordered(self, parent, start, end, dest_parent, dest_row):
        """React to a drag-reorder in the Pages thumbnail list by reordering the
        document to match, using an explicit page-order list (robust)."""
        if self.doc is None or getattr(self, "_suppress_page_reorder", False):
            return
        n = self.doc.page_count
        src = int(start)
        dst = int(dest_row)
        # Qt's destination is the insert row in the pre-removal list; when moving
        # downward the effective target shifts left by one after removal.
        if dst > src:
            dst -= 1
        if not (0 <= src < n):
            return
        dst = max(0, min(dst, n - 1))
        if src == dst:
            return
        order = list(range(n))
        order.insert(dst, order.pop(src))
        try:
            self._push_undo()
            self.doc.select(order)
            self.current_page_index = dst
            self._mark_page_structure_changed()
            self._mark_dirty(f"Page {src + 1} moved to position {dst + 1}", refresh_sidebars=True)
        except Exception as exc:
            QMessageBox.warning(self, "Reorder Pages", f"Could not reorder pages:\n{exc}")
            self.refresh_page_thumbnails()

    def move_page_up(self):
        if self.doc is None or self.current_page_index <= 0:
            return
        self.doc.move_page(self.current_page_index, self.current_page_index - 1)
        self.current_page_index -= 1
        self._mark_page_structure_changed()
        self._mark_dirty("Page moved up", refresh_sidebars=True)

    def move_page_down(self):
        if self.doc is None or self.current_page_index >= self.doc.page_count - 1:
            return
        self._push_undo()
        # no-op because the page is already right before index+1.
        target = self.current_page_index + 2
        if target >= self.doc.page_count:
            self.doc.move_page(self.current_page_index, -1)
        else:
            self.doc.move_page(self.current_page_index, target)
        self.current_page_index += 1
        self._mark_page_structure_changed()
        self._mark_dirty("Page moved down", refresh_sidebars=True)

    def insert_blank_page(self):
        if self.doc is None:
            return
        self._push_undo()
        current = self.doc[self.current_page_index]
        self.doc.new_page(pno=self.current_page_index + 1, width=current.rect.width, height=current.rect.height)
        self.current_page_index += 1
        self._mark_page_structure_changed()
        self._mark_dirty("Blank page inserted", refresh_sidebars=True)

    def insert_pdf_after_current(self):
        if self.doc is None:
            return
        path, _ = QFileDialog.getOpenFileName(self, "Merge PDF After Current Page", "", "PDF Files (*.pdf)")
        if not path:
            return
        try:
            src = fitz.open(path)
            if src.page_count == 0:
                raise ValueError("Selected PDF has no pages.")
            self._push_undo()
            self.doc.insert_pdf(src, start_at=self.current_page_index + 1)
            src.close()
            self._mark_page_structure_changed()
            self._mark_dirty("PDF merged", refresh_sidebars=True)
        except Exception as exc:
            QMessageBox.critical(self, "Merge Error", f"Could not merge PDF:\n{exc}")

    def insert_pdf_at(self, before: bool = False):
        """Insert another PDF's pages before/after the current page."""
        if self.doc is None:
            return
        where = "Before" if before else "After"
        path, _ = QFileDialog.getOpenFileName(self, f"Insert PDF {where} Current Page", "", "PDF Files (*.pdf)")
        if not path:
            return
        try:
            src = fitz.open(path)
            if src.page_count == 0:
                raise ValueError("Selected PDF has no pages.")
            start_at = self.current_page_index if before else self.current_page_index + 1
            self._push_undo()
            self.doc.insert_pdf(src, start_at=start_at)
            src.close()
            self.current_page_index = start_at
            self._mark_page_structure_changed()
            self._mark_dirty(f"Inserted {where.lower()} page {self.current_page_index}", refresh_sidebars=True)
        except Exception as exc:
            QMessageBox.critical(self, "Insert Error", f"Could not insert PDF:\n{exc}")

    def extract_page_range(self):
        if self.doc is None:
            return
        rng = page_range_from_user(self, self.doc.page_count)
        if not rng:
            return
        start, end = rng

        out, _ = QFileDialog.getSaveFileName(self, "Extract Page Range", f"pages_{start+1}_{end+1}.pdf", "PDF Files (*.pdf)")
        if not out:
            return
        if not out.lower().endswith(".pdf"):
            out += ".pdf"

        try:
            new_doc = fitz.open()
            new_doc.insert_pdf(self.doc, from_page=start, to_page=end)
            new_doc.save(out, garbage=4, deflate=True)
            new_doc.close()
            QMessageBox.information(self, "Extracted", f"Saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "Extract Error", f"Could not extract pages:\n{exc}")

    def split_every_page(self):
        if self.doc is None:
            return
        folder = QFileDialog.getExistingDirectory(self, "Choose Folder for Split Pages")
        if not folder:
            return
        try:
            base = safe_filename(Path(self.file_path).stem if self.file_path else "document")
            for i in range(self.doc.page_count):
                new_doc = fitz.open()
                new_doc.insert_pdf(self.doc, from_page=i, to_page=i)
                out = Path(folder) / f"{base}_page_{i + 1:03d}.pdf"
                new_doc.save(str(out), garbage=4, deflate=True)
                new_doc.close()
            QMessageBox.information(self, "Split Complete", f"Saved {self.doc.page_count} PDF files.")
        except Exception as exc:
            QMessageBox.critical(self, "Split Error", f"Could not split document:\n{exc}")

    def search_text(self):
        """Search in a worker thread so big documents stay responsive."""
        if self.doc is None:
            return
        query = self.search_input.text().strip()
        if not query:
            return

        self._search_job_id += 1
        job_id = self._search_job_id
        self.search_results.clear()
        self.search_results_list.clear()
        self.search_index = -1
        self.render_current_page()

        source = self._document_worker_source()
        if not source:
            return

        worker = SearchWorker(job_id, source, query)
        worker.progress.connect(lambda jid, done, total: self._worker_progress(jid, done, total, "Searching"))
        worker.finished_results.connect(self._search_finished_from_worker)
        worker.error.connect(lambda jid, msg: QMessageBox.critical(self, "Search Error", f"Search failed:\n{msg}"))
        worker.finished.connect(lambda jid: self._worker_finished(jid, "Search complete"))
        self._start_worker(worker)
        self.sidebar.setCurrentIndex(2)
        self.status.showMessage(f"Searching for '{query}'...")

    def go_to_search_result(self, index: int):
        if not self.search_results or self.doc is None:
            return
        self.search_index = index % len(self.search_results)
        page_index, _ = self.search_results[self.search_index]
        self.current_page_index = page_index
        self.search_results_list.setCurrentRow(self.search_index)
        self.render_current_page()

    def go_to_next_search_result(self):
        if not self.search_results:
            self.search_text()
            return
        self.go_to_search_result(self.search_index + 1)

    def go_to_search_item(self, item: QListWidgetItem):
        index = item.data(Qt.ItemDataRole.UserRole)
        if isinstance(index, int):
            self.go_to_search_result(index)

    def export_page_image(self, fmt: str):
        if self.doc is None:
            return
        ext = "jpg" if fmt.lower() in {"jpg", "jpeg"} else "png"
        out, _ = QFileDialog.getSaveFileName(self, f"Export Current Page as {ext.upper()}", f"page_{self.current_page_index+1}.{ext}", f"{ext.upper()} Image (*.{ext})")
        if not out:
            return
        if not out.lower().endswith("." + ext):
            out += "." + ext
        try:
            pix = self.doc[self.current_page_index].get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            pix.save(out)
            QMessageBox.information(self, "Exported", f"Saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", f"Could not export image:\n{exc}")

    def export_all_pages_images(self, fmt: str):
        """Export all pages in a worker thread to avoid locking the app."""
        if self.doc is None:
            return
        folder = QFileDialog.getExistingDirectory(self, f"Choose Folder for {fmt.upper()} Export")
        if not folder:
            return
        ext = "jpg" if fmt.lower() in {"jpg", "jpeg"} else "png"
        base = safe_filename(Path(self.file_path).stem if self.file_path else "document")
        source = self._document_worker_source()
        if not source:
            return

        self._export_job_id += 1
        job_id = self._export_job_id
        worker = ImageExportWorker(job_id, source, folder, base, ext, scale=2.0)
        worker.progress.connect(lambda jid, done, total: self._worker_progress(jid, done, total, f"Exporting {ext.upper()}"))
        worker.done.connect(self._image_export_finished_from_worker)
        worker.error.connect(lambda jid, msg: QMessageBox.critical(self, "Export Error", f"Could not export images:\n{msg}"))
        worker.finished.connect(lambda jid: self._worker_finished(jid, "Export complete"))
        self._start_worker(worker)
        self.status.showMessage(f"Exporting all pages as {ext.upper()}...")

    def export_current_page_text(self):
        if self.doc is None:
            return
        out, _ = QFileDialog.getSaveFileName(self, "Export Current Page Text", f"page_{self.current_page_index+1}.txt", "Text Files (*.txt)")
        if not out:
            return
        text = self.doc[self.current_page_index].get_text("text", sort=True)
        self._write_text_file(out, text)

    def export_full_document_text(self):
        if self.doc is None:
            return
        out, _ = QFileDialog.getSaveFileName(self, "Export Full Text", "document_text.txt", "Text Files (*.txt)")
        if not out:
            return
        text = "\n\n".join(f"--- Page {i+1} ---\n{p.get_text('text', sort=True)}" for i, p in enumerate(self.doc))
        self._write_text_file(out, text)

    def export_full_document_html(self):
        if self.doc is None:
            return
        out, _ = QFileDialog.getSaveFileName(self, "Export HTML", "document.html", "HTML Files (*.html)")
        if not out:
            return
        if not out.lower().endswith(".html"):
            out += ".html"
        body = []
        for i, page in enumerate(self.doc):
            text = html.escape(page.get_text("text", sort=True))
            body.append(f"<h2>Page {i+1}</h2><pre>{text}</pre>")
        content = "<!doctype html><html><head><meta charset='utf-8'><title>PDF Export</title></head><body>" + "\n".join(body) + "</body></html>"
        self._write_text_file(out, content)

    def export_full_document_markdown(self):
        if self.doc is None:
            return
        out, _ = QFileDialog.getSaveFileName(self, "Export Markdown", "document.md", "Markdown Files (*.md)")
        if not out:
            return
        if not out.lower().endswith(".md"):
            out += ".md"
        text = "\n\n".join(f"## Page {i+1}\n\n{p.get_text('text', sort=True)}" for i, p in enumerate(self.doc))
        self._write_text_file(out, text)

    def export_full_document_docx(self):
        if self.doc is None:
            return
        try:
            from docx import Document
        except ImportError:
            QMessageBox.warning(self, "Missing Dependency", "Install DOCX export support with:\n\npip install python-docx")
            return

        out, _ = QFileDialog.getSaveFileName(self, "Export DOCX", "document.docx", "Word Document (*.docx)")
        if not out:
            return
        if not out.lower().endswith(".docx"):
            out += ".docx"

        try:
            docx = Document()
            docx.add_heading(Path(self.file_path).name if self.file_path else "PDF Export", 0)
            for i, page in enumerate(self.doc):
                docx.add_heading(f"Page {i+1}", level=1)
                for para in page.get_text("text", sort=True).splitlines():
                    if para.strip():
                        docx.add_paragraph(para.strip())
                if i != self.doc.page_count - 1:
                    docx.add_page_break()
            docx.save(out)
            QMessageBox.information(self, "Exported", f"DOCX saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "DOCX Export Error", f"Could not export DOCX:\n{exc}")

    def export_metadata_json(self):
        if self.doc is None:
            return
        out, _ = QFileDialog.getSaveFileName(self, "Export Metadata JSON", "metadata.json", "JSON Files (*.json)")
        if not out:
            return
        if not out.lower().endswith(".json"):
            out += ".json"
        data = {"page_count": self.doc.page_count, "metadata": self.doc.metadata}
        self._write_text_file(out, json.dumps(data, indent=2, ensure_ascii=False))

    def _write_text_file(self, out: str, text: str):
        try:
            Path(out).write_text(text, encoding="utf-8")
            QMessageBox.information(self, "Exported", f"Saved:\n{out}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", f"Could not save file:\n{exc}")

    def show_document_info(self):
        if self.doc is None:
            return
        DocumentInfoDialog(self, self.doc.metadata or {}, self.doc.page_count).exec()

    def edit_metadata(self):
        if self.doc is None:
            return
        dialog = MetadataEditDialog(self, self.doc.metadata or {})
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        self._push_undo()
        self.doc.set_metadata(dialog.metadata())
        self._mark_dirty("Metadata updated")

    def zoom_in(self):
        if self.doc is None:
            return
        self.zoom = min(5.0, self.zoom + 0.15)
        self.render_current_page()

    def zoom_out(self):
        if self.doc is None:
            return
        self.zoom = max(0.25, self.zoom - 0.15)
        self.render_current_page()

    def fit_width(self):
        if self.doc is None:
            return
        page = self.doc[self.current_page_index]
        available = max(self.scroll.viewport().width() - 90, 200)
        self.zoom = clamp(available / page.rect.width, 0.25, 5.0)
        self.render_current_page()

    def fit_page(self):
        if self.doc is None:
            return
        page = self.doc[self.current_page_index]
        available_w = max(self.scroll.viewport().width() - 90, 200)
        available_h = max(self.scroll.viewport().height() - 90, 200)
        self.zoom = clamp(min(available_w / page.rect.width, available_h / page.rect.height), 0.25, 5.0)
        self.render_current_page()

    def handle_zoom_combo(self, text: str):
        if self.doc is None:
            return
        if text == "Fit Width":
            self.fit_width()
        elif text == "Fit Page":
            self.fit_page()
        elif text.endswith("%"):
            try:
                self.zoom = clamp(int(text[:-1]) / 100, 0.25, 5.0)
                self.render_current_page()
            except ValueError:
                pass

    def _update_zoom_combo_text(self):
        text = f"{int(round(self.zoom * 100))}%"
        presets = {
            "50%", "75%", "100%", "125%", "150%",
            "175%", "200%", "300%", "Fit Width", "Fit Page",
        }
        self.zoom_combo.blockSignals(True)
        # dropdown does not fill up with fractional percentages on wheel zoom.
        for i in range(self.zoom_combo.count() - 1, -1, -1):
            if self.zoom_combo.itemText(i) not in presets:
                self.zoom_combo.removeItem(i)
        if self.zoom_combo.findText(text) < 0:
            self.zoom_combo.insertItem(0, text)
        self.zoom_combo.setCurrentText(text)
        self.zoom_combo.blockSignals(False)

    def _push_undo(self):
        if self.doc is None:
            return
        try:
            self.undo_stack.append(self.doc.tobytes(garbage=4, deflate=True))
            if len(self.undo_stack) > self.max_history:
                self.undo_stack.pop(0)
            self.redo_stack.clear()
            self._refresh_history_actions()
        except Exception:
            pass

    def undo(self):
        if self.doc is None or not self.undo_stack:
            return
        try:
            self.redo_stack.append(self.doc.tobytes(garbage=4, deflate=True))
            data = self.undo_stack.pop()
            self.doc.close()
            self.doc = fitz.open(stream=data, filetype="pdf")
            self._doc_version += 1
            self._render_cache.clear()
            self._render_cache_order.clear()
            self.current_page_index = min(self.current_page_index, self.doc.page_count - 1)
            self.is_dirty = True
            self.refresh_sidebars()
            self.render_current_page()
            self._update_window_title()
            self._refresh_history_actions()
        except Exception as exc:
            QMessageBox.critical(self, "Undo Error", f"Could not undo:\n{exc}")

    def redo(self):
        if self.doc is None or not self.redo_stack:
            return
        try:
            self.undo_stack.append(self.doc.tobytes(garbage=4, deflate=True))
            data = self.redo_stack.pop()
            self.doc.close()
            self.doc = fitz.open(stream=data, filetype="pdf")
            self._doc_version += 1
            self._render_cache.clear()
            self._render_cache_order.clear()
            self.current_page_index = min(self.current_page_index, self.doc.page_count - 1)
            self.is_dirty = True
            self.refresh_sidebars()
            self.render_current_page()
            self._update_window_title()
            self._refresh_history_actions()
        except Exception as exc:
            QMessageBox.critical(self, "Redo Error", f"Could not redo:\n{exc}")

    def _refresh_history_actions(self):
        if hasattr(self, "undo_action"):
            self.undo_action.setEnabled(bool(self.undo_stack) and self.doc is not None)
        if hasattr(self, "redo_action"):
            self.redo_action.setEnabled(bool(self.redo_stack) and self.doc is not None)

    def _mark_dirty(self, message: str, refresh_sidebars: bool = False):
        self.is_dirty = True
        self._doc_version += 1
        self._render_cache.clear()
        self._render_cache_order.clear()
        self._xray_cache.clear()

        if self._page_structure_changed:
            self.refresh_sidebars()
            self._page_structure_changed = False
        else:
            self._refresh_light_sidebars()
            self._update_current_thumbnail_inline()

        self.render_current_page()
        self._update_window_title()
        self._refresh_history_actions()
        self.status.showMessage(message)

    def _update_window_title(self):
        name = Path(self.file_path).name if self.file_path else "Untitled"
        mark = " •" if self.is_dirty else ""
        self.setWindowTitle(f"Suketchi — {name}{mark}")
        # Keep the active tab's dirty marker in sync with the title.
        if getattr(self, "tab_bar", None) is not None and self.tabs:
            self._refresh_tab_bar()

    def _confirm_discard(self) -> bool:
        response = QMessageBox.question(
            self,
            "Unsaved Changes",
            "You have unsaved changes. Continue without saving?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        return response == QMessageBox.StandardButton.Yes

    def _image_point_to_pdf_xy(self, point: QPointF) -> Tuple[float, float]:
        return point.x() / self.zoom, point.y() / self.zoom

    def _image_rect_to_pdf_rect(self, rect: QRectF) -> fitz.Rect:
        x0 = rect.left() / self.zoom
        y0 = rect.top() / self.zoom
        x1 = rect.right() / self.zoom
        y1 = rect.bottom() / self.zoom
        return fitz.Rect(min(x0, x1), min(y0, y1), max(x0, x1), max(y0, y1))

    def _pdf_rect_to_image_rect(self, rect: fitz.Rect) -> QRectF:
        return QRectF(rect.x0 * self.zoom, rect.y0 * self.zoom, rect.width * self.zoom, rect.height * self.zoom)

    @staticmethod
    def _pdf_from_mime(mime) -> Optional[str]:
        if not mime.hasUrls():
            return None
        for url in mime.urls():
            if url.isLocalFile() and url.toLocalFile().lower().endswith(".pdf"):
                return url.toLocalFile()
        return None

    def dragEnterEvent(self, event):
        if self._pdf_from_mime(event.mimeData()):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if self._pdf_from_mime(event.mimeData()):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dropEvent(self, event):
        path = self._pdf_from_mime(event.mimeData())
        if not path:
            event.ignore()
            return
        event.acceptProposedAction()
        # Dropping a PDF opens it in a new tab; nothing is discarded.
        self.load_pdf(path)

    def createPopupMenu(self):
        # Suppress QMainWindow's default toolbar/dock context menu so a right
        # click on the ribbon (e.g. the X-ray button) can never hide it.
        return None

    def closeEvent(self, event):
        # Confirm once if ANY open tab has unsaved changes.
        self._snapshot_active_tab()
        if any(t.is_dirty for t in self.tabs) or self.is_dirty:
            if not self._confirm_discard():
                event.ignore()
                return
        for thread, worker in list(self._workers):
            try:
                worker.cancelled = True
            except Exception:
                pass
            thread.quit()
            thread.wait(2000)
        self._workers.clear()
        for tab in self.tabs:
            try:
                if tab.doc is not None:
                    tab.doc.close()
            except Exception:
                pass
        self.tabs.clear()
        if self.doc:
            try:
                self.doc.close()
            except Exception:
                pass
        event.accept()


def load_hand_font() -> str:
    """Load the bundled handwriting font, else fall back gracefully.

    Looks for PatrickHand-Regular.ttf in ./fonts (or ./assets, or next to the
    script). If it is missing, tries handwriting faces that ship with common
    systems, and finally settles for a normal sans so the app still runs
    everywhere.
    """
    global APP_FONT_FAMILY

    here = Path(__file__).resolve().parent
    for candidate in (
        here / "fonts" / "PatrickHand-Regular.ttf",
        here / "assets" / "fonts" / "PatrickHand-Regular.ttf",
        here / "PatrickHand-Regular.ttf",
        here / "assets" / "PatrickHand-Regular.ttf",
    ):
        if candidate.exists():
            font_id = QFontDatabase.addApplicationFont(str(candidate))
            families = QFontDatabase.applicationFontFamilies(font_id)
            if families:
                APP_FONT_FAMILY = families[0]
                return APP_FONT_FAMILY

    available = set(QFontDatabase.families())
    for name in ("Patrick Hand", "Comic Neue", "Segoe Print", "Bradley Hand",
                 "Chalkboard SE", "Comic Sans MS", "Ink Free"):
        if name in available:
            APP_FONT_FAMILY = name
            return APP_FONT_FAMILY

    APP_FONT_FAMILY = "Segoe UI"
    return APP_FONT_FAMILY


# Bundled signature typefaces (family name -> file), loaded on demand. These
# are open-licensed (SIL OFL) script fonts shipped in ./fonts.
SIGNATURE_FONT_FILES = [
    ("Great Vibes", "GreatVibes-Regular.ttf"),
    ("Allura", "Allura-Regular.ttf"),
    ("Dancing Script", "DancingScript-Regular.ttf"),
    ("Sacramento", "Sacramento-Regular.ttf"),
    ("Caveat", "Caveat-Regular.ttf"),
]

_SIGNATURE_FONTS_LOADED: List[str] = []
_SIGNATURE_FONT_PATHS: Dict[str, str] = {}


def load_signature_fonts() -> List[str]:
    """Register the bundled signature fonts and return the available family
    names (falling back to the app handwriting font if none are found)."""
    global _SIGNATURE_FONTS_LOADED
    if _SIGNATURE_FONTS_LOADED:
        return _SIGNATURE_FONTS_LOADED

    here = Path(__file__).resolve().parent
    families: List[str] = []
    for expected_family, filename in SIGNATURE_FONT_FILES:
        for base in (here / "fonts", here / "assets" / "fonts", here):
            path = base / filename
            if path.exists():
                font_id = QFontDatabase.addApplicationFont(str(path))
                fams = QFontDatabase.applicationFontFamilies(font_id)
                if fams:
                    families.append(fams[0])
                    # Remember the file for each reported family so we can embed
                    # exactly the right font into the PDF later.
                    _SIGNATURE_FONT_PATHS[fams[0]] = str(path)
                break

    if not families:
        # No bundled script fonts available: fall back to the app hand font so
        # the feature still works everywhere.
        families = [APP_FONT_FAMILY]

    _SIGNATURE_FONTS_LOADED = families
    return families


def signature_font_file(family: str) -> Optional[str]:
    """Return the on-disk path of the bundled signature font matching a family
    name, so it can be embedded in a PDF via insert_textbox(fontfile=...)."""
    if not _SIGNATURE_FONT_PATHS:
        load_signature_fonts()
    if family in _SIGNATURE_FONT_PATHS:
        return _SIGNATURE_FONT_PATHS[family]
    # Fallback: first available signature font file.
    if _SIGNATURE_FONT_PATHS:
        return next(iter(_SIGNATURE_FONT_PATHS.values()))
    return None


# ---- Spell checking --------------------------------------------------------
_SPELL_CHECKER = None          # a callable(word)->bool ("is this word known?")
_SPELL_READY = False
_WORD_RE = re.compile(r"[A-Za-z][A-Za-z'’]*")


def _load_spell_checker():
    """Lazily build a spell checker. Prefers the `pyspellchecker` package
    (offline dictionary); falls back to a system word list; if neither is
    available, spell checking is disabled gracefully."""
    global _SPELL_CHECKER, _SPELL_READY
    if _SPELL_READY:
        return _SPELL_CHECKER
    _SPELL_READY = True

    # 1. pyspellchecker (bundled dictionary, best quality)
    try:
        from spellchecker import SpellChecker
        sc = SpellChecker(distance=1)

        def _known(word: str) -> bool:
            return len(sc.unknown([word])) == 0

        _SPELL_CHECKER = _known
        return _SPELL_CHECKER
    except Exception:
        pass

    # 2. System word list fallback.
    for path in ("/usr/share/dict/words", "/usr/share/dict/american-english",
                 "/usr/dict/words"):
        try:
            if os.path.exists(path):
                with open(path, "r", errors="ignore") as fh:
                    words = {w.strip().lower() for w in fh if w.strip()}
                if words:
                    def _known(word: str, _words=words) -> bool:
                        return word.lower() in _words
                    _SPELL_CHECKER = _known
                    return _SPELL_CHECKER
        except Exception:
            continue

    _SPELL_CHECKER = None
    return None


def spell_checker_available() -> bool:
    return _load_spell_checker() is not None


# ---- Cryptographic digital signatures (pyHanko) ---------------------------
def pyhanko_available() -> bool:
    try:
        import pyhanko  # noqa: F401
        from pyhanko.sign import signers  # noqa: F401
        return True
    except Exception:
        return False


def is_word_misspelled(word: str) -> bool:
    """True if the token looks like an ordinary word that the dictionary does
    not recognise. Tokens with digits, all-caps acronyms, single letters, or
    mixed-case (e.g. camelCase / product names) are skipped to avoid noise."""
    checker = _load_spell_checker()
    if checker is None:
        return False
    core = word.strip("'’")
    if len(core) < 3:
        return False
    if any(ch.isdigit() for ch in core):
        return False
    # Skip acronyms / all-caps and odd mixed-caps tokens.
    if core.isupper():
        return False
    letters = core.replace("'", "").replace("’", "")
    if not letters.isalpha():
        return False
    try:
        return not checker(core.lower())
    except Exception:
        return False


def font_file_for_pdf_font(pdf_font_name: str) -> Optional[str]:
    """Match a font name found inside a PDF (e.g. 'GreatVibes-Regular', possibly
    with a subset prefix like 'ABCDEF+GreatVibes-Regular') back to a bundled
    signature font file, so a moved signature keeps its original typeface."""
    if not pdf_font_name:
        return None
    load_signature_fonts()
    # Strip a subset prefix ("ABCDEF+Name") if present.
    name = pdf_font_name.split("+", 1)[-1]
    name_l = name.lower().replace(" ", "").replace("-", "")
    here = Path(__file__).resolve().parent
    for _family, filename in SIGNATURE_FONT_FILES:
        stem = Path(filename).stem.lower().replace(" ", "").replace("-", "")
        if name_l == stem or name_l in stem or stem in name_l:
            for base in (here / "fonts", here / "assets" / "fonts", here):
                path = base / filename
                if path.exists():
                    return str(path)
    # Also try matching against the loaded family names.
    for family, path in _SIGNATURE_FONT_PATHS.items():
        fam_l = family.lower().replace(" ", "").replace("-", "")
        if name_l == fam_l or name_l in fam_l or fam_l in name_l:
            return path
    return None


def asset_path(*candidates) -> Optional[Path]:
    """Find a bundled asset next to the script, tolerating a few layouts."""
    here = Path(__file__).resolve().parent
    for rel in candidates:
        for base in (here, here / "logo", here / "assets"):
            path = base / rel
            if path.exists():
                return path
    return None


def render_svg(path: Path, size: int) -> Optional[QPixmap]:
    """Render an SVG crisply at an exact pixel size."""
    try:
        from PyQt6.QtSvg import QSvgRenderer
    except ImportError:
        return None
    renderer = QSvgRenderer(str(path))
    if not renderer.isValid():
        return None
    pixmap = QPixmap(size, size)
    pixmap.fill(Qt.GlobalColor.transparent)
    painter = QPainter(pixmap)
    painter.setRenderHint(QPainter.RenderHint.Antialiasing, True)
    renderer.render(painter)
    painter.end()
    return pixmap


def load_app_icon() -> QIcon:
    """Window / taskbar / launcher icon: the page mark.

    Uses the pre-rendered PNG set when present (its 16-32px sizes use the
    simplified mark, which survives where fine detail blurs). Otherwise renders
    the SVG at each size. Returns an empty icon if no assets exist at all, so
    the app still runs.
    """
    icon = QIcon()
    found = False
    for size in (16, 24, 32, 48, 64, 128, 256, 512):
        png = asset_path(f"png/icon_{size}.png", f"icon_{size}.png")
        if png:
            icon.addFile(str(png), QSize(size, size))
            found = True
    if found:
        return icon

    detailed = asset_path("suketchi_mark.svg")
    simplified = asset_path("suketchi_mark_small.svg")
    for size in (16, 24, 32, 48, 64, 128, 256, 512):
        source = simplified if (size <= 32 and simplified) else detailed
        if source:
            pixmap = render_svg(source, size)
            if pixmap:
                icon.addPixmap(pixmap)
    return icon


def load_seal_pixmap(size: int) -> Optional[QPixmap]:
    """The hanko seal, used as the small brand badge in the toolbar.

    The seal is a solid stamped square, so it stays legible at 20-30px where
    the detailed page mark would turn to grey mush.
    """
    svg = asset_path("suketchi_seal.svg")
    if svg:
        pixmap = render_svg(svg, size)
        if pixmap:
            return pixmap
    return load_brand_pixmap(size)


def load_brand_pixmap(size: int) -> Optional[QPixmap]:
    """The page mark, used large (empty state, about box)."""
    svg = None
    if size <= 40:
        svg = asset_path("suketchi_mark_small.svg")
    if svg is None:
        svg = asset_path("suketchi_mark.svg")
    if svg:
        pixmap = render_svg(svg, size)
        if pixmap:
            return pixmap
    png = asset_path("png/icon_256.png", "icon_256.png")
    if png:
        pixmap = QPixmap(str(png))
        if not pixmap.isNull():
            return pixmap.scaled(size, size, Qt.AspectRatioMode.KeepAspectRatio,
                                 Qt.TransformationMode.SmoothTransformation)
    return None


def _apply_bw_palette(app: QApplication):
    """Pure black-on-white palette, so native dialogs match the sketch look."""
    p = QPalette()
    p.setColor(QPalette.ColorRole.Window, PAPER)
    p.setColor(QPalette.ColorRole.WindowText, INK)
    p.setColor(QPalette.ColorRole.Base, PAPER)
    p.setColor(QPalette.ColorRole.AlternateBase, PAPER)
    p.setColor(QPalette.ColorRole.Text, INK)
    p.setColor(QPalette.ColorRole.Button, PAPER)
    p.setColor(QPalette.ColorRole.ButtonText, INK)
    p.setColor(QPalette.ColorRole.ToolTipBase, PAPER)
    p.setColor(QPalette.ColorRole.ToolTipText, INK)
    p.setColor(QPalette.ColorRole.Highlight, INK)
    p.setColor(QPalette.ColorRole.HighlightedText, PAPER)
    p.setColor(QPalette.ColorRole.PlaceholderText, GREY_FAINT)
    for role in (QPalette.ColorRole.Text, QPalette.ColorRole.ButtonText, QPalette.ColorRole.WindowText):
        p.setColor(QPalette.ColorGroup.Disabled, role, GREY_FAINT)
    app.setPalette(p)


_apply_light_palette = _apply_bw_palette


def main():
    # High-DPI adaptation must be configured before the QApplication exists.
    # PassThrough keeps fractional scaling exact so the UI stays crisp on any
    # display resolution / scale factor (100%, 125%, 150%, 200%, ...).
    try:
        QApplication.setHighDpiScaleFactorRoundingPolicy(
            Qt.HighDpiScaleFactorRoundingPolicy.PassThrough
        )
    except Exception:
        pass

    app = QApplication(sys.argv)
    app.setApplicationName("Suketchi PDF Reader")
    app.setStyle(SketchStyle("Fusion"))
    family = load_hand_font()
    load_signature_fonts()

    # Scale the base font relative to the primary screen's DPI so text is
    # readable on both low-DPI and high-DPI monitors.
    base_point_size = 11
    screen = QGuiApplication.primaryScreen()
    if screen is not None:
        dpi = screen.logicalDotsPerInch()
        if dpi > 0:
            scaled = round(base_point_size * (dpi / 96.0))
            base_point_size = max(9, min(scaled, 18))
    app.setFont(QFont(family, base_point_size))

    _apply_bw_palette(app)
    icon = load_app_icon()
    if not icon.isNull():
        app.setWindowIcon(icon)
    window = PdfStudioOverhaulPro()
    if not icon.isNull():
        window.setWindowIcon(icon)
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()

